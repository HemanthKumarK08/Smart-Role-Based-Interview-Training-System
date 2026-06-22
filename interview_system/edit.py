import os
import sys
import tkinter as tk
import tkinter.messagebox as messagebox
from tkinter import ttk, scrolledtext, filedialog, simpledialog
import customtkinter as ctk

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from database import init_db
import auth
import admin
import student
import interview_engine
import feedback_engine

#---------------------------------------------------------------------------
# Design tokens
#---------------------------------------------------------------------------
FONT  = "Inter"
SZ    = dict(xs=10, sm=11, md=13, lg=15, xl=18, xxl=24, hero=32)

SB_BG      = "#0f0f11"
SB_ACTIVE  = "#1e1e2e"
SB_ACCENT  = "#6366f1"
SB_TEXT    = "#e2e2e8"
SB_MUTED   = "#6b6b80"
SB_HOVER   = "#1a1a2a"
SB_BORDER  = "#1e1e2e"

CONTENT_BG   = ("#f4f5f7", "#16161e")
CARD_BG      = ("#ffffff",  "#1e1e2e")
CARD_BORDER  = ("#e2e4ea",  "#2a2a3e")
ACCENT       = "#6366f1"
SUCCESS      = "#22c55e"
TEAL         = "#14b8a6"
PINK         = "#ec4899"
AMBER        = "#f59e0b"
DANGER       = "#ef4444"
TEXT_PRIMARY = ("#111118", "#e2e2e8")
TEXT_MUTED   = ("#6b6b80", "#6b6b80")

STAT_PALETTES = [
    ("#ede9fe", "#6d28d9", "#4c1d95"),
    ("#d1fae5", "#059669", "#064e3b"),
    ("#fef3c7", "#d97706", "#78350f"),
]

#---------------------------------------------------------------------------
# Helpers
#---------------------------------------------------------------------------
def F(size_key, weight="normal"):
    return ctk.CTkFont(family=FONT, size=SZ[size_key], weight=weight)


def _darken(hex_color, factor=0.85):
    if isinstance(hex_color, tuple):
        hex_color = hex_color[-1]
    if not isinstance(hex_color, str):
        return hex_color
    if not hex_color.startswith("#"):
        return hex_color
    hex_color = hex_color.lstrip("#")
    if len(hex_color) != 6:
        return "#" + hex_color
    try:
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        r = max(0, min(255, int(r * factor)))
        g = max(0, min(255, int(g * factor)))
        b = max(0, min(255, int(b * factor)))
        return f"#{r:02x}{g:02x}{b:02x}"
    except ValueError:
        return "#" + hex_color


def _style_tree():
    style = ttk.Style()
    try:
        style.theme_use("clam")
    except Exception:
        pass
    style.configure("Treeview", rowheight=30, font=(FONT, 11), borderwidth=0)
    style.configure("Treeview.Heading", font=(FONT, 11, "bold"),
                    background="#6366f1", foreground="white")
    style.map("Treeview", background=[("selected", "#4338ca")])


#---------------------------------------------------------------------------
# Application
#---------------------------------------------------------------------------
class InterviewApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Smart Interview Prep")
        self.geometry("1400x860")
        self.minsize(1200, 750)

        self._sidebar_frame   = None
        self._content_frame   = None
        self._sidebar_btns    = {}

        self._selected_role_id    = None
        self._selected_difficulty = None
        self._last_result_id      = None

        self._show_login_home()

    # ------------------------------------------------------------------ #
    #  Low-level widget factories                                        #
    # ------------------------------------------------------------------ #

    def _lbl(self, parent, text, size="md", weight="normal",
             color=None, anchor="w", **kw):
        return ctk.CTkLabel(
            parent, text=text, font=F(size, weight),
            text_color=color or TEXT_PRIMARY,
            anchor=anchor, **kw,
        )

    def _entry(self, parent, show=None, width=320, placeholder=""):
        opts = dict(width=width, font=F("md"),
                    border_color=CARD_BORDER, corner_radius=8)
        if show:        opts["show"]             = show
        if placeholder: opts["placeholder_text"] = placeholder
        return ctk.CTkEntry(parent, **opts)

    def _btn(self, parent, text, command, color=ACCENT,
             hover=None, width=140, height=38, **kw):
        return ctk.CTkButton(
            parent, text=text, command=command,
            font=F("sm", "bold"),
            fg_color=color,
            hover_color=hover or _darken(color),
            corner_radius=8,
            width=width, height=height, **kw,
        )

    def _ghost_btn(self, parent, text, command, width=120, height=34):
        return ctk.CTkButton(
            parent, text=text, command=command,
            font=F("sm"), fg_color="transparent",
            hover_color=("gray85", "gray25"),
            text_color=TEXT_MUTED,
            border_width=1, border_color=CARD_BORDER,
            corner_radius=8, width=width, height=height,
        )

    def _divider(self, parent):
        ctk.CTkFrame(parent, height=1,
                     fg_color=CARD_BORDER).pack(fill="x", pady=12)

    def _form_row(self, parent, label_text, show=None, width=340):
        block = ctk.CTkFrame(parent, fg_color="transparent")
        block.pack(fill="x", pady=6)
        self._lbl(block, label_text, size="sm", weight="bold",
                  color=TEXT_MUTED).pack(anchor="w", pady=(0, 4))
        e = self._entry(block, show=show, width=width)
        e.pack(fill="x")
        return e

    def _section_title(self, parent, text, subtitle=None):
        self._lbl(parent, text, size="xl", weight="bold").pack(
            anchor="w", pady=(0, 2))
        if subtitle:
            self._lbl(parent, subtitle, size="sm",
                      color=TEXT_MUTED).pack(anchor="w", pady=(0, 16))

    # ------------------------------------------------------------------ #
    #  Card primitives                                                   #
    # ------------------------------------------------------------------ #

    def _card_frame(self, parent, padx=20, pady=20, **kw):
        outer = ctk.CTkFrame(
            parent, corner_radius=12,
            fg_color=CARD_BG,
            border_width=1, border_color=CARD_BORDER, **kw,
        )
        inner = ctk.CTkFrame(outer, fg_color="transparent")
        inner.pack(fill="both", expand=True, padx=padx, pady=pady)
        return outer, inner

    def _action_card(self, parent, emoji, title, subtitle, command,
                     accent=ACCENT, width=220, height=130):
        outer = ctk.CTkFrame(
            parent, corner_radius=12,
            fg_color=CARD_BG,
            border_width=1, border_color=CARD_BORDER,
            width=width, height=height, cursor="hand2",
        )
        outer.pack_propagate(False)
        ctk.CTkFrame(outer, fg_color=accent,
                     height=3, corner_radius=0).place(x=0, y=0, relwidth=1)
        inner = ctk.CTkFrame(outer, fg_color="transparent")
        inner.pack(fill="both", expand=True, padx=16, pady=14)
        ctk.CTkLabel(inner, text=emoji, font=F("xl"), anchor="w").pack(anchor="w")
        ctk.CTkLabel(inner, text=title, font=F("md", "bold"),
                     anchor="w").pack(anchor="w", pady=(4, 1))
        ctk.CTkLabel(inner, text=subtitle, font=F("xs"),
                     text_color=TEXT_MUTED, anchor="w").pack(anchor="w")

        def _hover_in(e):  outer.configure(border_color=accent)
        def _hover_out(e): outer.configure(border_color=CARD_BORDER)
        def _click(e):     command()

        for w in [outer, inner]:
            w.bind("<Enter>",    _hover_in)
            w.bind("<Leave>",    _hover_out)
            w.bind("<Button-1>", _click)

        self.after(50, lambda: [
            c.bind("<Button-1>", _click)
            for c in inner.winfo_children()
        ])
        return outer

    # ------------------------------------------------------------------ #
    #  Sidebar layout                                                    #
    # ------------------------------------------------------------------ #

    def _sidebar_layout(self, items, active_key,
                         logo_line1="Interview", logo_line2="Prep",
                         user_label=None):
        self._clear()
        root_row = ctk.CTkFrame(self, fg_color="transparent")
        root_row.pack(fill="both", expand=True)

        sb = ctk.CTkFrame(root_row, width=250, corner_radius=0, fg_color=SB_BG)
        sb.pack(side="left", fill="y")
        sb.pack_propagate(False)
        self._sidebar_frame = sb

        logo_wrap = ctk.CTkFrame(sb, fg_color="transparent")
        logo_wrap.pack(fill="x", padx=20, pady=(28, 24))
        dot = ctk.CTkFrame(logo_wrap, width=8, height=8,
                           corner_radius=4, fg_color=SB_ACCENT)
        dot.pack(side="left", pady=(4, 0))
        ctk.CTkFrame(logo_wrap, width=6, fg_color="transparent").pack(side="left")
        txt_col = ctk.CTkFrame(logo_wrap, fg_color="transparent")
        txt_col.pack(side="left")
        ctk.CTkLabel(txt_col, text=logo_line1,
                     font=ctk.CTkFont(family=FONT, size=15, weight="bold"),
                     text_color=SB_TEXT, anchor="w").pack(anchor="w")
        ctk.CTkLabel(txt_col, text=logo_line2,
                     font=ctk.CTkFont(family=FONT, size=11),
                     text_color=SB_MUTED, anchor="w").pack(anchor="w")

        ctk.CTkFrame(sb, height=1, fg_color=SB_BORDER).pack(
            fill="x", padx=20, pady=(0, 12))

        self._sidebar_btns = {}

        def _set_active(key):
            for k, btn in self._sidebar_btns.items():
                if k == key:
                    btn.configure(fg_color=SB_ACTIVE, text_color=SB_TEXT)
                else:
                    btn.configure(fg_color="transparent", text_color=SB_MUTED)

        for key, emoji, label, command in items:
            def _make_cmd(k, cmd):
                def _inner():
                    _set_active(k)
                    cmd()
                return _inner

            btn = ctk.CTkButton(
                sb,
                text=f"  {emoji}  {label}",
                command=_make_cmd(key, command),
                font=ctk.CTkFont(family=FONT, size=13),
                fg_color="transparent",
                hover_color=SB_HOVER,
                text_color=SB_MUTED,
                anchor="w",
                corner_radius=8,
                height=40,
                border_width=0,
            )
            btn.pack(fill="x", padx=12, pady=2)
            self._sidebar_btns[key] = btn

        _set_active(active_key)

        if user_label:
            ctk.CTkFrame(sb, height=1, fg_color=SB_BORDER).pack(
                fill="x", padx=20, pady=(0, 8), side="bottom")
            ctk.CTkLabel(
                sb, text=f"  👤  {user_label}",
                font=ctk.CTkFont(family=FONT, size=11),
                text_color=SB_MUTED, anchor="w",
            ).pack(fill="x", padx=16, pady=(0, 16), side="bottom")

        content_wrap = ctk.CTkFrame(root_row, fg_color=CONTENT_BG, corner_radius=0)
        content_wrap.pack(side="left", fill="both", expand=True)
        self._content_frame = content_wrap
        return content_wrap

    def _sidebar_nav(self, key, builder_fn):
        for k, btn in self._sidebar_btns.items():
            if k == key:
                btn.configure(fg_color=SB_ACTIVE, text_color=SB_TEXT)
            else:
                btn.configure(fg_color="transparent", text_color=SB_MUTED)
        for w in self._content_frame.winfo_children():
            w.destroy()
        builder_fn(self._content_frame)

    # ------------------------------------------------------------------ #
    #  Auth shell                                                        #
    # ------------------------------------------------------------------ #

    def _auth_shell(self, header_title, header_sub):
        self._clear()
        root = ctk.CTkFrame(self, fg_color="transparent")
        root.pack(fill="both", expand=True)

        left = ctk.CTkFrame(root, width=420, corner_radius=0, fg_color=SB_BG)
        left.pack(side="left", fill="y")
        left.pack_propagate(False)

        brand = ctk.CTkFrame(left, fg_color="transparent")
        brand.place(relx=0.5, rely=0.42, anchor="center")
        ctk.CTkLabel(brand, text="🎯", font=ctk.CTkFont(size=52)).pack()
        ctk.CTkLabel(brand, text="Interview\nPrep",
                     font=ctk.CTkFont(family=FONT, size=28, weight="bold"),
                     text_color=SB_TEXT, justify="center").pack(pady=(8, 4))
        ctk.CTkLabel(brand,
                     text="Role-based smart practice\nfor your dream job",
                     font=ctk.CTkFont(family=FONT, size=12),
                     text_color=SB_MUTED, justify="center").pack()

        right = ctk.CTkFrame(root, fg_color=CONTENT_BG, corner_radius=0)
        right.pack(side="left", fill="both", expand=True)

        scroll = ctk.CTkScrollableFrame(right, fg_color="transparent")
        scroll.pack(fill="both", expand=True, padx=60, pady=40)

        ctk.CTkLabel(scroll, text=header_title,
                     font=ctk.CTkFont(family=FONT, size=26, weight="bold"),
                     anchor="w").pack(anchor="w", pady=(0, 4))
        ctk.CTkLabel(scroll, text=header_sub,
                     font=ctk.CTkFont(family=FONT, size=12),
                     text_color=TEXT_MUTED, anchor="w").pack(
            anchor="w", pady=(0, 28))
        return scroll

    # ------------------------------------------------------------------ #
    #  Login / register screens                                          #
    # ------------------------------------------------------------------ #

    def _show_login_home(self):
        panel = self._auth_shell("Welcome back.", "Choose how you'd like to continue.")
        cards_row = ctk.CTkFrame(panel, fg_color="transparent")
        cards_row.pack(anchor="w", pady=(0, 32))
        for emoji, title, sub, cmd, accent in [
            ("🎓", "Student Login",  "Practice & track progress",
             self._show_student_login,    ACCENT),
            ("✨", "Create Account", "Register as a new student",
             self._show_student_register, SUCCESS),
            ("🛡", "Admin Console",  "Manage roles & analytics",
             self._show_admin_login,      PINK),
        ]:
            c = self._action_card(cards_row, emoji, title, sub, cmd,
                                  accent=accent, width=190, height=120)
            c.pack(side="left", padx=(0, 12))
        self.update_idletasks()

    def _show_student_login(self):
        panel = self._auth_shell("Student login.", "Sign in to start practicing.")
        self._ghost_btn(panel, "← Back", self._show_login_home,
                        width=100, height=32).pack(anchor="w", pady=(0, 20))
        self.stu_email = self._form_row(panel, "Email address", width=380)
        self.stu_pass  = self._form_row(panel, "Password", show="*", width=380)
        row = ctk.CTkFrame(panel, fg_color="transparent")
        row.pack(anchor="w", pady=20)
        self._btn(row, "Sign in", self._student_login,
                  width=140).pack(side="left", padx=(0, 10))
        self._ghost_btn(row, "Register instead",
                        self._show_student_register, width=160).pack(side="left")
        self.stu_email.focus_set()
        self.unbind("<Return>")
        self.bind("<Return>", lambda e: self._student_login())

    def _show_student_register(self):
        panel = self._auth_shell("Create account.", "Free forever — no credit card needed.")
        self._ghost_btn(panel, "← Back", self._show_login_home,
                        width=100, height=32).pack(anchor="w", pady=(0, 20))
        self.reg_name    = self._form_row(panel, "Full name",       width=380)
        self.reg_email   = self._form_row(panel, "Email address",   width=380)
        self.reg_pass    = self._form_row(panel, "Password", show="*", width=380)
        self.reg_confirm = self._form_row(panel, "Confirm password", show="*", width=380)
        row = ctk.CTkFrame(panel, fg_color="transparent")
        row.pack(anchor="w", pady=20)
        self._btn(row, "Create account", self._register_student_action,
                  color=SUCCESS, width=160).pack(side="left", padx=(0, 10))
        self._ghost_btn(row, "Sign in instead",
                        self._show_student_login, width=150).pack(side="left")
        self.reg_name.focus_set()

    def _show_admin_login(self):
        panel = self._auth_shell("Admin console.", "Manage roles, questions & analytics.")
        self._ghost_btn(panel, "← Back", self._show_login_home,
                        width=100, height=32).pack(anchor="w", pady=(0, 20))
        self.adm_user = self._form_row(panel, "Username", width=380)
        self.adm_pass = self._form_row(panel, "Password", show="*", width=380)
        self._btn(panel, "Sign in", self._admin_login,
                  color=PINK, width=380).pack(anchor="w", pady=16)
        self._divider(panel)
        self._lbl(panel, "First-time setup", size="lg",
                  weight="bold").pack(anchor="w", pady=(0, 4))
        self._lbl(panel, "No admin yet? Create one below.",
                  size="sm", color=TEXT_MUTED).pack(anchor="w", pady=(0, 14))
        self.new_adm_user = self._form_row(panel, "New admin username", width=380)
        self.new_adm_pass = self._form_row(panel, "New admin password",
                                           show="*", width=380)
        self._btn(panel, "Create admin account", self._create_admin,
                  color=("#7c3aed", "#6d28d9"), width=380).pack(anchor="w", pady=12)
        self.adm_user.focus_set()
        self.unbind("<Return>")
        self.bind("<Return>", lambda e: self._admin_login())

    # ------------------------------------------------------------------ #
    #  Auth actions                                                      #
    # ------------------------------------------------------------------ #

    def _student_login(self):
        try:
            email    = self.stu_email.get().strip()
            password = self.stu_pass.get()
        except (tk.TclError, AttributeError):
            messagebox.showerror("Error", "Please open Student Login first.")
            return
        if not email or not password:
            messagebox.showwarning("Missing fields", "Enter email and password.")
            return
        ok, msg = auth.student_login(email, password)
        if ok:
            self.unbind("<Return>")
            self._show_student_dashboard()
        else:
            messagebox.showerror("Login Failed", msg)

    def _register_student_action(self):
        try:
            name     = self.reg_name.get().strip()
            email    = self.reg_email.get().strip()
            password = self.reg_pass.get()
            confirm  = self.reg_confirm.get()
        except (tk.TclError, AttributeError):
            messagebox.showerror("Error", "Please open Register screen first.")
            return
        if not all([name, email, password, confirm]):
            messagebox.showwarning("Missing fields", "Fill in all fields.")
            return
        ok, msg = auth.student_register(name, email, password, confirm)
        if ok:
            messagebox.showinfo("Success", msg)
            self._show_student_login()
        else:
            messagebox.showerror("Registration Failed", msg)

    def _admin_login(self):
        try:
            username = self.adm_user.get().strip()
            password = self.adm_pass.get()
        except (tk.TclError, AttributeError):
            messagebox.showerror("Error", "Please open Admin Login first.")
            return
        if not username or not password:
            messagebox.showwarning("Missing fields", "Enter username and password.")
            return
        ok, msg = auth.admin_login(username, password)
        if ok:
            self.unbind("<Return>")
            self._show_admin_dashboard()
        else:
            messagebox.showerror("Login Failed", msg)

    def _create_admin(self):
        try:
            username = self.new_adm_user.get().strip()
            password = self.new_adm_pass.get()
        except (tk.TclError, AttributeError):
            messagebox.showerror("Error", "Please open Admin Login first.")
            return
        if not username or not password:
            messagebox.showwarning("Missing fields", "Enter username and password.")
            return
        ok, msg = auth.create_admin(username, password)
        if ok:
            messagebox.showinfo("Success", msg + "\nYou can now log in.")
        else:
            messagebox.showerror("Error", msg)

    def _logout(self):
        auth.logout()
        self.unbind("<Return>")
        self._show_login_home()

    # ------------------------------------------------------------------ #
    #  Student dashboard                                                 #
    # ------------------------------------------------------------------ #

    def _show_student_dashboard(self):
        profile = student.load_profile()
        uid     = auth.get_session()["user"]["id"]

        content = self._sidebar_layout(
            items=[
                ("dashboard", "🏠", "Dashboard",      lambda: None),
                ("interview", "🎯", "Start Interview", lambda: None),
                ("reports",   "📊", "Reports",         lambda: None),
                ("logout",    "→",  "Logout",          self._logout),
            ],
            active_key="dashboard",
            logo_line1="Interview",
            logo_line2="Prep",
            user_label=profile["name"],
        )

        def build_dashboard(parent):
            for w in parent.winfo_children(): w.destroy()
            scroll = ctk.CTkScrollableFrame(parent, fg_color="transparent")
            scroll.pack(fill="both", expand=True, padx=32, pady=28)

            self._lbl(scroll, f"Hi, {profile['name']} 👋",
                      size="xxl", weight="bold").pack(anchor="w")
            self._lbl(scroll, "Here's your interview summary.",
                      size="sm", color=TEXT_MUTED).pack(anchor="w", pady=(2, 20))

            history   = feedback_engine.view_history(uid)
            completed = len(history)
            avg_score = (
                round(sum(h["overall_score"] for h in history) / completed)
                if completed else 0
            )
            best_role = "—"
            if history:
                best      = max(history, key=lambda h: h["overall_score"])
                best_role = best["role_name"]

            stats_row = ctk.CTkFrame(scroll, fg_color="transparent")
            stats_row.pack(fill="x", pady=(0, 24))
            stats_row.columnconfigure((0, 1, 2), weight=1)

            for col_i, (label, value, sub) in enumerate([
                ("Completed Interviews", completed,        ""),
                ("Average Score",        f"{avg_score}%",  "across all rounds"),
                ("Best Role",            best_role,         "highest score"),
            ]):
                outer, inner = self._card_frame(stats_row, padx=18, pady=16)
                outer.grid(row=0, column=col_i, padx=(0, 12), sticky="ew")
                self._lbl(inner, label, size="xs", weight="bold",
                          color=TEXT_MUTED).pack(anchor="w")
                self._lbl(inner, str(value), size="xxl",
                          weight="bold").pack(anchor="w", pady=(4, 0))
                if sub:
                    self._lbl(inner, sub, size="xs", color=TEXT_MUTED).pack(anchor="w")

            self._lbl(scroll, "Quick actions", size="lg",
                      weight="bold").pack(anchor="w", pady=(4, 12))
            cards_row = ctk.CTkFrame(scroll, fg_color="transparent")
            cards_row.pack(anchor="w")

            for emoji, title, sub, cmd, accent in [
                ("🎯", "Start Interview", "Pick a role and begin",
                 lambda: self._sidebar_nav("interview", build_interview), ACCENT),
                ("📊", "View Reports",   "See your score history",
                 lambda: self._sidebar_nav("reports",   build_reports),   TEAL),
            ]:
                c = self._action_card(cards_row, emoji, title, sub,
                                      cmd, accent=accent, width=200, height=120)
                c.pack(side="left", padx=(0, 12))

            self._lbl(scroll, "Your profile", size="lg",
                      weight="bold").pack(anchor="w", pady=(24, 12))
            p_outer, p_inner = self._card_frame(scroll, padx=20, pady=16)
            p_outer.pack(fill="x")
            for field, val in [("Name",         profile["name"]),
                                ("Email",        profile["email"]),
                                ("Member since", profile["created_at"])]:
                row = ctk.CTkFrame(p_inner, fg_color="transparent")
                row.pack(fill="x", pady=5)
                self._lbl(row, field, size="xs", weight="bold",
                          color=TEXT_MUTED).pack(side="left", padx=(0, 12))
                self._lbl(row, val, size="sm").pack(side="left")

        def build_interview(parent):
            for w in parent.winfo_children(): w.destroy()
            scroll = ctk.CTkScrollableFrame(parent, fg_color="transparent")
            scroll.pack(fill="both", expand=True, padx=32, pady=28)

            self._section_title(scroll, "Start Interview",
                                "Pick a role and difficulty, then go.")

            roles      = admin.fetch_roles()
            self.role_var  = tk.StringVar()
            role_names = [f"{r['id']}: {r['role_name']}" for r in roles]

            outer, inner = self._card_frame(scroll, padx=24, pady=20)
            outer.pack(fill="x", pady=(0, 20))

            self._lbl(inner, "Job role", size="sm", weight="bold",
                      color=TEXT_MUTED).pack(anchor="w", pady=(0, 6))
            if role_names:
                self.role_var.set(role_names[0])
                ctk.CTkComboBox(
                    inner, variable=self.role_var, values=role_names,
                    width=380, state="readonly", font=F("md"),
                    corner_radius=8,
                ).pack(anchor="w")
            else:
                self._lbl(inner, "No roles yet — ask your admin.",
                          size="sm", color=TEXT_MUTED).pack(anchor="w")

            ctk.CTkFrame(inner, height=1, fg_color=CARD_BORDER).pack(fill="x", pady=16)

            self._lbl(inner, "Difficulty", size="sm", weight="bold",
                      color=TEXT_MUTED).pack(anchor="w", pady=(0, 8))
            self.diff_var = tk.StringVar(value="Easy")
            diff_row      = ctk.CTkFrame(inner, fg_color="transparent")
            diff_row.pack(anchor="w")
            for d, accent in [("Easy", SUCCESS), ("Medium", AMBER), ("Hard", DANGER)]:
                ctk.CTkRadioButton(
                    diff_row, text=f"  {d}",
                    variable=self.diff_var, value=d,
                    font=F("md"), fg_color=accent, hover_color=accent,
                ).pack(side="left", padx=(0, 20))

            ctk.CTkFrame(inner, height=1, fg_color=CARD_BORDER).pack(fill="x", pady=16)
            self._btn(inner, "🚀  Start Interview",
                      self._start_interview, width=200, height=42).pack(anchor="w")

        def build_reports(parent):
            for w in parent.winfo_children(): w.destroy()
            scroll = ctk.CTkScrollableFrame(parent, fg_color="transparent")
            scroll.pack(fill="both", expand=True, padx=32, pady=28)

            self._section_title(scroll, "Your Reports",
                                "All your past interview sessions.")

            history = feedback_engine.view_history(uid)
            _style_tree()
            cols = ("id", "role", "mcq", "tech", "hr", "overall", "date")
            tree = ttk.Treeview(scroll, columns=cols, show="headings", height=12)
            for col, heading, w in zip(
                cols,
                ("ID", "Role", "MCQ", "Tech", "HR", "Overall", "Date"),
                (40,  180,    60,    60,    60,   70,      120),
            ):
                tree.heading(col, text=heading)
                tree.column(col, width=w, anchor="center")
            tree.pack(fill="both", expand=True, pady=8)

            for h in history:
                tree.insert("", "end", values=(
                    h["id"], h["role_name"],
                    h["mcq_score"], h["technical_score"],
                    h["hr_score"], h["overall_score"], h["created_at"],
                ))

            btn_row = ctk.CTkFrame(scroll, fg_color="transparent")
            btn_row.pack(anchor="w", pady=8)
            self._btn(btn_row, "View Feedback",
                      lambda: self._view_feedback(tree),
                      width=140).pack(side="left", padx=(0, 10))
            self._btn(btn_row, "Perfect answers",
                      lambda: self._show_perfect_answers(tree),
                      color=SUCCESS, width=150).pack(side="left", padx=(0, 10))
            self._btn(btn_row, "Export CSV",
                      lambda: self._export_report(tree),
                      color=TEAL, width=130).pack(side="left")

        # Wire sidebar
        self._sidebar_btns["dashboard"].configure(
            command=lambda: self._sidebar_nav("dashboard", build_dashboard))
        self._sidebar_btns["interview"].configure(
            command=lambda: self._sidebar_nav("interview", build_interview))
        self._sidebar_btns["reports"].configure(
            command=lambda: self._sidebar_nav("reports",   build_reports))

        build_dashboard(content)

    # ------------------------------------------------------------------ #
    #  Feedback / export helpers                                         #
    # ------------------------------------------------------------------ #

    def _show_perfect_answers(self, tree_or_result_id):
        if isinstance(tree_or_result_id, int):
            result_id = tree_or_result_id
        else:
            tree = tree_or_result_id
            sel = tree.selection()
            if not sel:
                messagebox.showwarning("Select", "Select a report first.")
                return
            result_id = tree.item(sel[0])["values"][0]

        uid = auth.get_session()["user"]["id"]
        review = feedback_engine.get_session_review(result_id, uid)
        text = feedback_engine.format_session_review_text(review)

        win = ctk.CTkToplevel(self)
        win.title("Perfect answers — your session")
        win.geometry("720x520")
        self._lbl(win,
                  "Compare your answers with the correct / model responses.",
                  size="sm", color=TEXT_MUTED).pack(
            anchor="w", padx=14, pady=(12, 4))
        tb = ctk.CTkTextbox(win, wrap="word",
                            font=ctk.CTkFont(family=FONT, size=12))
        tb.pack(fill="both", expand=True, padx=12, pady=(0, 12))
        tb.insert("1.0", text)
        tb.configure(state="disabled")

    def _view_feedback(self, tree):
        sel = tree.selection()
        if not sel:
            messagebox.showwarning("Select", "Select a report first.")
            return
        rid     = tree.item(sel[0])["values"][0]
        uid     = auth.get_session()["user"]["id"]
        history = feedback_engine.view_history(uid)
        row     = next((h for h in history if h["id"] == rid), None)
        if row:
            win = ctk.CTkToplevel(self)
            win.title("Feedback Report")
            win.geometry("640x440")
            tb = ctk.CTkTextbox(win, wrap="word",
                                font=ctk.CTkFont(family=FONT, size=11))
            tb.pack(fill="both", expand=True, padx=12, pady=12)
            tb.insert("1.0", row["feedback"])
            tb.configure(state="disabled")
            btn_row = ctk.CTkFrame(win, fg_color="transparent")
            btn_row.pack(fill="x", padx=12, pady=(0, 12))
            self._btn(btn_row, "View perfect answers",
                      lambda: self._show_perfect_answers(rid),
                      color=SUCCESS, width=180).pack(side="left")

    def _export_report(self, tree):
        sel = tree.selection()
        if not sel:
            messagebox.showwarning("Select", "Select a report first.")
            return
        rid  = tree.item(sel[0])["values"][0]
        path = filedialog.asksaveasfilename(
            defaultextension=".csv", filetypes=[("CSV", "*.csv")])
        if path:
            ok, msg = feedback_engine.export_report(rid, path)
            if ok:
                messagebox.showinfo("Exported", f"Saved to:\n{msg}")
            else:
                messagebox.showerror("Error", msg)

    def _export_latest(self, result_id):
        path = filedialog.asksaveasfilename(
            defaultextension=".csv", filetypes=[("CSV", "*.csv")])
        if path:
            ok, msg = feedback_engine.export_report(result_id, path)
            if ok:
                messagebox.showinfo("Exported", f"Saved to:\n{msg}")

    # ------------------------------------------------------------------ #
    #  Start interview                                                   #
    # ------------------------------------------------------------------ #

    def _start_interview(self):
        role_sel = self.role_var.get().strip()
        if not role_sel or ":" not in role_sel:
            messagebox.showwarning(
                "Role", "No valid job role available. Ask admin to create one.")
            return
        try:
            role_id = int(role_sel.split(":")[0])
        except ValueError:
            messagebox.showwarning("Role", "Invalid role selected.")
            return

        role       = student.select_role(role_id)
        difficulty = student.select_difficulty(self.diff_var.get())
        if not role or not difficulty:
            messagebox.showerror("Error", "Invalid role or difficulty.")
            return

        session = auth.get_session()
        self.update_idletasks()
        ok, msg = interview_engine.start_interview(
            session["user"]["id"], role_id, role["role_name"], difficulty)
        if not ok:
            messagebox.showerror("Cannot Start", msg)
            return
        self._show_interview_round()

    def _advance_to_next_round(self):
        sess = interview_engine.get_active_session()
        if sess.current_round == 1:
            interview_engine.calculate_mcq_score()
        elif sess.current_round == 2:
            interview_engine.calculate_technical_score()
        completed_round = sess.current_round
        interview_engine.advance_round()
        messagebox.showinfo("Next Round",
                            f"Round {completed_round} complete. Starting next round.")
        self._show_interview_round()
    def _show_interview_round(self):
        self._clear()

        # Get active session
        sess = interview_engine.get_active_session()

        # FIX: prevent crash if session expired
        if not sess:
            messagebox.showerror(
                "Session Expired",
                "No active interview session found."
            )
            self._show_student_dashboard()
            return

        round_meta = {
            1: ("MCQ Round", ACCENT, "#312e81"),
            2: ("Technical Round", TEAL, "#0d9488"),
            3: ("HR Round", PINK, "#9d174d"),
        }

        rname, badge_fg, badge_hover = round_meta.get(
            sess.current_round,
            ("Complete", ACCENT, "#312e81")
        )

        # =========================
        # MAIN ROOT
        # =========================
        root = ctk.CTkFrame(
            self,
            fg_color="transparent"
        )
        root.pack(fill="both", expand=True)

        ctk.CTkFrame(
            root,
            width=6,
            corner_radius=0,
            fg_color=badge_fg
        ).pack(side="left", fill="y")

        main = ctk.CTkFrame(
            root,
            fg_color=CONTENT_BG,
            corner_radius=0
        )
        main.pack(side="left", fill="both", expand=True)

        # =========================
        # TOP BAR
        # =========================
        topbar = ctk.CTkFrame(
            main,
            fg_color=("white", "#1e1e2e"),
            height=60,
            corner_radius=0
        )
        topbar.pack(fill="x")
        topbar.pack_propagate(False)

        tb = ctk.CTkFrame(
            topbar,
            fg_color="transparent"
        )
        tb.pack(fill="both", expand=True, padx=24, pady=10)

        ctk.CTkLabel(
            tb,
            text=rname,
            font=ctk.CTkFont(
                family=FONT,
                size=16,
                weight="bold"
            )
        ).pack(side="left")

        ctk.CTkLabel(
            tb,
            text=f"  ·  {sess.role_name}  ·  {sess.difficulty}",
            font=ctk.CTkFont(
                family=FONT,
                size=12
            ),
            text_color=TEXT_MUTED
        ).pack(side="left")

        # EXIT BUTTON (reset session)
        self._ghost_btn(
            tb,
            "✕ Exit",
            lambda: (
                interview_engine.reset_session(),
                self._show_student_dashboard()
            ),
            width=90,
            height=30
        ).pack(side="right")

        # =========================
        # CURRENT QUESTION
        # =========================
        q = interview_engine.get_current_question()

        if not q:
            if sess.current_round < 3:
                self._advance_to_next_round()
                return
            else:
                interview_engine.calculate_hr_score()
                self._show_results()
                return

        # =========================
        # TOTAL QUESTIONS
        # =========================
        if sess.current_round == 1:
            total = len(sess.mcq_questions)
        elif sess.current_round == 2:
            total = len(sess.technical_questions)
        else:
            total = len(sess.hr_questions)

        idx = sess.current_index + 1

        # =========================
        # PROGRESS BAR
        # =========================
        prog = ctk.CTkProgressBar(
            main,
            fg_color=("gray85", "gray25"),
            progress_color=badge_fg,
            height=4,
            corner_radius=0
        )
        prog.set(idx / max(total, 1))
        prog.pack(fill="x")

        # =========================
        # SCROLL AREA
        # =========================
        scroll = ctk.CTkScrollableFrame(
            main,
            fg_color="transparent"
        )
        scroll.pack(
            fill="both",
            expand=True,
            padx=40,
            pady=24
        )

        # =========================
        # QUESTION CARD
        # =========================
        q_outer, q_inner = self._card_frame(
            scroll,
            padx=28,
            pady=24
        )
        q_outer.pack(fill="x", pady=(0, 16))

        badge_frame = ctk.CTkFrame(
            q_inner,
            fg_color=badge_fg,
            corner_radius=6
        )
        badge_frame.pack(anchor="w", pady=(0, 14))

        ctk.CTkLabel(
            badge_frame,
            text=f"  Question {idx} of {total}  ",
            font=ctk.CTkFont(
                family=FONT,
                size=10,
                weight="bold"
            ),
            text_color="white"
        ).pack(padx=6, pady=4)

        # Question text
        ctk.CTkLabel(
            q_inner,
            text=q["question"],
            font=ctk.CTkFont(
                family=FONT,
                size=15
            ),
            wraplength=700,
            justify="left",
            anchor="w"
        ).pack(
            anchor="w",
            pady=(0, 20),
            fill="x"
        )

        self.answer_var = tk.StringVar()

        # =========================
        # MCQ ROUND
        # =========================
        if sess.current_round == 1:
            for opt, letter in [
                (q["option_a"], "A"),
                (q["option_b"], "B"),
                (q["option_c"], "C"),
                (q["option_d"], "D"),
            ]:
                opt_frame = ctk.CTkFrame(
                    q_inner,
                    fg_color=("gray97", "gray20"),
                    corner_radius=8
                )
                opt_frame.pack(fill="x", pady=4)

                ctk.CTkRadioButton(
                    opt_frame,
                    text=f"  {letter}.   {opt}",
                    variable=self.answer_var,
                    value=letter,
                    font=ctk.CTkFont(
                        family=FONT,
                        size=13
                    ),
                    fg_color=badge_fg,
                    hover_color=badge_hover
                ).pack(
                    anchor="w",
                    padx=16,
                    pady=10
                )

        # =========================
        # TECHNICAL / HR ROUND
        # =========================
        else:
            self._lbl(
                q_inner,
                "Your answer",
                size="sm",
                weight="bold",
                color=TEXT_MUTED
            ).pack(anchor="w", pady=(0, 6))

            self.answer_text = ctk.CTkTextbox(
                q_inner,
                height=160,
                wrap="word",
                font=ctk.CTkFont(
                    family=FONT,
                    size=13
                ),
                corner_radius=8
            )
            self.answer_text.pack(fill="x")

        # =========================
        # SUBMIT BUTTON
        # =========================
        self._btn(
            q_inner,
            "Submit answer →",
            lambda q=q: self._submit_answer(q),
            width=180,
            height=42
        ).pack(
            anchor="w",
            pady=(20, 0)
        )
        
    def _submit_answer(self, q):
        sess = interview_engine.get_active_session()

        if sess.current_round == 1:
            ans = self.answer_var.get()
            if not ans:
                messagebox.showwarning("Answer", "Select an option.")
                return
            interview_engine.submit_mcq_answer(q["id"], ans)

        elif sess.current_round == 2:
            ans = self.answer_text.get("1.0", "end").strip()
            if not ans:
                messagebox.showwarning("Answer", "Enter your answer.")
                return
            result = interview_engine.submit_technical_answer(
                q["id"], ans, q["keywords"])
            messagebox.showinfo(
                "Feedback",
                f"Score: {result['score']}%\n{result['feedback']}\n"
                f"Keywords matched: {result['matched_keywords']}/{result['total_keywords']}")
        else:
            ans = self.answer_text.get("1.0", "end").strip()
            if not ans:
                messagebox.showwarning("Answer", "Enter your answer.")
                return
            result = interview_engine.submit_hr_answer(q["id"], ans, q["keywords"])
            messagebox.showinfo(
                "Feedback",
                f"Score: {result['final_score']}%\n{result['feedback']}")

        has_more = interview_engine.next_question()
        if has_more:
            self._show_interview_round()
        elif sess.current_round < 3:
            self._advance_to_next_round()
        else:
            interview_engine.calculate_hr_score()
            self._show_results()

    # ------------------------------------------------------------------ #
    #  Results screen                                                    #
    # ------------------------------------------------------------------ #

    def _show_results(self):
        report = feedback_engine.finalize_interview()
        if not report:
            messagebox.showerror("Error", "Could not generate report.")
            self._show_student_dashboard()
            return

        self._last_result_id = report["result_id"]
        self._clear()

        root = ctk.CTkFrame(self, fg_color=CONTENT_BG, corner_radius=0)
        root.pack(fill="both", expand=True)
        scroll = ctk.CTkScrollableFrame(root, fg_color="transparent")
        scroll.pack(fill="both", expand=True, padx=48, pady=36)

        self._lbl(scroll, "Interview complete! 🎉",
                  size="xxl", weight="bold").pack(anchor="w")
        self._lbl(scroll, f"{report['role_name']}  ·  {report['difficulty']}",
                  size="sm", color=TEXT_MUTED).pack(anchor="w", pady=(2, 24))

        scores_row = ctk.CTkFrame(scroll, fg_color="transparent")
        scores_row.pack(fill="x", pady=(0, 24))
        scores_row.columnconfigure((0, 1, 2, 3), weight=1)

        for col_i, (label, val, color) in enumerate([
            ("MCQ",     report["mcq_score"],       ACCENT),
            ("Tech",    report["technical_score"],  TEAL),
            ("HR",      report["hr_score"],         PINK),
            ("Overall", report["overall_score"],    AMBER),
        ]):
            box = ctk.CTkFrame(scores_row, fg_color=color, corner_radius=12)
            box.grid(row=0, column=col_i, padx=(0, 12), sticky="ew")
            ctk.CTkLabel(box, text=label,
                         font=ctk.CTkFont(family=FONT, size=10, weight="bold"),
                         text_color="white").pack(pady=(14, 2))
            ctk.CTkLabel(box, text=f"{val}%",
                         font=ctk.CTkFont(family=FONT, size=28, weight="bold"),
                         text_color="white").pack(pady=(0, 14))

        self._lbl(scroll, "Detailed feedback", size="lg",
                  weight="bold").pack(anchor="w", pady=(4, 10))
        fb_outer, fb_inner = self._card_frame(scroll, padx=20, pady=16)
        fb_outer.pack(fill="x")
        tb = ctk.CTkTextbox(fb_inner, height=200, wrap="word",
                            font=ctk.CTkFont(family=FONT, size=12))
        tb.pack(fill="both", expand=True)
        tb.insert("1.0", report["feedback"])
        tb.configure(state="disabled")

        btn_row = ctk.CTkFrame(scroll, fg_color="transparent")
        btn_row.pack(anchor="w", pady=16)
        self._btn(btn_row, "Back to Dashboard",
                  self._show_student_dashboard, width=180).pack(
            side="left", padx=(0, 10))
        self._btn(btn_row, "View perfect answers",
                  lambda: self._show_perfect_answers(report["result_id"]),
                  color=SUCCESS, width=180).pack(side="left", padx=(0, 10))
        self._btn(btn_row, "Export CSV",
                  lambda: self._export_latest(report["result_id"]),
                  color=TEAL, width=140).pack(side="left")

    # ------------------------------------------------------------------ #
    #  Admin dashboard                                                   #
    # ------------------------------------------------------------------ #

    def _show_admin_dashboard(self):
        session = auth.get_session()

        content = self._sidebar_layout(
            items=[
                ("roles",     "📋", "Roles",     lambda: None),
                ("questions", "❓", "Questions", lambda: None),
                ("analytics", "📈", "Analytics", lambda: None),
                ("logout",    "→",  "Logout",    self._logout),
            ],
            active_key="roles",
            logo_line1="Admin",
            logo_line2="Console",
            user_label=session["user"]["username"],
        )

        def build_roles(parent):
            for w in parent.winfo_children(): w.destroy()
            scroll = ctk.CTkScrollableFrame(parent, fg_color="transparent")
            scroll.pack(fill="both", expand=True, padx=28, pady=24)
            self._section_title(scroll, "Job Roles", "Add, rename or remove roles.")
            form = ctk.CTkFrame(scroll, fg_color="transparent")
            form.pack(anchor="w", pady=(0, 16))
            self.new_role_entry = self._entry(form, width=280)
            self.new_role_entry.pack(side="left", padx=(0, 10))
            self._btn(form, "+ Add Role", self._add_role,
                      color=SUCCESS, width=120).pack(side="left")

            _style_tree()
            self.roles_tree = ttk.Treeview(scroll, columns=("id", "name"),
                                           show="headings", height=10)
            self.roles_tree.heading("id",   text="ID")
            self.roles_tree.heading("name", text="Role Name")
            self.roles_tree.column("id",   width=50,  anchor="center")
            self.roles_tree.column("name", width=340)
            self.roles_tree.pack(fill="x", pady=(0, 12))

            btns = ctk.CTkFrame(scroll, fg_color="transparent")
            btns.pack(anchor="w")
            self._ghost_btn(btns, "Refresh", self._refresh_roles,
                            width=100).pack(side="left", padx=(0, 8))
            self._btn(btns, "Rename", self._update_role,
                      width=100).pack(side="left", padx=(0, 8))
            self._btn(btns, "Delete", self._delete_role,
                      color=DANGER, width=100).pack(side="left")
            self._refresh_roles()

        def build_questions(parent):
            for w in parent.winfo_children(): w.destroy()
            scroll = ctk.CTkScrollableFrame(parent, fg_color="transparent")
            scroll.pack(fill="both", expand=True, padx=28, pady=24)
            self._section_title(scroll, "Question Bank",
                                "Add MCQ, Technical or HR questions.")
            self._admin_questions_panel(scroll)

        def build_analytics(parent):
            for w in parent.winfo_children(): w.destroy()
            scroll = ctk.CTkScrollableFrame(parent, fg_color="transparent")
            scroll.pack(fill="both", expand=True, padx=28, pady=24)
            self._section_title(scroll, "Analytics",
                                "View student performance data.")
            self._admin_analytics_panel(scroll)

        self._sidebar_btns["roles"].configure(
            command=lambda: self._sidebar_nav("roles",     build_roles))
        self._sidebar_btns["questions"].configure(
            command=lambda: self._sidebar_nav("questions", build_questions))
        self._sidebar_btns["analytics"].configure(
            command=lambda: self._sidebar_nav("analytics", build_analytics))

        build_roles(content)

    # ------------------------------------------------------------------ #
    #  Admin — roles CRUD                                                #
    # ------------------------------------------------------------------ #

    def _refresh_roles(self):
        for i in self.roles_tree.get_children():
            self.roles_tree.delete(i)
        for r in admin.fetch_roles():
            self.roles_tree.insert("", "end", values=(r["id"], r["role_name"]))

    def _add_role(self):
        ok, msg = admin.add_role(self.new_role_entry.get())
        if ok:
            self.new_role_entry.delete(0, tk.END)
            self._refresh_roles()
        messagebox.showinfo("Role", msg)

    def _update_role(self):
        sel = self.roles_tree.selection()
        if not sel:
            messagebox.showwarning("Select", "Select a role first.")
            return
        rid, name = self.roles_tree.item(sel[0])["values"]
        new_name  = simpledialog.askstring(
            "Rename Role", "New role name:", initialvalue=name)
        if new_name:
            ok, msg = admin.update_role(rid, new_name)
            self._refresh_roles()
            messagebox.showinfo("Role", msg)

    def _delete_role(self):
        sel = self.roles_tree.selection()
        if not sel:
            return
        rid = self.roles_tree.item(sel[0])["values"][0]
        if messagebox.askyesno("Confirm", "Delete this role and all its questions?"):
            ok, msg = admin.delete_role(rid)
            self._refresh_roles()
            messagebox.showinfo("Role", msg)

    # ------------------------------------------------------------------ #
    #  Admin — questions panel                                           #
    # ------------------------------------------------------------------ #

    def _admin_questions_panel(self, parent):
        tab_row = ctk.CTkFrame(parent, fg_color="transparent")
        tab_row.pack(anchor="w", pady=(0, 16))
        q_content = ctk.CTkFrame(parent, fg_color="transparent")
        q_content.pack(fill="both", expand=True)

        def clear_q():
            for w in q_content.winfo_children(): w.destroy()

        for label, cmd_fn, color in [
            ("MCQ",       lambda: (clear_q(), self._build_mcq_form(q_content)),  ACCENT),
            ("Technical", lambda: (clear_q(), self._build_tech_form(q_content)), TEAL),
            ("HR",        lambda: (clear_q(), self._build_hr_form(q_content)),   PINK),
            ("Bulk Upload", lambda: (clear_q(), self._build_bulk_upload(q_content)), AMBER),
            ("View All",  lambda: (clear_q(), self._build_q_list(q_content)),
             ("gray40", "gray50")),
        ]:
            self._btn(tab_row, label, cmd_fn,
                      color=color, width=110, height=34).pack(
                side="left", padx=(0, 8))

        clear_q()
        self._build_mcq_form(q_content)

    def _role_combo_widget(self, parent):
        roles = admin.fetch_roles()
        var   = tk.StringVar()
        vals  = [f"{r['id']}: {r['role_name']}" for r in roles]
        if not vals:
            vals = ["No roles available"]
        cb = ctk.CTkComboBox(parent, variable=var, values=vals,
                             width=340, state="readonly", font=F("md"))
        var.set(vals[0])
        return var, cb

    def _get_role_id_from_var(self, var):
        value = var.get().strip()
        if not value or ":" not in value:
            return None
        try:
            return int(value.split(":")[0])
        except ValueError:
            return None

    def _grid_label(self, parent, text, row):
        ctk.CTkLabel(parent, text=text, font=F("sm", "bold"),
                     text_color=TEXT_MUTED, anchor="e").grid(
            row=row, column=0, sticky="e", padx=(0, 10), pady=5)

    def _grid_entry(self, parent, row):
        e = self._entry(parent, width=340)
        e.grid(row=row, column=1, pady=5, sticky="ew")
        return e

    def _build_mcq_form(self, parent):
        parent.columnconfigure(1, weight=1)
        self.mcq_role_var, cb = self._role_combo_widget(parent)
        self._grid_label(parent, "Role:", 0)
        cb.grid(row=0, column=1, pady=5, sticky="ew")
        self._grid_label(parent, "Difficulty:", 1)
        self.mcq_diff = ctk.CTkComboBox(parent,
                                         values=["Easy", "Medium", "Hard"],
                                         width=340, state="readonly", font=F("md"))
        self.mcq_diff.set("Easy")
        self.mcq_diff.grid(row=1, column=1, pady=5)
        self.mcq_q   = self._grid_labeled_entry(parent, "Question:", 2)
        self.mcq_a   = self._grid_labeled_entry(parent, "Option A:", 3)
        self.mcq_b   = self._grid_labeled_entry(parent, "Option B:", 4)
        self.mcq_c   = self._grid_labeled_entry(parent, "Option C:", 5)
        self.mcq_d   = self._grid_labeled_entry(parent, "Option D:", 6)
        self.mcq_ans = self._grid_labeled_entry(parent, "Correct (A/B/C/D):", 7)
        self._btn(parent, "Add MCQ", self._add_mcq,
                  width=130).grid(row=8, column=1, pady=14, sticky="w")

    def _build_tech_form(self, parent):
        parent.columnconfigure(1, weight=1)
        self.tech_role_var, cb = self._role_combo_widget(parent)
        self._grid_label(parent, "Role:", 0)
        cb.grid(row=0, column=1, pady=5, sticky="ew")
        self._grid_label(parent, "Difficulty:", 1)
        self.tech_diff = ctk.CTkComboBox(parent,
                                          values=["Easy", "Medium", "Hard"],
                                          width=340, state="readonly", font=F("md"))
        self.tech_diff.set("Easy")
        self.tech_diff.grid(row=1, column=1, pady=5)
        self.tech_q  = self._grid_labeled_entry(parent, "Question:", 2)
        self.tech_kw = self._grid_labeled_entry(parent,
                                                 "Keywords (comma-sep, min 3):", 3)
        self._btn(parent, "Add Technical", self._add_technical,
                  color=TEAL, width=150).grid(row=4, column=1, pady=14, sticky="w")

    # def _build_hr_form(self, parent):
    #     parent.columnconfigure(1, weight=1)
    #     self.hr_role_var, cb = self._role_combo_widget(parent)
    #     self._grid_label(parent, "Role:", 0)
    #     cb.grid(row=0, column=1, pady=5, sticky="ew")
    #     self.hr_q  = self._grid_labeled_entry(parent, "Question:", 1)
    #     self.hr_kw = self._grid_labeled_entry(parent, "Expected Keywords:", 2)
    #     self._btn(parent, "Add HR", self._add_hr,
    #               color=PINK, width=120).grid(row=3, column=1, pady=14, sticky="w")
    
    def _build_hr_form(self, parent):
        parent.columnconfigure(1, weight=1)

        # Role dropdown
        self.hr_role_var, cb = self._role_combo_widget(parent)
        self._grid_label(parent, "Role:", 0)
        cb.grid(row=0, column=1, pady=5, sticky="ew")

        # Difficulty dropdown
        self._grid_label(parent, "Difficulty:", 1)

        self.hr_diff = ctk.CTkComboBox(
            parent,
            values=["Easy", "Medium", "Hard"],
            width=340,
            state="readonly",
            font=F("md")
        )
        self.hr_diff.set("Easy")
        self.hr_diff.grid(
            row=1,
            column=1,
            pady=5,
            sticky="ew"
        )

        # Question
        self.hr_q = self._grid_labeled_entry(
            parent,
            "Question:",
            2
        )

        # Expected Keywords
        self.hr_kw = self._grid_labeled_entry(
            parent,
            "Expected Keywords:",
            3
        )

        # Add button
        self._btn(
            parent,
            "Add HR",
            self._add_hr,
            color=PINK,
            width=120
        ).grid(
            row=4,
            column=1,
            pady=14,
            sticky="w"
        )

    def _download_template(self):
        import csv
        import json
        from docx import Document

        win = ctk.CTkToplevel(self)
        win.title("Download Template")
        win.geometry("420x220")
        win.resizable(False, False)
        win.grab_set()

        ctk.CTkLabel(win, text="Choose Template Format", font=F("lg")).pack(pady=(20, 10))
        ctk.CTkLabel(win, text="Click any format to download instantly.",
                     font=F("sm"), text_color=TEXT_MUTED).pack(pady=(0, 20))

        def download_csv():
            filepath = filedialog.asksaveasfilename(
                defaultextension=".csv", filetypes=[("CSV File", "*.csv")],
                initialfile="questions_template.csv")
            if not filepath:
                return
            rows = [
                ["type", "question", "option_a", "option_b", "option_c",
                 "option_d", "correct_answer", "keywords", "ideal_answer"],
                ["mcq", "What is polymorphism?", "Inheritance", "Method Overloading",
                 "Abstraction", "Encapsulation", "B", "", ""],
                ["technical", "Explain deadlock in OS.", "", "", "", "", "",
                 "deadlock, locking, process", "Should explain circular wait."],
                ["hr", "Tell me about a challenge you faced.", "", "", "", "", "",
                 "teamwork, communication", "Should explain challenge and learning."]
            ]
            with open(filepath, "w", newline="", encoding="utf-8") as f:
                csv.writer(f).writerows(rows)
            win.destroy()
            messagebox.showinfo("Success", "CSV template downloaded.")

        def download_json():
            filepath = filedialog.asksaveasfilename(
                defaultextension=".json", filetypes=[("JSON File", "*.json")],
                initialfile="questions_template.json")
            if not filepath:
                return
            import json
            data = [
                {"type": "mcq", "question": "What is polymorphism?",
                 "option_a": "Inheritance", "option_b": "Method Overloading",
                 "option_c": "Abstraction", "option_d": "Encapsulation",
                 "correct_answer": "B", "keywords": "", "ideal_answer": ""},
                {"type": "technical", "question": "Explain deadlock in OS.",
                 "option_a": "", "option_b": "", "option_c": "", "option_d": "",
                 "correct_answer": "", "keywords": "deadlock, locking, process",
                 "ideal_answer": "Should explain circular wait."},
                {"type": "hr", "question": "Tell me about a challenge you faced.",
                 "option_a": "", "option_b": "", "option_c": "", "option_d": "",
                 "correct_answer": "", "keywords": "teamwork, communication",
                 "ideal_answer": "Should explain challenge and learning."}
            ]
            with open(filepath, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=4)
            win.destroy()
            messagebox.showinfo("Success", "JSON template downloaded.")

        def download_docx():
            filepath = filedialog.asksaveasfilename(
                defaultextension=".docx", filetypes=[("Word File", "*.docx")],
                initialfile="questions_template.docx")
            if not filepath:
                return
            doc = Document()
            doc.add_heading("Question Import Template", level=1)
            doc.add_paragraph("Separate each question with ----")
            doc.add_paragraph(
                "TYPE: mcq\nQUESTION: What is polymorphism?\n"
                "A: Inheritance\nB: Method Overloading\nC: Abstraction\n"
                "D: Encapsulation\nANSWER: B\nKEYWORDS:\nIDEAL_ANSWER:\n----\n"
                "TYPE: technical\nQUESTION: Explain deadlock in OS.\n"
                "KEYWORDS: deadlock, locking, process\n"
                "IDEAL_ANSWER: Should explain circular wait.\n----\n"
                "TYPE: hr\nQUESTION: Tell me about a challenge you faced.\n"
                "KEYWORDS: teamwork, communication\n"
                "IDEAL_ANSWER: Should explain challenge and learning.\n----"
            )
            doc.save(filepath)
            win.destroy()
            messagebox.showinfo("Success", "DOCX template downloaded.")

        btn_row = ctk.CTkFrame(win, fg_color="transparent")
        btn_row.pack(pady=10)
        self._btn(btn_row, "CSV Template",  download_csv,  width=120).grid(
            row=0, column=0, padx=8)
        self._btn(btn_row, "JSON Template", download_json, width=120).grid(
            row=0, column=1, padx=8)
        self._btn(btn_row, "DOCX Template", download_docx, width=120).grid(
            row=0, column=2, padx=8)

    def _build_bulk_upload(self, parent):
        parent.columnconfigure(0, weight=0)
        parent.columnconfigure(1, weight=0)
        parent.columnconfigure(2, weight=1)

        self.bulk_role_var, cb = self._role_combo_widget(parent)
        self._grid_label(parent, "Job role:", 0)
        cb.grid(row=0, column=1, columnspan=2, padx=(12, 0), pady=(8, 12), sticky="ew")

        self.bulk_diff_var = tk.StringVar(value="Easy")
        self._grid_label(parent, "Difficulty:", 1)
        diff_cb = ctk.CTkComboBox(parent, variable=self.bulk_diff_var,
                                   values=["Easy", "Medium", "Hard"],
                                   width=180, state="readonly", font=F("md"))
        diff_cb.grid(row=1, column=1, padx=(12, 0), pady=(6, 12), sticky="w")

        self._grid_label(parent, "Question file:", 2)
        self._btn(parent, "Choose CSV / JSON / DOCX", self._pick_bulk_file,
                  width=190).grid(row=2, column=1, padx=(12, 10), pady=8, sticky="w")
        self.bulk_file_label = ctk.CTkLabel(parent, text="No file selected",
                                             font=F("sm"), text_color=TEXT_MUTED,
                                             anchor="w")
        self.bulk_file_label.grid(row=2, column=2, padx=(8, 0), pady=8, sticky="w")

        self.bulk_replace_var = tk.BooleanVar(value=False)
        ctk.CTkCheckBox(
            parent,
            text="Replace all existing questions for this role + difficulty",
            variable=self.bulk_replace_var, font=F("sm"),
        ).grid(row=3, column=1, columnspan=2, padx=(12, 0), pady=(10, 16), sticky="w")

        hint = (
            "Upload one file per role + difficulty (.csv, .json, .docx).\n"
            "Columns: type, question, option_a, option_b,\n"
            "option_c, option_d, correct_answer,\n"
            "keywords, ideal_answer\n"
            "Difficulty is selected above.\n"
            "Template: assets/questions_template.csv"
        )
        ctk.CTkLabel(parent, text=hint, font=F("xs"), text_color=TEXT_MUTED,
                     justify="left", wraplength=520).grid(
            row=4, column=1, columnspan=2, padx=(12, 0), pady=(4, 18), sticky="w")

        btn_row = ctk.CTkFrame(parent, fg_color="transparent")
        btn_row.grid(row=5, column=1, columnspan=2, padx=(12, 0), pady=8, sticky="w")
        self._btn(btn_row, "Upload questions", self._run_bulk_upload,
                  color=AMBER, width=170).grid(row=0, column=0, padx=(0, 12))
        self._ghost_btn(btn_row, "Open Template", self._open_bulk_template,
                        width=150).grid(row=0, column=1, padx=(0, 12))
        self._ghost_btn(btn_row, "Download Template", self._download_template,
                        width=170).grid(row=0, column=2)

        self._bulk_filepath = None

    def _pick_bulk_file(self):
        path = filedialog.askopenfilename(
            title="Select question bank",
            filetypes=[
                ("Question Files", "*.csv *.json *.docx"),
                ("CSV files", "*.csv"),
                ("JSON files", "*.json"),
                ("Word files", "*.docx"),
            ],
        )
        if path:
            self._bulk_filepath = path
            self.bulk_file_label.configure(text=os.path.basename(path))

    def _open_bulk_template(self):
        import subprocess
        template = "assets/questions_template.csv"
        if not os.path.exists(template):
            messagebox.showerror("Template Missing",
                                 "questions_template.csv not found in assets folder.")
            return
        try:
            if sys.platform.startswith("win"):
                os.startfile(template)
            elif sys.platform == "darwin":
                subprocess.Popen(["open", template])
            else:
                subprocess.Popen(["xdg-open", template])
        except Exception as exc:
            messagebox.showerror("Open Failed", f"Could not open template:\n{exc}")

    def _run_bulk_upload(self):
        role_id = self._get_role_id_from_var(self.bulk_role_var)
        if not role_id:
            messagebox.showwarning("Role", "Select a valid job role.")
            return

        difficulty = self.bulk_diff_var.get().strip()
        if difficulty not in ["Easy", "Medium", "Hard"]:
            messagebox.showwarning("Difficulty", "Select a valid difficulty level.")
            return

        filepath = getattr(self, "_bulk_filepath", None)
        if not filepath:
            messagebox.showwarning("File", "Choose a CSV, JSON, or DOCX file first.")
            return

        ext = filepath.lower().split(".")[-1]
        if ext not in {"csv", "json", "docx"}:
            messagebox.showerror("Invalid File", "Supported formats: CSV, JSON, DOCX")
            return

        replace = self.bulk_replace_var.get()
        if replace and not messagebox.askyesno(
            "Replace questions",
            f"This will delete ALL existing '{difficulty}' questions for the selected role "
            f"and import from the file.\n\nOther difficulties are NOT affected.\n\nContinue?"
        ):
            return

        try:
            ok, msg = admin.import_questions_from_file(
                role_id, filepath,
                difficulty=difficulty,
                replace_existing=replace,
            )
            if ok:
                self._refresh_questions()
                self._bulk_filepath = None
                if hasattr(self, "bulk_file_label"):
                    self.bulk_file_label.configure(text="No file selected")
                # Reset the checkbox so the next upload doesn't accidentally replace
                if hasattr(self, "bulk_replace_var"):
                    self.bulk_replace_var.set(False)
            messagebox.showinfo("Bulk Upload", msg)
        except Exception as exc:
            messagebox.showerror("Bulk Upload Failed", f"Error: {exc}")

    def _build_q_list(self, parent):
        _style_tree()
        tree_frame = tk.Frame(parent, bg="#1e1e2e")
        tree_frame.pack(fill="both", expand=True, pady=(0, 10))

        vsb = ttk.Scrollbar(tree_frame, orient="vertical")
        vsb.pack(side="right", fill="y")

        self.q_tree = ttk.Treeview(
            tree_frame,
            columns=("id", "role", "type", "diff", "question"),
            show="headings", height=14,
            yscrollcommand=vsb.set,
        )
        vsb.configure(command=self.q_tree.yview)
        for col, h, w in [
            ("id",       "ID",       40),
            ("role",     "Role",    140),
            ("type",     "Type",     90),
            ("diff",     "Diff",     60),
            ("question", "Question", 300),
        ]:
            self.q_tree.heading(col, text=h)
            self.q_tree.column(col, width=w)
        self.q_tree.pack(side="left", fill="both", expand=True)

        bf = ctk.CTkFrame(parent, fg_color="transparent")
        bf.pack(anchor="w")
        self._ghost_btn(bf, "Refresh", self._refresh_questions,
                        width=100).pack(side="left", padx=(0, 8))
        self._btn(bf, "Delete", self._delete_question,
                  color=DANGER, width=100).pack(side="left")
        self._refresh_questions()

    def _grid_labeled_entry(self, parent, label, row):
        self._grid_label(parent, label, row)
        e = self._entry(parent, width=340)
        e.grid(row=row, column=1, pady=5, sticky="ew")
        return e

    def _add_mcq(self):
        ok, msg = admin.add_mcq_question(
            self._get_role_id_from_var(self.mcq_role_var),
            self.mcq_diff.get(),
            self.mcq_q.get(), self.mcq_a.get(), self.mcq_b.get(),
            self.mcq_c.get(), self.mcq_d.get(), self.mcq_ans.get(),
        )
        messagebox.showinfo("MCQ", msg)
        if ok: self._refresh_questions()

    def _add_technical(self):
        ok, msg = admin.add_technical_question(
            self._get_role_id_from_var(self.tech_role_var),
            self.tech_diff.get(), self.tech_q.get(), self.tech_kw.get(),
        )
        messagebox.showinfo("Technical", msg)
        if ok: self._refresh_questions()

    def _add_hr(self):
        ok, msg = admin.add_hr_question(
    self._get_role_id_from_var(self.hr_role_var),
    self.hr_diff.get(),
    self.hr_q.get(),
    self.hr_kw.get(),
)
        messagebox.showinfo("HR", msg)
        if ok: self._refresh_questions()

    def _refresh_questions(self):
        if not hasattr(self, "q_tree") or not self.q_tree.winfo_exists():
            return
        for i in self.q_tree.get_children():
            self.q_tree.delete(i)
        for q in admin.view_questions():
            self.q_tree.insert("", "end", values=(
                q["id"], q["role_name"], q["question_type"],
                q["difficulty"] or "—", (q["question"] or "")[:55],
            ))

    def _delete_question(self):
        sel = self.q_tree.selection()
        if not sel: return
        qid = self.q_tree.item(sel[0])["values"][0]
        if messagebox.askyesno("Confirm", "Delete this question?"):
            ok, msg = admin.delete_question(qid)
            self._refresh_questions()
            messagebox.showinfo("Question", msg)

    # ------------------------------------------------------------------ #
    #  Admin — analytics panel  (UPDATED)                               #
    # ------------------------------------------------------------------ #

    def _admin_analytics_panel(self, parent):
        stats     = admin.get_dashboard_stats()
        avg       = admin.get_average_scores()
        role_perf = admin.get_role_wise_performance()
        recent    = admin.get_recent_interviews()

        # ── Top stats cards ──────────────────────────────────────────────
        cards = ctk.CTkFrame(parent, fg_color="transparent")
        cards.pack(fill="x", pady=(0, 20))
        for i in range(4):
            cards.columnconfigure(i, weight=1)
        for col, (label, val) in enumerate([
            ("Students",   stats["students"]),
            ("Roles",      stats["roles"]),
            ("Questions",  stats["questions"]),
            ("Interviews", stats["interviews"]),
        ]):
            box = ctk.CTkFrame(cards, corner_radius=10)
            box.grid(row=0, column=col, padx=8, sticky="ew")
            ctk.CTkLabel(box, text=label, font=F("sm", "bold"),
                         text_color=TEXT_MUTED).pack(pady=(10, 2))
            ctk.CTkLabel(box, text=str(val),
                         font=F("xxl", "bold")).pack(pady=(0, 12))

        # ── Average scores ───────────────────────────────────────────────
        score_box = ctk.CTkFrame(parent, corner_radius=10)
        score_box.pack(fill="x", pady=(0, 20))
        ctk.CTkLabel(score_box, text="Average Scores",
                     font=F("lg", "bold")).pack(anchor="w", padx=16, pady=(12, 8))
        for label, val in [
            ("MCQ",       avg["mcq"]),
            ("Technical", avg["technical"]),
            ("HR",        avg["hr"]),
            ("Overall",   avg["overall"]),
        ]:
            ctk.CTkLabel(score_box,
                         text=f"{label}: {val}%").pack(anchor="w", padx=18, pady=2)
        ctk.CTkFrame(score_box, height=10, fg_color="transparent").pack()

        # ── Role-wise performance ────────────────────────────────────────
        ctk.CTkLabel(parent, text="Role Performance",
                     font=F("lg", "bold")).pack(anchor="w", pady=(10, 8))
        _style_tree()
        role_tree = ttk.Treeview(
            parent,
            columns=("role", "attempts", "score"),
            show="headings", height=8,
        )
        role_tree.heading("role",     text="Role")
        role_tree.heading("attempts", text="Attempts")
        role_tree.heading("score",    text="Avg Score")
        role_tree.column("role",     width=220)
        role_tree.column("attempts", width=120, anchor="center")
        role_tree.column("score",    width=120, anchor="center")
        role_tree.pack(fill="x", pady=(0, 20))
        for row in role_perf:
            role_tree.insert("", "end", values=(
                row["role_name"],
                row["total_attempts"],
                round(row["avg_score"] or 0, 2),
            ))

        # ── Recent interviews ────────────────────────────────────────────
        ctk.CTkLabel(parent, text="Recent Interviews",
                     font=F("lg", "bold")).pack(anchor="w", pady=(10, 8))
        recent_tree = ttk.Treeview(
            parent,
            columns=("name", "role", "score", "date"),
            show="headings", height=8,
        )
        recent_tree.heading("name",  text="Student")
        recent_tree.heading("role",  text="Role")
        recent_tree.heading("score", text="Score")
        recent_tree.heading("date",  text="Date")
        recent_tree.column("name",  width=180)
        recent_tree.column("role",  width=180)
        recent_tree.column("score", width=100, anchor="center")
        recent_tree.column("date",  width=160, anchor="center")
        recent_tree.pack(fill="both", expand=True)
        for row in recent:
            recent_tree.insert("", "end", values=(
                row["name"],
                row["role_name"],
                row["overall_score"],
                row["created_at"],
            ))

    def _clear(self):
        for w in self.winfo_children(): w.destroy()


#---------------------------------------------------------------------------
# Seed & entry point
#---------------------------------------------------------------------------
def _seed_questions_for_role(role_id, role_name):
    import database
    conn = database.get_connection()
    cur = conn.cursor()
    cur.execute("SELECT COUNT(*) FROM questions WHERE role_id = ?", (role_id,))
    if cur.fetchone()[0] > 0:
        conn.close()
        return

    if role_name == "Software Developer":
        admin.add_mcq_question(
            role_id, "Easy",
            "What does OOP stand for?",
            "Object Oriented Programming", "Open Object Protocol",
            "Operational Output Process", "Ordered Operand Pointer", "A",
        )
        admin.add_mcq_question(
            role_id, "Easy",
            "Which keyword defines a function in Python?",
            "func", "def", "function", "fn", "B",
        )
        admin.add_technical_question(
            role_id, "Easy",
            "Explain the concept of inheritance in OOP.",
            "inheritance,class,object,parent,child",
        )
        admin.add_hr_question(
            role_id,
            "Tell us about a time you worked in a team.",
            "teamwork,collaboration,communication,goals,learning",
        )
    elif role_name == "Web Developer":
        admin.add_mcq_question(
            role_id, "Easy",
            "Which language runs in the browser?",
            "Python", "JavaScript", "Java", "C++", "B",
        )
        admin.add_technical_question(
            role_id, "Easy",
            "What is the difference between HTML and CSS?",
            "structure,style,markup,presentation,layout",
        )
        admin.add_hr_question(
            role_id,
            "How do you handle tight deadlines on a project?",
            "prioritize,communication,deadline,planning,team",
        )
    elif role_name == "Data Analyst":
        admin.add_mcq_question(
            role_id, "Easy",
            "Which tool is commonly used for data visualization?",
            "Git", "Tableau", "Docker", "Nginx", "B",
        )
        admin.add_mcq_question(
            role_id, "Easy",
            "What does SQL primarily manage?",
            "Images", "Relational data", "Network packets", "Audio", "B",
        )
        admin.add_technical_question(
            role_id, "Easy",
            "What is the difference between mean and median?",
            "average,outlier,central,skew,distribution",
        )
        admin.add_hr_question(
            role_id,
            "How do you explain insights to non-technical stakeholders?",
            "communication,simple,visual,story,business",
        )
    conn.close()


def seed_sample_data():
    import database
    conn = database.get_connection()
    cur  = conn.cursor()
    cur.execute("SELECT id FROM admins WHERE username = ?", ("admin",))
    if not cur.fetchone():
        auth.create_admin("admin", "admin123")
    cur.execute("SELECT COUNT(*) FROM roles")
    if cur.fetchone()[0] == 0:
        for role in ("Software Developer", "Web Developer", "Data Analyst"):
            admin.add_role(role)

    cur.execute("SELECT COUNT(*) FROM questions")
    if cur.fetchone()[0] == 0:
        roles = {r["role_name"]: r["id"] for r in admin.fetch_roles()}
        dev  = roles.get("Software Developer")
        web  = roles.get("Web Developer")
        data = roles.get("Data Analyst")

        if dev:
            admin.add_mcq_question(
                dev, "Easy", "What does OOP stand for?",
                "Object Oriented Programming", "Open Object Protocol",
                "Operational Output Process", "Ordered Operand Pointer", "A")
            admin.add_mcq_question(
                dev, "Easy", "Which keyword defines a function in Python?",
                "func", "def", "function", "fn", "B")
            admin.add_mcq_question(
                dev, "Medium", "Which data structure uses LIFO?",
                "Queue", "Stack", "Tree", "Graph", "B")
            admin.add_technical_question(
                dev, "Easy", "Explain the concept of inheritance in OOP.",
                "inheritance,class,object,parent,child")
            admin.add_technical_question(
                dev, "Medium", "What is polymorphism? Give an example.",
                "polymorphism,method,override,interface,class")
            admin.add_hr_question(
                dev, "Tell us about a time you worked in a team.",
                "teamwork,collaboration,communication,goals,learning")
            admin.add_hr_question(
                dev, "Where do you see yourself in five years?",
                "growth,learning,skills,goals,career")

        if web:
            admin.add_mcq_question(
                web, "Easy", "Which language runs in the browser?",
                "Python", "JavaScript", "Java", "C++", "B")
            admin.add_mcq_question(
                web, "Easy", "What does CSS stand for?",
                "Computer Style Sheets", "Cascading Style Sheets",
                "Creative Style System", "Colorful Style Sheets", "B")
            admin.add_technical_question(
                web, "Easy", "What is the difference between HTML and CSS?",
                "structure,style,markup,presentation,layout")
            admin.add_technical_question(
                web, "Medium", "Explain responsive web design.",
                "responsive,mobile,media,queries,layout")
            admin.add_hr_question(
                web, "How do you handle tight deadlines on a project?",
                "prioritize,communication,deadline,planning,team")
            admin.add_hr_question(
                web, "Describe a bug you fixed recently.",
                "debugging,testing,problem,solution,learning")

        if data:
            admin.add_mcq_question(
                data, "Easy", "Which tool is commonly used for data visualization?",
                "Git", "Tableau", "Docker", "Nginx", "B")
            admin.add_mcq_question(
                data, "Easy", "What does SQL primarily manage?",
                "Images", "Relational data", "Network packets", "Audio", "B")
            admin.add_mcq_question(
                data, "Medium", "Which measure shows data spread around the mean?",
                "Median", "Mode", "Standard deviation", "Count", "C")
            admin.add_technical_question(
                data, "Easy", "What is the difference between mean and median?",
                "average,outlier,central,skew,distribution")
            admin.add_technical_question(
                data, "Medium", "When would you use a JOIN in SQL?",
                "tables,relationship,keys,query,combine")
            admin.add_hr_question(
                data, "How do you explain insights to non-technical stakeholders?",
                "communication,simple,visual,story,business")
            admin.add_hr_question(
                data, "Tell us about a data project you are proud of.",
                "analysis,impact,metrics,learning,results")

    for role in admin.fetch_roles():
        _seed_questions_for_role(role["id"], role["role_name"])

    conn.close()


def main():
    ctk.set_appearance_mode("dark")
    ctk.set_default_color_theme("blue")
    init_db()
    seed_sample_data()
    app = InterviewApp()
    app.mainloop()


if __name__ == "__main__":
    main()