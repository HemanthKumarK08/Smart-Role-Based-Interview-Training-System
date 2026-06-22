
import os
import sys
import tkinter as tk
import tkinter.messagebox as messagebox
from tkinter import ttk, scrolledtext, filedialog, simpledialog
import customtkinter as ctk
from PIL import Image as PILImage

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from database import init_db
import auth
import admin
import student
import interview_engine
import feedback_engine

#---------------------------------------------------------------------------
# Design tokens
#---------------------------------------------------------------------------
FONT  = "Inter"
SZ    = dict(xs=10, sm=11, md=13, lg=15, xl=18, xxl=24, hero=32)

SB_BG      = "#0F172A"
SB_ACTIVE  = "#1E293B"
SB_ACCENT  = "#2563EB"
SB_TEXT    = "#F8FAFC"
SB_MUTED   = "#94A3B8"
SB_HOVER   = "#1E293B"
SB_BORDER  = "#334155"

CONTENT_BG   = ("#0B1120", "#0B1120") # Slate-950 background
CARD_BG      = ("#1E293B", "#1E293B") # Slate-800 card color
CARD_BORDER  = ("#334155", "#334155") # Slate-700 border color
ACCENT       = "#2563EB"              # Primary: blue-600
SECONDARY_BLUE = "#3B82F6"            # Secondary Blue
DARK_BLUE    = "#1E3A8A"              # Dark Blue: blue-900
SUCCESS      = "#10B981"              # Success Green
TEAL         = "#14B8A6"
PINK         = "#EC4899"
AMBER        = "#F59E0B"              # Warning
DANGER       = "#EF4444"
TEXT_PRIMARY = ("#F8FAFC", "#F8FAFC") # Text Primary: slate-50
TEXT_MUTED   = ("#94A3B8", "#94A3B8") # Text Muted: slate-400
TEXT_SECONDARY = ("#CBD5E1", "#CBD5E1") # Slate-300

STAT_PALETTES = [
    ("#1E293B", "#3B82F6", "#F8FAFC"), # Slate card, Blue accent, White text
    ("#1E293B", "#10B981", "#F8FAFC"), # Slate card, Green accent, White text
    ("#1E293B", "#F59E0B", "#F8FAFC"), # Slate card, Amber accent, White text
]

#---------------------------------------------------------------------------
# Helpers
#---------------------------------------------------------------------------
def F(size_key, weight="normal"):
    return ctk.CTkFont(family=FONT, size=SZ[size_key], weight=weight)


def _darken(hex_color, factor=0.85):
    if isinstance(hex_color, tuple):
        hex_color = hex_color[-1]
    if not isinstance(hex_color, str):
        return hex_color
    if not hex_color.startswith("#"):
        return hex_color
    hex_color = hex_color.lstrip("#")
    if len(hex_color) != 6:
        return "#" + hex_color
    try:
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        r = max(0, min(255, int(r * factor)))
        g = max(0, min(255, int(g * factor)))
        b = max(0, min(255, int(b * factor)))
        return f"#{r:02x}{g:02x}{b:02x}"
    except ValueError:
        return "#" + hex_color


def _style_tree():
    style = ttk.Style()
    try:
        style.theme_use("clam")
    except Exception:
        pass
    # Configure tree widget for a clean dark theme matching cards
    style.configure("Treeview", 
                    background="#1E293B", 
                    fieldbackground="#1E293B", 
                    foreground="#F8FAFC", 
                    rowheight=30, 
                    font=(FONT, 11), 
                    borderwidth=0)
    style.configure("Treeview.Heading", 
                    font=(FONT, 11, "bold"),
                    background="#2563EB", 
                    foreground="white")
    style.map("Treeview", background=[("selected", "#3B82F6")])


#---------------------------------------------------------------------------
# Application
#---------------------------------------------------------------------------
class InterviewApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Smart Interview Prep")
        self.geometry("1400x860")
        self.minsize(1200, 750)

        self._sidebar_frame   = None
        self._content_frame   = None
        self._sidebar_btns    = {}

        self._selected_role_id    = None
        self._selected_difficulty = None
        self._last_result_id      = None

        self.time_left = 0
        self.timer_job = None
        self.timer_label = None

        self._show_login_home()

    def _load_illustration(self, filename, size=(360, 360)):
        try:
            path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets", filename)
            pil_img = PILImage.open(path)
            return ctk.CTkImage(light_image=pil_img, dark_image=pil_img, size=size)
        except Exception as e:
            print(f"Error loading image {filename}: {e}")
            return None

    def _split_layout_shell(self, left_ratio=0.55):
        self._clear()
        
        # Root layout frame covering the entire window
        root = ctk.CTkFrame(self, fg_color=CONTENT_BG)
        root.pack(fill="both", expand=True)
        
        # Grid layout with column weights for split screen
        left_w = int(left_ratio * 100)
        right_w = 100 - left_w
        
        root.columnconfigure(0, weight=left_w)
        root.columnconfigure(1, weight=right_w)
        root.rowconfigure(0, weight=1)
        
        # Left side panel for branding and illustration
        left_frame = ctk.CTkFrame(root, fg_color=CONTENT_BG, corner_radius=0)
        left_frame.grid(row=0, column=0, sticky="nsew")
        
        # Right side panel for cards
        right_frame = ctk.CTkFrame(root, fg_color=CONTENT_BG, corner_radius=0)
        right_frame.grid(row=0, column=1, sticky="nsew")
        
        return left_frame, right_frame

    # ------------------------------------------------------------------ #
    #  Low-level widget factories                                        #
    # ------------------------------------------------------------------ #

    def _lbl(self, parent, text, size="md", weight="normal",
             color=None, anchor="w", **kw):
        return ctk.CTkLabel(
            parent, text=text, font=F(size, weight),
            text_color=color or TEXT_PRIMARY,
            anchor=anchor, **kw,
        )

    def _entry(self, parent, show=None, width=320, placeholder=""):
        opts = dict(width=width, font=F("md"),
                    border_color=CARD_BORDER, corner_radius=8)
        if show:        opts["show"]             = show
        if placeholder: opts["placeholder_text"] = placeholder
        return ctk.CTkEntry(parent, **opts)

    def _btn(self, parent, text, command, color=ACCENT,
             hover=None, width=140, height=38, **kw):
        return ctk.CTkButton(
            parent, text=text, command=command,
            font=F("sm", "bold"),
            fg_color=color,
            hover_color=hover or _darken(color),
            corner_radius=8,
            width=width, height=height, **kw,
        )

    def _ghost_btn(self, parent, text, command, width=120, height=34):
        return ctk.CTkButton(
            parent, text=text, command=command,
            font=F("sm"), fg_color="transparent",
            hover_color=("gray85", "gray25"),
            text_color=TEXT_MUTED,
            border_width=1, border_color=CARD_BORDER,
            corner_radius=8, width=width, height=height,
        )

    def _divider(self, parent):
        ctk.CTkFrame(parent, height=1,
                     fg_color=CARD_BORDER).pack(fill="x", pady=12)

    def _form_row(self, parent, label_text, show=None, width=340):
        block = ctk.CTkFrame(parent, fg_color="transparent")
        block.pack(fill="x", pady=6)
        self._lbl(block, label_text, size="sm", weight="bold",
                  color=TEXT_MUTED).pack(anchor="w", pady=(0, 4))
        e = self._entry(block, show=show, width=width)
        e.pack(fill="x")
        return e

    def _section_title(self, parent, text, subtitle=None):
        self._lbl(parent, text, size="xl", weight="bold").pack(
            anchor="w", pady=(0, 2))
        if subtitle:
            self._lbl(parent, subtitle, size="sm",
                      color=TEXT_MUTED).pack(anchor="w", pady=(0, 16))

    # ------------------------------------------------------------------ #
    #  Card primitives                                                   #
    # ------------------------------------------------------------------ #

    def _card_frame(self, parent, padx=20, pady=20, **kw):
        outer = ctk.CTkFrame(
            parent, corner_radius=12,
            fg_color=CARD_BG,
            border_width=1, border_color=CARD_BORDER, **kw,
        )
        inner = ctk.CTkFrame(outer, fg_color="transparent")
        inner.pack(fill="both", expand=True, padx=padx, pady=pady)
        return outer, inner

    def _action_card(self, parent, emoji, title, subtitle, command,
                     accent=ACCENT, width=220, height=130):
        outer = ctk.CTkFrame(
            parent, corner_radius=12,
            fg_color=CARD_BG,
            border_width=1, border_color=CARD_BORDER,
            width=width, height=height, cursor="hand2",
        )
        outer.pack_propagate(False)
        ctk.CTkFrame(outer, fg_color=accent,
                     height=3, corner_radius=0).place(x=0, y=0, relwidth=1)
        inner = ctk.CTkFrame(outer, fg_color="transparent")
        inner.pack(fill="both", expand=True, padx=16, pady=14)
        ctk.CTkLabel(inner, text=emoji, font=F("xl"), anchor="w").pack(anchor="w")
        ctk.CTkLabel(inner, text=title, font=F("md", "bold"),
                     anchor="w").pack(anchor="w", pady=(4, 1))
        ctk.CTkLabel(inner, text=subtitle, font=F("xs"),
                     text_color=TEXT_MUTED, anchor="w").pack(anchor="w")

        def _hover_in(e):  outer.configure(border_color=accent)
        def _hover_out(e): outer.configure(border_color=CARD_BORDER)
        def _click(e):     command()

        for w in [outer, inner]:
            w.bind("<Enter>",    _hover_in)
            w.bind("<Leave>",    _hover_out)
            w.bind("<Button-1>", _click)

        self.after(50, lambda: [
            c.bind("<Button-1>", _click)
            for c in inner.winfo_children()
        ])
        return outer

    # ------------------------------------------------------------------ #
    #  Sidebar layout                                                    #
    # ------------------------------------------------------------------ #

    def _sidebar_layout(self, items, active_key,
                         logo_line1="Interview", logo_line2="Prep",
                         user_label=None):
        self._clear()
        root_row = ctk.CTkFrame(self, fg_color="transparent")
        root_row.pack(fill="both", expand=True)

        sb = ctk.CTkFrame(root_row, width=250, corner_radius=0, fg_color=SB_BG)
        sb.pack(side="left", fill="y")
        sb.pack_propagate(False)
        self._sidebar_frame = sb

        logo_wrap = ctk.CTkFrame(sb, fg_color="transparent")
        logo_wrap.pack(fill="x", padx=20, pady=(28, 24))
        dot = ctk.CTkFrame(logo_wrap, width=8, height=8,
                           corner_radius=4, fg_color=SB_ACCENT)
        dot.pack(side="left", pady=(4, 0))
        ctk.CTkFrame(logo_wrap, width=6, fg_color="transparent").pack(side="left")
        txt_col = ctk.CTkFrame(logo_wrap, fg_color="transparent")
        txt_col.pack(side="left")
        ctk.CTkLabel(txt_col, text=logo_line1,
                     font=ctk.CTkFont(family=FONT, size=15, weight="bold"),
                     text_color=SB_TEXT, anchor="w").pack(anchor="w")
        ctk.CTkLabel(txt_col, text=logo_line2,
                     font=ctk.CTkFont(family=FONT, size=11),
                     text_color=SB_MUTED, anchor="w").pack(anchor="w")

        ctk.CTkFrame(sb, height=1, fg_color=SB_BORDER).pack(
            fill="x", padx=20, pady=(0, 12))

        self._sidebar_btns = {}

        def _set_active(key):
            for k, btn in self._sidebar_btns.items():
                if k == key:
                    btn.configure(fg_color=SB_ACTIVE, text_color=SB_TEXT)
                else:
                    btn.configure(fg_color="transparent", text_color=SB_MUTED)

        for key, emoji, label, command in items:
            def _make_cmd(k, cmd):
                def _inner():
                    _set_active(k)
                    cmd()
                return _inner

            btn = ctk.CTkButton(
                sb,
                text=f"  {emoji}  {label}",
                command=_make_cmd(key, command),
                font=ctk.CTkFont(family=FONT, size=13),
                fg_color="transparent",
                hover_color=SB_HOVER,
                text_color=SB_MUTED,
                anchor="w",
                corner_radius=8,
                height=40,
                border_width=0,
            )
            btn.pack(fill="x", padx=12, pady=2)
            self._sidebar_btns[key] = btn

        _set_active(active_key)

        if user_label:
            ctk.CTkFrame(sb, height=1, fg_color=SB_BORDER).pack(
                fill="x", padx=20, pady=(0, 8), side="bottom")
            ctk.CTkLabel(
                sb, text=f"  👤  {user_label}",
                font=ctk.CTkFont(family=FONT, size=11),
                text_color=SB_MUTED, anchor="w",
            ).pack(fill="x", padx=16, pady=(0, 16), side="bottom")

        content_wrap = ctk.CTkFrame(root_row, fg_color=CONTENT_BG, corner_radius=0)
        content_wrap.pack(side="left", fill="both", expand=True)
        self._content_frame = content_wrap
        return content_wrap

    def _sidebar_nav(self, key, builder_fn):
        for k, btn in self._sidebar_btns.items():
            if k == key:
                btn.configure(fg_color=SB_ACTIVE, text_color=SB_TEXT)
            else:
                btn.configure(fg_color="transparent", text_color=SB_MUTED)
        for w in self._content_frame.winfo_children():
            w.destroy()
        builder_fn(self._content_frame)

    # ------------------------------------------------------------------ #
    #  Auth shell                                                        #
    # ------------------------------------------------------------------ #

    def _auth_shell(self, header_title, header_sub):
        self._clear()
        root = ctk.CTkFrame(self, fg_color="transparent")
        root.pack(fill="both", expand=True)

        left = ctk.CTkFrame(root, width=420, corner_radius=0, fg_color=SB_BG)
        left.pack(side="left", fill="y")
        left.pack_propagate(False)

        brand = ctk.CTkFrame(left, fg_color="transparent")
        brand.place(relx=0.5, rely=0.42, anchor="center")
        ctk.CTkLabel(brand, text="🎯", font=ctk.CTkFont(size=52)).pack()
        ctk.CTkLabel(brand, text="Interview\nPrep",
                     font=ctk.CTkFont(family=FONT, size=28, weight="bold"),
                     text_color=SB_TEXT, justify="center").pack(pady=(8, 4))
        ctk.CTkLabel(brand,
                     text="Role-based smart practice\nfor your dream job",
                     font=ctk.CTkFont(family=FONT, size=12),
                     text_color=SB_MUTED, justify="center").pack()

        right = ctk.CTkFrame(root, fg_color=CONTENT_BG, corner_radius=0)
        right.pack(side="left", fill="both", expand=True)

        scroll = ctk.CTkScrollableFrame(right, fg_color="transparent")
        scroll.pack(fill="both", expand=True, padx=60, pady=40)

        ctk.CTkLabel(scroll, text=header_title,
                     font=ctk.CTkFont(family=FONT, size=26, weight="bold"),
                     anchor="w").pack(anchor="w", pady=(0, 4))
        ctk.CTkLabel(scroll, text=header_sub,
                     font=ctk.CTkFont(family=FONT, size=12),
                     text_color=TEXT_MUTED, anchor="w").pack(
        anchor="w", pady=(0, 28))
        return scroll

    # ------------------------------------------------------------------ #
    #  Login / register screens                                          #
    # ------------------------------------------------------------------ #

    def _show_login_home(self):
        # ------------------------------------------------------------------ #
        #  Redesigned Welcome Page - Premium Dark Theme Split-Screen Layout  #
        # ------------------------------------------------------------------ #
        left, right = self._split_layout_shell(left_ratio=0.55)
        
        # Left Side (55%): Branding, Feature Badges, Stats & Illustration (Scrollable)
        l_scroll = ctk.CTkScrollableFrame(left, fg_color="transparent")
        l_scroll.pack(fill="both", expand=True, padx=40, pady=30)
        
        # Logo + Project Title
        header_row = ctk.CTkFrame(l_scroll, fg_color="transparent")
        header_row.pack(fill="x", anchor="w", pady=(0, 10))
        
        logo = ctk.CTkLabel(header_row, text="🎯", font=ctk.CTkFont(size=36))
        logo.pack(side="left", padx=(0, 15))
        
        title_label = ctk.CTkLabel(header_row, text="Smart Role-Based Interview\nTraining System", 
                                   font=F("xl", "bold"), text_color=TEXT_PRIMARY, justify="left", anchor="w")
        title_label.pack(side="left")
        
        # Tagline
        tagline = ctk.CTkLabel(l_scroll, text="Prepare • Practice • Perform", 
                               font=F("lg", "bold"), text_color=SECONDARY_BLUE, anchor="w")
        tagline.pack(anchor="w", pady=(5, 20))
        
        # Premium Illustration
        img = self._load_illustration("welcome_illustration.png", size=(360, 360))
        if img:
            img_lbl = ctk.CTkLabel(l_scroll, image=img, text="")
            img_lbl.pack(anchor="w", pady=(0, 20))
            
        # Core Features Section Title
        ctk.CTkLabel(l_scroll, text="Core Features", font=F("md", "bold"), text_color=TEXT_PRIMARY, anchor="w").pack(anchor="w", pady=(10, 8))
        
        feat_frame = ctk.CTkFrame(l_scroll, fg_color="transparent")
        feat_frame.pack(fill="x", pady=(0, 25))
        feat_frame.columnconfigure((0, 1), weight=1)
        
        features = [
            "Role-Based Interview Practice",
            "Multiple Difficulty Levels",
            "Coding Output Round",
            "Technical & HR Evaluation",
            "Instant Performance Reports",
            "Random Question Selection"
        ]
        
        for idx, feat in enumerate(features):
            r = idx // 2
            c = idx % 2
            f_box = ctk.CTkFrame(feat_frame, fg_color=CARD_BG, border_width=1, border_color=CARD_BORDER, corner_radius=10)
            f_box.grid(row=r, column=c, padx=5, pady=5, sticky="ew")
            
            lbl_check = ctk.CTkLabel(f_box, text="✓", font=F("md", "bold"), text_color=SUCCESS)
            lbl_check.pack(side="left", padx=(12, 6), pady=10)
            
            lbl_txt = ctk.CTkLabel(f_box, text=feat, font=F("sm", "bold"), text_color=TEXT_SECONDARY)
            lbl_txt.pack(side="left", padx=5, pady=10)
            
        # Statistics Panel
        ctk.CTkLabel(l_scroll, text="Overview Statistics", font=F("md", "bold"), text_color=TEXT_PRIMARY, anchor="w").pack(anchor="w", pady=(10, 8))
        
        stats_frame = ctk.CTkFrame(l_scroll, fg_color="transparent")
        stats_frame.pack(fill="x", pady=(0, 15))
        stats_frame.columnconfigure((0, 1, 2, 3), weight=1)
        
        stats = [
            ("500+", "Questions"),
            ("4 Rounds", "Interview"),
            ("5+", "Job Roles"),
            ("Analytics", "Performance")
        ]
        
        for idx, (val, label) in enumerate(stats):
            s_box = ctk.CTkFrame(stats_frame, fg_color=CARD_BG, border_width=1, border_color=CARD_BORDER, corner_radius=12)
            s_box.grid(row=0, column=idx, padx=5, pady=5, sticky="ew")
            
            ctk.CTkLabel(s_box, text=val, font=F("md", "bold"), text_color=SECONDARY_BLUE).pack(pady=(12, 2))
            ctk.CTkLabel(s_box, text=label, font=F("xs"), text_color=TEXT_MUTED).pack(pady=(0, 12))

        # Right Side (45%): Welcome Portals Card (Scrollable)
        r_scroll = ctk.CTkScrollableFrame(right, fg_color="transparent")
        r_scroll.pack(fill="both", expand=True, padx=30, pady=30)
        
        card_outer, card_body = self._card_frame(r_scroll, padx=30, pady=30)
        card_outer.pack(fill="both", expand=True, pady=(20, 10))
        
        ctk.CTkLabel(card_body, text="Welcome", font=F("xxl", "bold"), text_color=TEXT_PRIMARY).pack(anchor="w", pady=(10, 2))
        ctk.CTkLabel(card_body, text="Choose your portal", font=F("sm"), text_color=TEXT_MUTED).pack(anchor="w", pady=(0, 25))
        
        # Portal Action Buttons with Blue Gradient Look and hover animations
        self._btn(card_body, "👨🎓  Student Login", self._show_student_login,
                  color=ACCENT, hover=SECONDARY_BLUE, width=320, height=48).pack(anchor="w", pady=10)
                  
        self._btn(card_body, "👨🏫  Admin Login", self._show_admin_login,
                  color=DARK_BLUE, hover=ACCENT, width=320, height=48).pack(anchor="w", pady=10)
        
        self._divider(card_body)
        
        # New Registration Section
        ctk.CTkLabel(card_body, text="New Student?", font=F("sm", "bold"), text_color=TEXT_MUTED).pack(anchor="w", pady=(10, 4))
        self._btn(card_body, "📝  Register Now", self._show_student_register,
                  color=SUCCESS, hover=_darken(SUCCESS, 0.8), width=320, height=48).pack(anchor="w", pady=(2, 20))
                  
        # Footer
        footer_frame = ctk.CTkFrame(card_body, fg_color="transparent")
        footer_frame.pack(fill="x", side="bottom", pady=(40, 0))
        
        ctk.CTkLabel(footer_frame, text="©️ 2026 Smart Role-Based Interview Training System", 
                     font=F("xs", "bold"), text_color=TEXT_MUTED, justify="center").pack(fill="x")
        ctk.CTkLabel(footer_frame, text="Version 2.0 • Developed by Hemanth Kumar K", 
                     font=F("xs"), text_color=TEXT_MUTED, justify="center").pack(fill="x", pady=2)
        
        self.update_idletasks()

    def _show_student_login(self):
        # ------------------------------------------------------------------ #
        #  Redesigned Student Login Page - Split Screen Layout               #
        # ------------------------------------------------------------------ #
        left, right = self._split_layout_shell(left_ratio=0.55)
        
        # Left Side (55%): Welcome Info, Highlights & Flowchart (Scrollable)
        l_scroll = ctk.CTkScrollableFrame(left, fg_color="transparent")
        l_scroll.pack(fill="both", expand=True, padx=40, pady=30)
        
        self._ghost_btn(l_scroll, "←  Back", self._show_login_home, width=100, height=32).pack(anchor="w", pady=(0, 20))
        
        ctk.CTkLabel(l_scroll, text="Welcome Back! 👋", font=F("xxl", "bold"), text_color=TEXT_PRIMARY, anchor="w").pack(anchor="w", pady=(0, 10))
        
        subtexts = [
            "Practice Interviews.",
            "Improve Your Skills.",
            "Become Placement Ready."
        ]
        for sub in subtexts:
            ctk.CTkLabel(l_scroll, text=sub, font=F("md", "bold"), text_color=TEXT_SECONDARY, anchor="w").pack(anchor="w", pady=2)
            
        img = self._load_illustration("student_login_illustration.png", size=(340, 340))
        if img:
            img_lbl = ctk.CTkLabel(l_scroll, image=img, text="")
            img_lbl.pack(anchor="w", pady=(20, 20))
            
        # Highlights Checkmarks
        ctk.CTkLabel(l_scroll, text="Core Highlights", font=F("sm", "bold"), text_color=TEXT_MUTED, anchor="w").pack(anchor="w", pady=(10, 5))
        
        highlights = [
            "Personalized Interviews",
            "Random Questions",
            "Coding Output Round",
            "Instant Feedback",
            "Performance Reports"
        ]
        
        for hl in highlights:
            hl_row = ctk.CTkFrame(l_scroll, fg_color="transparent")
            hl_row.pack(fill="x", pady=3)
            ctk.CTkLabel(hl_row, text="✔", font=F("sm", "bold"), text_color=SECONDARY_BLUE).pack(side="left", padx=(5, 8))
            ctk.CTkLabel(hl_row, text=hl, font=F("sm", "bold"), text_color=TEXT_SECONDARY).pack(side="left")
            
        # Interview Process flowchart widget
        ctk.CTkLabel(l_scroll, text="Interview Process Flow", font=F("sm", "bold"), text_color=TEXT_MUTED, anchor="w").pack(anchor="w", pady=(20, 5))
        
        flow_frame = ctk.CTkFrame(l_scroll, fg_color="transparent")
        flow_frame.pack(fill="x", pady=10)
        
        flow_steps = [
            ("📝", "Register"),
            ("💼", "Choose Role"),
            ("⚙️", "Choose Diff"),
            ("🎯", "Attend"),
            ("📊", "Results")
        ]
        
        for idx, (emoji, step_txt) in enumerate(flow_steps):
            step_box = ctk.CTkFrame(flow_frame, fg_color=CARD_BG, border_width=1, border_color=CARD_BORDER, corner_radius=10, width=72, height=72)
            step_box.pack(side="left", padx=2)
            step_box.pack_propagate(False)
            
            ctk.CTkLabel(step_box, text=emoji, font=F("lg")).pack(pady=(8, 1))
            ctk.CTkLabel(step_box, text=step_txt, font=ctk.CTkFont(size=9, weight="bold"), text_color=TEXT_SECONDARY).pack()
            
            if idx < len(flow_steps) - 1:
                arrow = ctk.CTkLabel(flow_frame, text="➔", font=F("md", "bold"), text_color=SECONDARY_BLUE)
                arrow.pack(side="left", padx=4)

        # Right Side (45%): Premium Login Card (Scrollable)
        r_scroll = ctk.CTkScrollableFrame(right, fg_color="transparent")
        r_scroll.pack(fill="both", expand=True, padx=30, pady=30)
        
        card_outer, card_body = self._card_frame(r_scroll, padx=25, pady=25)
        card_outer.pack(fill="both", expand=True, pady=20)
        
        ctk.CTkLabel(card_body, text="Student Login", font=F("xl", "bold"), text_color=TEXT_PRIMARY).pack(anchor="w", pady=(0, 20))
        
        self.stu_email = self._form_row(card_body, "Email address", width=340)
        self.stu_pass  = self._form_row(card_body, "Password", show="*", width=340)
        
        # Options row (Show Password & Forgot Password)
        opts_row = ctk.CTkFrame(card_body, fg_color="transparent")
        opts_row.pack(fill="x", pady=10)
        
        show_pass_var = tk.IntVar(value=0)
        def toggle_pass():
            if show_pass_var.get() == 1:
                self.stu_pass.configure(show="")
            else:
                self.stu_pass.configure(show="*")
                
        chk_show = ctk.CTkCheckBox(opts_row, text="Show Password", variable=show_pass_var, 
                                   command=toggle_pass, font=F("sm"), text_color=TEXT_SECONDARY,
                                   checkbox_width=18, checkbox_height=18)
        chk_show.pack(side="left")
        
        def forgot_password_action():
            messagebox.showinfo("Password Recovery", 
                                "To recover or reset your password, please contact the System Administrator.")
                                
        forgot_btn = ctk.CTkButton(opts_row, text="Forgot Password?", command=forgot_password_action,
                                   font=F("sm", "bold"), fg_color="transparent", text_color=SECONDARY_BLUE,
                                   hover_color=("gray85", "gray25"), width=100, height=20)
        forgot_btn.pack(side="right")
        
        # Action Buttons
        self._btn(card_body, "🔑  Login", self._student_login,
                  color=ACCENT, hover=SECONDARY_BLUE, width=340, height=42).pack(anchor="w", pady=(20, 10))
                  
        self._divider(card_body)
        
        ctk.CTkLabel(card_body, text="Don't have an account?", font=F("sm"), text_color=TEXT_MUTED).pack(anchor="w", pady=(5, 2))
        
        self._btn(card_body, "📝  Register Now", self._show_student_register,
                  color=SUCCESS, hover=_darken(SUCCESS, 0.8), width=340, height=40).pack(anchor="w")
                  
        self.stu_email.focus_set()
        self.unbind("<Return>")
        self.bind("<Return>", lambda e: self._student_login())

    def _show_student_register(self):
        # ------------------------------------------------------------------ #
        #  Redesigned Student Registration Page - Split Layout               #
        # ------------------------------------------------------------------ #
        left, right = self._split_layout_shell(left_ratio=0.55)
        
        # Left Side (55%): Registration Details & Benefits (Scrollable)
        l_scroll = ctk.CTkScrollableFrame(left, fg_color="transparent")
        l_scroll.pack(fill="both", expand=True, padx=40, pady=30)
        
        self._ghost_btn(l_scroll, "←  Back", self._show_login_home, width=100, height=32).pack(anchor="w", pady=(0, 20))
        
        ctk.CTkLabel(l_scroll, text="Create Your Interview Profile 🚀", font=F("xxl", "bold"), text_color=TEXT_PRIMARY, anchor="w").pack(anchor="w", pady=(0, 10))
        
        img = self._load_illustration("register_illustration.png", size=(340, 340))
        if img:
            img_lbl = ctk.CTkLabel(l_scroll, image=img, text="")
            img_lbl.pack(anchor="w", pady=(0, 20))
            
        ctk.CTkLabel(l_scroll, text="Profile Benefits", font=F("sm", "bold"), text_color=TEXT_MUTED, anchor="w").pack(anchor="w", pady=(10, 5))
        
        benefits = [
            "Unlimited Interview Practice",
            "Coding Output Round",
            "Instant Feedback",
            "Performance Tracking",
            "Download Reports",
            "Multiple Job Roles"
        ]
        
        for ben in benefits:
            ben_row = ctk.CTkFrame(l_scroll, fg_color="transparent")
            ben_row.pack(fill="x", pady=3)
            ctk.CTkLabel(ben_row, text="✔", font=F("sm", "bold"), text_color=SUCCESS).pack(side="left", padx=(5, 8))
            ctk.CTkLabel(ben_row, text=ben, font=F("sm", "bold"), text_color=TEXT_SECONDARY).pack(side="left")

        # Right Side (45%): Premium Registration Card (Scrollable)
        r_scroll = ctk.CTkScrollableFrame(right, fg_color="transparent")
        r_scroll.pack(fill="both", expand=True, padx=30, pady=30)
        
        card_outer, card_body = self._card_frame(r_scroll, padx=25, pady=25)
        card_outer.pack(fill="both", expand=True, pady=20)
        
        ctk.CTkLabel(card_body, text="Student Registration", font=F("xl", "bold"), text_color=TEXT_PRIMARY).pack(anchor="w", pady=(0, 20))
        
        self.reg_name    = self._form_row(card_body, "Full name",       width=340)
        self.reg_email   = self._form_row(card_body, "Email address",   width=340)
        self.reg_pass    = self._form_row(card_body, "Password", show="*", width=340)
        self.reg_confirm = self._form_row(card_body, "Confirm password", show="*", width=340)
        
        # Show Password Toggle
        show_pass_var = tk.IntVar(value=0)
        def toggle_pass():
            show_val = "" if show_pass_var.get() == 1 else "*"
            self.reg_pass.configure(show=show_val)
            self.reg_confirm.configure(show=show_val)
            
        chk_show = ctk.CTkCheckBox(card_body, text="Show Passwords", variable=show_pass_var, 
                                   command=toggle_pass, font=F("sm"), text_color=TEXT_SECONDARY,
                                   checkbox_width=18, checkbox_height=18)
        chk_show.pack(anchor="w", pady=10)
        
        # Action Buttons
        self._btn(card_body, "📝  Register", self._register_student_action,
                  color=SUCCESS, hover=_darken(SUCCESS, 0.8), width=340, height=42).pack(anchor="w", pady=(15, 10))
                  
        self._divider(card_body)
        
        ctk.CTkLabel(card_body, text="Already have an account?", font=F("sm"), text_color=TEXT_MUTED).pack(anchor="w", pady=(5, 2))
        
        self._btn(card_body, "🔑  Login Here", self._show_student_login,
                  color=ACCENT, hover=SECONDARY_BLUE, width=340, height=40).pack(anchor="w")
                  
        self.reg_name.focus_set()

    def _show_admin_login(self):
        # ------------------------------------------------------------------ #
        #  Redesigned Admin Login Page - Split Screen Layout                 #
        # ------------------------------------------------------------------ #
        left, right = self._split_layout_shell(left_ratio=0.55)
        
        # Left Side (55%): Administrator Portal Info (Scrollable)
        l_scroll = ctk.CTkScrollableFrame(left, fg_color="transparent")
        l_scroll.pack(fill="both", expand=True, padx=40, pady=30)
        
        self._ghost_btn(l_scroll, "←  Back", self._show_login_home, width=100, height=32).pack(anchor="w", pady=(0, 20))
        
        ctk.CTkLabel(l_scroll, text="Administrator Portal 🛡", font=F("xxl", "bold"), text_color=TEXT_PRIMARY, anchor="w").pack(anchor="w", pady=(0, 10))
        
        img = self._load_illustration("admin_login_illustration.png", size=(340, 340))
        if img:
            img_lbl = ctk.CTkLabel(l_scroll, image=img, text="")
            img_lbl.pack(anchor="w", pady=(0, 20))
            
        ctk.CTkLabel(l_scroll, text="Management Categories", font=F("sm", "bold"), text_color=TEXT_MUTED, anchor="w").pack(anchor="w", pady=(10, 5))
        
        areas = [
            "Manage Roles",
            "Manage Questions",
            "Student Management",
            "Analytics Dashboard",
            "Reports",
            "Performance Monitoring"
        ]
        
        for area in areas:
            area_row = ctk.CTkFrame(l_scroll, fg_color="transparent")
            area_row.pack(fill="x", pady=3)
            ctk.CTkLabel(area_row, text="✔", font=F("sm", "bold"), text_color=PINK).pack(side="left", padx=(5, 8))
            ctk.CTkLabel(area_row, text=area, font=F("sm", "bold"), text_color=TEXT_SECONDARY).pack(side="left")

        # Right Side (45%): Premium Admin Login & Setup Card (Scrollable)
        r_scroll = ctk.CTkScrollableFrame(right, fg_color="transparent")
        r_scroll.pack(fill="both", expand=True, padx=30, pady=30)
        
        card_outer, card_body = self._card_frame(r_scroll, padx=25, pady=25)
        card_outer.pack(fill="both", expand=True, pady=10)
        
        ctk.CTkLabel(card_body, text="Admin Login", font=F("xl", "bold"), text_color=TEXT_PRIMARY).pack(anchor="w", pady=(0, 20))
        
        self.adm_user = self._form_row(card_body, "Username", width=340)
        self.adm_pass = self._form_row(card_body, "Password", show="*", width=340)
        
        # Show Password Toggle
        show_pass_var = tk.IntVar(value=0)
        def toggle_pass():
            if show_pass_var.get() == 1:
                self.adm_pass.configure(show="")
            else:
                self.adm_pass.configure(show="*")
                
        chk_show = ctk.CTkCheckBox(card_body, text="Show Password", variable=show_pass_var, 
                                   command=toggle_pass, font=F("sm"), text_color=TEXT_SECONDARY,
                                   checkbox_width=18, checkbox_height=18)
        chk_show.pack(anchor="w", pady=10)
        
        # Sign In Button
        self._btn(card_body, "🔑  Sign In", self._admin_login,
                  color=PINK, hover=_darken(PINK, 0.8), width=340, height=42).pack(anchor="w", pady=(10, 15))
                  
        self._divider(card_body)
        
        # First-Time Setup Section
        ctk.CTkLabel(card_body, text="First-time Setup", font=F("lg", "bold"), text_color=TEXT_PRIMARY).pack(anchor="w", pady=(5, 2))
        ctk.CTkLabel(card_body, text="No admin yet? Create one below.", font=F("sm"), text_color=TEXT_MUTED).pack(anchor="w", pady=(0, 15))
        
        self.new_adm_user = self._form_row(card_body, "New admin username", width=340)
        self.new_adm_pass = self._form_row(card_body, "New admin password", show="*", width=340)
        
        # Show Setup Password Toggle
        show_setup_pass_var = tk.IntVar(value=0)
        def toggle_setup_pass():
            if show_setup_pass_var.get() == 1:
                self.new_adm_pass.configure(show="")
            else:
                self.new_adm_pass.configure(show="*")
                
        chk_setup_show = ctk.CTkCheckBox(card_body, text="Show Password", variable=show_setup_pass_var, 
                                         command=toggle_setup_pass, font=F("sm"), text_color=TEXT_SECONDARY,
                                         checkbox_width=18, checkbox_height=18)
        chk_setup_show.pack(anchor="w", pady=10)
        
        self._btn(card_body, "🔑  Create Admin Account", self._create_admin,
                  color=("#7c3aed", "#6d28d9"), hover=("#6d28d9", "#5b21b6"), width=340, height=42).pack(anchor="w", pady=(10, 10))
                  
        self.adm_user.focus_set()
        self.unbind("<Return>")
        self.bind("<Return>", lambda e: self._admin_login())

    # ------------------------------------------------------------------ #
    #  Auth actions                                                      #

    def _student_login(self):
        try:
            email    = self.stu_email.get().strip()
            password = self.stu_pass.get()
        except (tk.TclError, AttributeError):
            messagebox.showerror("Error", "Please open Student Login first.")
            return
        if not email or not password:
            messagebox.showwarning("Missing fields", "Enter email and password.")
            return
        ok, msg = auth.student_login(email, password)
        if ok:
            self.unbind("<Return>")
            self._show_student_dashboard()
        else:
            messagebox.showerror("Login Failed", msg)

    def _register_student_action(self):
        try:
            name     = self.reg_name.get().strip()
            email    = self.reg_email.get().strip()
            password = self.reg_pass.get()
            confirm  = self.reg_confirm.get()
        except (tk.TclError, AttributeError):
            messagebox.showerror("Error", "Please open Register screen first.")
            return
        if not all([name, email, password, confirm]):
            messagebox.showwarning("Missing fields", "Fill in all fields.")
            return
        ok, msg = auth.student_register(name, email, password, confirm)
        if ok:
            messagebox.showinfo("Success", msg)
            self._show_student_login()
        else:
            messagebox.showerror("Registration Failed", msg)

    def _admin_login(self):
        try:
            username = self.adm_user.get().strip()
            password = self.adm_pass.get()
        except (tk.TclError, AttributeError):
            messagebox.showerror("Error", "Please open Admin Login first.")
            return
        if not username or not password:
            messagebox.showwarning("Missing fields", "Enter username and password.")
            return
        ok, msg = auth.admin_login(username, password)
        if ok:
            self.unbind("<Return>")
            self._show_admin_dashboard()
        else:
            messagebox.showerror("Login Failed", msg)

    def _create_admin(self):
        try:
            username = self.new_adm_user.get().strip()
            password = self.new_adm_pass.get()
        except (tk.TclError, AttributeError):
            messagebox.showerror("Error", "Please open Admin Login first.")
            return
        if not username or not password:
            messagebox.showwarning("Missing fields", "Enter username and password.")
            return
        ok, msg = auth.create_admin(username, password)
        if ok:
            messagebox.showinfo("Success", msg + "\nYou can now log in.")
        else:
            messagebox.showerror("Error", msg)

    def _logout(self):
        auth.logout()
        self.unbind("<Return>")
        self._show_login_home()

    # ------------------------------------------------------------------ #
    #  Student dashboard                                                 #
    # ------------------------------------------------------------------ #

    def _show_student_dashboard(self):
        profile = student.load_profile()
        uid     = auth.get_session()["user"]["id"]

        # ── CHANGE 1: Added "profile" item to sidebar ──────────────────
        content = self._sidebar_layout(
            items=[
                ("dashboard", "🏠", "Dashboard",      lambda: None),
                ("interview", "🎯", "Start Interview", lambda: None),
                ("reports",   "📊", "Reports",         lambda: None),
                ("profile",   "👤", "My Profile",      lambda: None),
                ("logout",    "→",  "Logout",          self._logout),
            ],
            active_key="dashboard",
            logo_line1="Interview",
            logo_line2="Prep",
            user_label=profile["name"],
        )

        def build_dashboard(parent):
            for w in parent.winfo_children(): w.destroy()
            scroll = ctk.CTkScrollableFrame(parent, fg_color="transparent")
            scroll.pack(fill="both", expand=True, padx=32, pady=28)

            self._lbl(scroll, f"Hi, {profile['name']} 👋",
                      size="xxl", weight="bold").pack(anchor="w")
            self._lbl(scroll, "Here's your interview summary.",
                      size="sm", color=TEXT_MUTED).pack(anchor="w", pady=(2, 20))

            history   = feedback_engine.view_history(uid)
            completed = len(history)
            avg_score = (
                round(sum(h["overall_score"] for h in history) / completed)
                if completed else 0
            )
            best_role = "—"
            if history:
                best      = max(history, key=lambda h: h["overall_score"])
                best_role = best["role_name"]

            stats_row = ctk.CTkFrame(scroll, fg_color="transparent")
            stats_row.pack(fill="x", pady=(0, 24))
            stats_row.columnconfigure((0, 1, 2), weight=1)

            for col_i, (label, value, sub) in enumerate([
                ("Completed Interviews", completed,        ""),
                ("Average Score",        f"{avg_score}%",  "across all rounds"),
                ("Best Role",            best_role,         "highest score"),
            ]):
                outer, inner = self._card_frame(stats_row, padx=18, pady=16)
                outer.grid(row=0, column=col_i, padx=(0, 12), sticky="ew")
                self._lbl(inner, label, size="xs", weight="bold",
                          color=TEXT_MUTED).pack(anchor="w")
                self._lbl(inner, str(value), size="xxl",
                          weight="bold").pack(anchor="w", pady=(4, 0))
                if sub:
                    self._lbl(inner, sub, size="xs", color=TEXT_MUTED).pack(anchor="w")

            self._lbl(scroll, "Quick actions", size="lg",
                      weight="bold").pack(anchor="w", pady=(4, 12))
            cards_row = ctk.CTkFrame(scroll, fg_color="transparent")
            cards_row.pack(anchor="w")

            for emoji, title, sub, cmd, accent in [
                ("🎯", "Start Interview", "Pick a role and begin",
                 lambda: self._sidebar_nav("interview", build_interview), ACCENT),
                ("📊", "View Reports",   "See your score history",
                 lambda: self._sidebar_nav("reports",   build_reports),   TEAL),
            ]:
                c = self._action_card(cards_row, emoji, title, sub,
                                      cmd, accent=accent, width=200, height=120)
                c.pack(side="left", padx=(0, 12))

            self._lbl(scroll, "Your profile", size="lg",
                      weight="bold").pack(anchor="w", pady=(24, 12))
            p_outer, p_inner = self._card_frame(scroll, padx=20, pady=16)
            p_outer.pack(fill="x")
            for field, val in [("Name",         profile["name"]),
                                ("Email",        profile["email"]),
                                ("Member since", profile["created_at"])]:
                row = ctk.CTkFrame(p_inner, fg_color="transparent")
                row.pack(fill="x", pady=5)
                self._lbl(row, field, size="xs", weight="bold",
                          color=TEXT_MUTED).pack(side="left", padx=(0, 12))
                self._lbl(row, val, size="sm").pack(side="left")

        def build_interview(parent):
            for w in parent.winfo_children(): w.destroy()
            scroll = ctk.CTkScrollableFrame(parent, fg_color="transparent")
            scroll.pack(fill="both", expand=True, padx=32, pady=28)

            self._section_title(scroll, "Start Interview",
                                "Pick a role and difficulty, then go.")

            roles      = admin.fetch_roles()
            self.role_var  = tk.StringVar()
            role_names = [f"{r['id']}: {r['role_name']}" for r in roles]

            outer, inner = self._card_frame(scroll, padx=24, pady=20)
            outer.pack(fill="x", pady=(0, 20))

            self._lbl(inner, "Job role", size="sm", weight="bold",
                      color=TEXT_MUTED).pack(anchor="w", pady=(0, 6))
            if role_names:
                self.role_var.set(role_names[0])
                ctk.CTkComboBox(
                    inner, variable=self.role_var, values=role_names,
                    width=380, state="readonly", font=F("md"),
                    corner_radius=8,
                ).pack(anchor="w")
            else:
                self._lbl(inner, "No roles yet — ask your admin.",
                          size="sm", color=TEXT_MUTED).pack(anchor="w")

            ctk.CTkFrame(inner, height=1, fg_color=CARD_BORDER).pack(fill="x", pady=16)

            self._lbl(inner, "Difficulty", size="sm", weight="bold",
                      color=TEXT_MUTED).pack(anchor="w", pady=(0, 8))
            self.diff_var = tk.StringVar(value="Easy")
            diff_row      = ctk.CTkFrame(inner, fg_color="transparent")
            diff_row.pack(anchor="w")
            for d, accent in [("Easy", SUCCESS), ("Medium", AMBER), ("Hard", DANGER)]:
                ctk.CTkRadioButton(
                    diff_row, text=f"  {d}",
                    variable=self.diff_var, value=d,
                    font=F("md"), fg_color=accent, hover_color=accent,
                ).pack(side="left", padx=(0, 20))

            ctk.CTkFrame(inner, height=1, fg_color=CARD_BORDER).pack(fill="x", pady=16)
            self._btn(inner, "🚀  Start Interview",
                      self._start_interview, width=200, height=42).pack(anchor="w")

        def build_reports(parent):
            for w in parent.winfo_children(): w.destroy()
            scroll = ctk.CTkScrollableFrame(parent, fg_color="transparent")
            scroll.pack(fill="both", expand=True, padx=32, pady=28)

            self._section_title(scroll, "Your Reports",
                                "All your past interview sessions.")

            history = feedback_engine.view_history(uid)
            _style_tree()
            cols = ("id", "role", "mcq", "tech", "coding", "hr", "overall", "date")
            tree = ttk.Treeview(scroll, columns=cols, show="headings", height=12)
            for col, heading, w in zip(
                cols,
                ("ID", "Role", "MCQ", "Tech", "Coding", "HR", "Overall", "Date"),
                (40,  180,    60,    60,    60,      60,   70,      120),
            ):
                tree.heading(col, text=heading)
                tree.column(col, width=w, anchor="center")
            tree.pack(fill="both", expand=True, pady=8)

            for h in history:
                has_coding_session = False
                try:
                    if h.get("session_data"):
                        snap = json.loads(h["session_data"])
                        if snap.get("coding") and len(snap["coding"]) > 0:
                            has_coding_session = True
                except Exception:
                    pass
                coding_val = f"{h['coding_score']}%" if has_coding_session else "—"
                
                tree.insert("", "end", values=(
                    h["id"], h["role_name"],
                    f"{h['mcq_score']}%", f"{h['technical_score']}%",
                    coding_val, f"{h['hr_score']}%", f"{h['overall_score']}%", h["created_at"],
                ))

            btn_row = ctk.CTkFrame(scroll, fg_color="transparent")
            btn_row.pack(anchor="w", pady=8)
            self._btn(btn_row, "View Feedback",
                      lambda: self._view_feedback(tree),
                      width=140).pack(side="left", padx=(0, 10))
            self._btn(btn_row, "Perfect answers",
                      lambda: self._show_perfect_answers(tree),
                      color=SUCCESS, width=150).pack(side="left", padx=(0, 10))
            self._btn(btn_row, "Export CSV",
                      lambda: self._export_report(tree),
                      color=TEAL, width=130).pack(side="left")

        # ── CHANGE 2: Student profile page ─────────────────────────────
        def build_profile(parent):
            for w in parent.winfo_children():
                w.destroy()

            scroll = ctk.CTkScrollableFrame(parent, fg_color="transparent")
            scroll.pack(fill="both", expand=True, padx=32, pady=28)

            # Reload fresh profile data every time the page is opened
            current_profile = student.load_profile()

            self._section_title(
                scroll,
                "My Profile",
                "Update your account details."
            )

            card, body = self._card_frame(scroll)
            card.pack(fill="x")

            name_entry = self._form_row(body, "Name", width=400)
            name_entry.insert(0, current_profile["name"])

            email_entry = self._form_row(body, "Email", width=400)
            email_entry.insert(0, current_profile["email"])

            password_entry = self._form_row(
                body, "New Password", show="*", width=400)

            confirm_entry = self._form_row(
                body, "Confirm Password", show="*", width=400)

            def save_profile():
                name    = name_entry.get().strip()
                email   = email_entry.get().strip()
                password = password_entry.get()
                confirm  = confirm_entry.get()

                if not name or not email:
                    messagebox.showwarning(
                        "Missing Fields",
                        "Name and Email are required."
                    )
                    return

                current_uid = auth.get_session()["user"]["id"]
                student.update_profile(current_uid, name, email)

                if password:
                    if password != confirm:
                        messagebox.showerror(
                            "Password Mismatch",
                            "Passwords do not match."
                        )
                        return
                    student.update_password(current_uid, password)

                messagebox.showinfo(
                    "Success",
                    "Profile updated successfully."
                )
                # Refresh dashboard so sidebar user label updates
                self._show_student_dashboard()

            self._btn(
                body,
                "Save Changes",
                save_profile,
                color=SUCCESS,
                width=180
            ).pack(anchor="w", pady=20)

        # ── Wire sidebar buttons ────────────────────────────────────────
        self._sidebar_btns["dashboard"].configure(
            command=lambda: self._sidebar_nav("dashboard", build_dashboard))
        self._sidebar_btns["interview"].configure(
            command=lambda: self._sidebar_nav("interview", build_interview))
        self._sidebar_btns["reports"].configure(
            command=lambda: self._sidebar_nav("reports",   build_reports))
        # ── CHANGE 3: Wire profile button ───────────────────────────────
        self._sidebar_btns["profile"].configure(
            command=lambda: self._sidebar_nav("profile",  build_profile))

        build_dashboard(content)

    # ------------------------------------------------------------------ #
    #  Feedback / export helpers                                         #
    # ------------------------------------------------------------------ #

    def _show_perfect_answers(self, tree_or_result_id):
        if isinstance(tree_or_result_id, int):
            result_id = tree_or_result_id
        else:
            tree = tree_or_result_id
            sel = tree.selection()
            if not sel:
                messagebox.showwarning("Select", "Select a report first.")
                return
            result_id = tree.item(sel[0])["values"][0]

        uid = auth.get_session()["user"]["id"]
        review = feedback_engine.get_session_review(result_id, uid)
        text = feedback_engine.format_session_review_text(review)

        win = ctk.CTkToplevel(self)
        win.title("Perfect answers — your session")
        win.geometry("720x520")
        self._lbl(win,
                  "Compare your answers with the correct / model responses.",
                  size="sm", color=TEXT_MUTED).pack(
            anchor="w", padx=14, pady=(12, 4))
        tb = ctk.CTkTextbox(win, wrap="word",
                            font=ctk.CTkFont(family=FONT, size=12))
        tb.pack(fill="both", expand=True, padx=12, pady=(0, 12))
        tb.insert("1.0", text)
        tb.configure(state="disabled")

    def _view_feedback(self, tree):
        sel = tree.selection()
        if not sel:
            messagebox.showwarning("Select", "Select a report first.")
            return
        rid     = tree.item(sel[0])["values"][0]
        uid     = auth.get_session()["user"]["id"]
        history = feedback_engine.view_history(uid)
        row     = next((h for h in history if h["id"] == rid), None)
        if row:
            win = ctk.CTkToplevel(self)
            win.title("Feedback Report")
            win.geometry("640x440")
            tb = ctk.CTkTextbox(win, wrap="word",
                                font=ctk.CTkFont(family=FONT, size=11))
            tb.pack(fill="both", expand=True, padx=12, pady=12)
            tb.insert("1.0", row["feedback"])
            tb.configure(state="disabled")
            btn_row = ctk.CTkFrame(win, fg_color="transparent")
            btn_row.pack(fill="x", padx=12, pady=(0, 12))
            self._btn(btn_row, "View perfect answers",
                      lambda: self._show_perfect_answers(rid),
                      color=SUCCESS, width=180).pack(side="left")

    def _export_report(self, tree):
        sel = tree.selection()
        if not sel:
            messagebox.showwarning("Select", "Select a report first.")
            return
        rid  = tree.item(sel[0])["values"][0]
        path = filedialog.asksaveasfilename(
            defaultextension=".csv", filetypes=[("CSV", "*.csv")])
        if path:
            ok, msg = feedback_engine.export_report(rid, path)
            if ok:
                messagebox.showinfo("Exported", f"Saved to:\n{msg}")
            else:
                messagebox.showerror("Error", msg)

    def _export_latest(self, result_id):
        path = filedialog.asksaveasfilename(
            defaultextension=".csv", filetypes=[("CSV", "*.csv")])
        if path:
            ok, msg = feedback_engine.export_report(result_id, path)
            if ok:
                messagebox.showinfo("Exported", f"Saved to:\n{msg}")

    # ------------------------------------------------------------------ #
    #  Start interview                                                   #
    # ------------------------------------------------------------------ #

    def _start_interview(self):
        role_sel = self.role_var.get().strip()
        if not role_sel or ":" not in role_sel:
            messagebox.showwarning(
                "Role", "No valid job role available. Ask admin to create one.")
            return
        try:
            role_id = int(role_sel.split(":")[0])
        except ValueError:
            messagebox.showwarning("Role", "Invalid role selected.")
            return

        role       = student.select_role(role_id)
        difficulty = student.select_difficulty(self.diff_var.get())
        if not role or not difficulty:
            messagebox.showerror("Error", "Invalid role or difficulty.")
            return

        session = auth.get_session()
        self.update_idletasks()
        ok, msg = interview_engine.start_interview(
            session["user"]["id"], role_id, role["role_name"], difficulty)
        if not ok:
            messagebox.showerror("Cannot Start", msg)
            return
        self._show_interview_round()

    def _start_timer(self, seconds, q):
        if self.timer_job:
            self.after_cancel(self.timer_job)
            self.timer_job = None
        self.time_left = seconds
        self._update_timer(q)

    def _update_timer(self, q):
        mins = self.time_left // 60
        secs = self.time_left % 60
        if self.timer_label:
            self.timer_label.configure(text=f"⏱ {mins:02}:{secs:02}")
        if self.time_left <= 0:
            self.timer_job = None
            self._time_up(q)
            return
        self.time_left -= 1
        self.timer_job = self.after(1000, lambda: self._update_timer(q))

    def _time_up(self, q):
        messagebox.showinfo("Time Up", "Time is over. Moving to next question.")
        sess = interview_engine.get_active_session()
        if not sess:
            return
        has_coding = getattr(sess, "has_coding_round", False)
        if sess.current_round == 1:
            interview_engine.submit_mcq_answer(q["id"], "")
        elif sess.current_round == 2:
            interview_engine.submit_technical_answer(q["id"], "", q["keywords"])
        elif sess.current_round == 3 and has_coding:
            interview_engine.submit_coding_answer(q["id"], "")
        else:
            interview_engine.submit_hr_answer(q["id"], "", q["keywords"])
        has_more = interview_engine.next_question()
        if has_more:
            self._show_interview_round()
        elif sess.current_round < (4 if has_coding else 3):
            self._advance_to_next_round()
        else:
            interview_engine.calculate_hr_score()
            self._show_results()

    def _advance_to_next_round(self):
        sess = interview_engine.get_active_session()
        has_coding = getattr(sess, "has_coding_round", False)
        if sess.current_round == 1:
            interview_engine.calculate_mcq_score()
        elif sess.current_round == 2:
            interview_engine.calculate_technical_score()
        elif sess.current_round == 3 and has_coding:
            interview_engine.calculate_coding_score()
        completed_round = sess.current_round
        interview_engine.advance_round()
        messagebox.showinfo("Next Round",
                            f"Round {completed_round} complete. Starting next round.")
        self._show_interview_round()

    def _show_interview_round(self):
        self._clear()
        sess = interview_engine.get_active_session()
        if not sess:
            messagebox.showerror("Session Expired",
                                 "No active interview session found.")
            self._show_student_dashboard()
            return

        has_coding = getattr(sess, "has_coding_round", False)
        if has_coding:
            round_meta = {
                1: ("MCQ Round", ACCENT, "#312e81"),
                2: ("Technical Round", TEAL, "#0d9488"),
                3: ("Coding Output Round", AMBER, "#78350f"),
                4: ("HR Round", PINK, "#9d174d"),
            }
        else:
            round_meta = {
                1: ("MCQ Round", ACCENT, "#312e81"),
                2: ("Technical Round", TEAL, "#0d9488"),
                3: ("HR Round", PINK, "#9d174d"),
            }
        rname, badge_fg, badge_hover = round_meta.get(
            sess.current_round, ("Complete", ACCENT, "#312e81"))

        root = ctk.CTkFrame(self, fg_color="transparent")
        root.pack(fill="both", expand=True)
        ctk.CTkFrame(root, width=6, corner_radius=0,
                     fg_color=badge_fg).pack(side="left", fill="y")
        main = ctk.CTkFrame(root, fg_color=CONTENT_BG, corner_radius=0)
        main.pack(side="left", fill="both", expand=True)

        topbar = ctk.CTkFrame(main, fg_color=("white", "#1e1e2e"),
                              height=60, corner_radius=0)
        topbar.pack(fill="x")
        topbar.pack_propagate(False)
        tb = ctk.CTkFrame(topbar, fg_color="transparent")
        tb.pack(fill="both", expand=True, padx=24, pady=10)
        ctk.CTkLabel(tb, text=rname,
                     font=ctk.CTkFont(family=FONT, size=16,
                                      weight="bold")).pack(side="left")
        ctk.CTkLabel(tb, text=f"  ·  {sess.role_name}  ·  {sess.difficulty}",
                     font=ctk.CTkFont(family=FONT, size=12),
                     text_color=TEXT_MUTED).pack(side="left")

        self.timer_label = ctk.CTkLabel(
            tb, text="⏱ 00:00",
            font=ctk.CTkFont(family=FONT, size=12, weight="bold"),
            text_color=DANGER)
        self.timer_label.pack(side="right", padx=(0, 16))

        self._ghost_btn(
            tb, "✕ Exit",
            lambda: (
                self.after_cancel(self.timer_job) if self.timer_job else None,
                setattr(self, "timer_job", None),
                interview_engine.reset_session(),
                self._show_student_dashboard()
            ),
            width=90, height=30
        ).pack(side="right")

        q = interview_engine.get_current_question()
        if not q:
            max_round = 4 if has_coding else 3
            if sess.current_round < max_round:
                self._advance_to_next_round()
                return
            else:
                interview_engine.calculate_hr_score()
                self._show_results()
                return

        if sess.current_round == 1:
            total = len(sess.mcq_questions)
        elif sess.current_round == 2:
            total = len(sess.technical_questions)
        elif sess.current_round == 3 and has_coding:
            total = len(sess.coding_questions)
        else:
            total = len(sess.hr_questions)
        idx = sess.current_index + 1

        prog = ctk.CTkProgressBar(main, fg_color=("gray85", "gray25"),
                                  progress_color=badge_fg, height=4, corner_radius=0)
        prog.set(idx / max(total, 1))
        prog.pack(fill="x")

        scroll = ctk.CTkScrollableFrame(main, fg_color="transparent")
        scroll.pack(fill="both", expand=True, padx=40, pady=24)

        q_outer, q_inner = self._card_frame(scroll, padx=28, pady=24)
        q_outer.pack(fill="x", pady=(0, 16))

        badge_frame = ctk.CTkFrame(q_inner, fg_color=badge_fg, corner_radius=6)
        badge_frame.pack(anchor="w", pady=(0, 14))
        ctk.CTkLabel(badge_frame, text=f"  Question {idx} of {total}  ",
                     font=ctk.CTkFont(family=FONT, size=10, weight="bold"),
                     text_color="white").pack(padx=6, pady=4)

        ctk.CTkLabel(q_inner, text=q["question"],
                     font=ctk.CTkFont(family=FONT, size=15),
                     wraplength=700, justify="left", anchor="w").pack(
            anchor="w", pady=(0, 20), fill="x")

        self.answer_var = tk.StringVar()

        if sess.current_round == 1:
            self._start_timer(30, q)
        elif sess.current_round == 2:
            self._start_timer(90, q)
        elif sess.current_round == 3 and has_coding:
            self._start_timer(30, q)
        else:
            self._start_timer(120, q)

        if sess.current_round == 3 and has_coding:
            code_outer = ctk.CTkFrame(q_inner, fg_color=("#f1f1f4", "#181825"), border_width=1, border_color=("gray75", "gray30"), corner_radius=8)
            code_outer.pack(fill="x", pady=(0, 16))
            code_label = ctk.CTkLabel(
                code_outer, text=q.get("code_snippet") or "",
                font=ctk.CTkFont(family="Courier", size=13),
                justify="left", anchor="w"
            )
            code_label.pack(padx=16, pady=12, fill="both", expand=True)

        if sess.current_round == 1 or (sess.current_round == 3 and has_coding):
            for opt, letter in [
                (q["option_a"], "A"), (q["option_b"], "B"),
                (q["option_c"], "C"), (q["option_d"], "D"),
            ]:
                opt_frame = ctk.CTkFrame(q_inner, fg_color=("gray97", "gray20"),
                                         corner_radius=8)
                opt_frame.pack(fill="x", pady=4)
                ctk.CTkRadioButton(
                    opt_frame, text=f"  {letter}.   {opt}",
                    variable=self.answer_var, value=letter,
                    font=ctk.CTkFont(family=FONT, size=13),
                    fg_color=badge_fg, hover_color=badge_hover,
                ).pack(anchor="w", padx=16, pady=10)
        else:
            self._lbl(q_inner, "Your answer", size="sm", weight="bold",
                      color=TEXT_MUTED).pack(anchor="w", pady=(0, 6))
            self.answer_text = ctk.CTkTextbox(
                q_inner, height=160, wrap="word",
                font=ctk.CTkFont(family=FONT, size=13), corner_radius=8)
            self.answer_text.pack(fill="x")

        self._btn(q_inner, "Submit answer →",
                  lambda q=q: self._submit_answer(q),
                  width=180, height=42).pack(anchor="w", pady=(20, 0))

    def _submit_answer(self, q):
        if self.timer_job:
            self.after_cancel(self.timer_job)
            self.timer_job = None

        sess = interview_engine.get_active_session()
        if not sess:
            return

        has_coding = getattr(sess, "has_coding_round", False)

        if sess.current_round == 1 or (sess.current_round == 3 and has_coding):
            ans = self.answer_var.get()
            if not ans:
                messagebox.showwarning("Answer", "Select an option.")
                return
            if sess.current_round == 1:
                interview_engine.submit_mcq_answer(q["id"], ans)
            else:
                interview_engine.submit_coding_answer(q["id"], ans)

        elif sess.current_round == 2:
            ans = self.answer_text.get("1.0", "end").strip()
            if not ans:
                messagebox.showwarning("Answer", "Enter your answer.")
                return
            result = interview_engine.submit_technical_answer(
                q["id"], ans, q["keywords"])
            messagebox.showinfo(
                "Feedback",
                f"Score: {result['score']}%\n{result['feedback']}\n"
                f"Keywords matched: {result['matched_keywords']}/{result['total_keywords']}")
        else:
            ans = self.answer_text.get("1.0", "end").strip()
            if not ans:
                messagebox.showwarning("Answer", "Enter your answer.")
                return
            result = interview_engine.submit_hr_answer(q["id"], ans, q["keywords"])
            messagebox.showinfo(
                "Feedback",
                f"Score: {result['final_score']}%\n{result['feedback']}")

        has_more = interview_engine.next_question()
        if has_more:
            self._show_interview_round()
        elif sess.current_round < (4 if has_coding else 3):
            self._advance_to_next_round()
        else:
            interview_engine.calculate_hr_score()
            self._show_results()

    # ------------------------------------------------------------------ #
    #  Results screen                                                    #
    # ------------------------------------------------------------------ #

    def _show_results(self):
        if self.timer_job:
            self.after_cancel(self.timer_job)
            self.timer_job = None

        report = feedback_engine.finalize_interview()
        if not report:
            messagebox.showerror("Error", "Could not generate report.")
            self._show_student_dashboard()
            return

        self._last_result_id = report["result_id"]
        self._clear()

        root = ctk.CTkFrame(self, fg_color=CONTENT_BG, corner_radius=0)
        root.pack(fill="both", expand=True)
        scroll = ctk.CTkScrollableFrame(root, fg_color="transparent")
        scroll.pack(fill="both", expand=True, padx=48, pady=36)

        self._lbl(scroll, "Interview complete! 🎉",
                  size="xxl", weight="bold").pack(anchor="w")
        self._lbl(scroll, f"{report['role_name']}  ·  {report['difficulty']}",
                  size="sm", color=TEXT_MUTED).pack(anchor="w", pady=(2, 24))

        sess = interview_engine.get_active_session()
        has_coding = getattr(sess, "has_coding_round", False)

        scorecards = [
            ("MCQ",     report["mcq_score"],       ACCENT),
            ("Tech",    report["technical_score"],  TEAL),
        ]
        if has_coding:
            scorecards.append(("Coding", report["coding_score"], AMBER))
        scorecards.extend([
            ("HR",      report["hr_score"],         PINK),
            ("Overall", report["overall_score"],    AMBER),
        ])

        scores_row = ctk.CTkFrame(scroll, fg_color="transparent")
        scores_row.pack(fill="x", pady=(0, 24))
        num_cards = len(scorecards)
        scores_row.columnconfigure(list(range(num_cards)), weight=1)

        for col_i, (label, val, color) in enumerate(scorecards):
            box = ctk.CTkFrame(scores_row, fg_color=color, corner_radius=12)
            box.grid(row=0, column=col_i, padx=(0, 12), sticky="ew")
            ctk.CTkLabel(box, text=label,
                         font=ctk.CTkFont(family=FONT, size=10, weight="bold"),
                         text_color="white").pack(pady=(14, 2))
            ctk.CTkLabel(box, text=f"{val}%",
                         font=ctk.CTkFont(family=FONT, size=28, weight="bold"),
                         text_color="white").pack(pady=(0, 14))

        self._lbl(scroll, "Detailed feedback", size="lg",
                  weight="bold").pack(anchor="w", pady=(4, 10))
        fb_outer, fb_inner = self._card_frame(scroll, padx=20, pady=16)
        fb_outer.pack(fill="x")
        tb = ctk.CTkTextbox(fb_inner, height=200, wrap="word",
                            font=ctk.CTkFont(family=FONT, size=12))
        tb.pack(fill="both", expand=True)
        tb.insert("1.0", report["feedback"])
        tb.configure(state="disabled")

        btn_row = ctk.CTkFrame(scroll, fg_color="transparent")
        btn_row.pack(anchor="w", pady=16)
        self._btn(btn_row, "Back to Dashboard",
                  lambda: (interview_engine.reset_session(), self._show_student_dashboard()), width=180).pack(
            side="left", padx=(0, 10))
        self._btn(btn_row, "View perfect answers",
                  lambda: self._show_perfect_answers(report["result_id"]),
                  color=SUCCESS, width=180).pack(side="left", padx=(0, 10))
        self._btn(btn_row, "Export CSV",
                  lambda: self._export_latest(report["result_id"]),
                  color=TEAL, width=140).pack(side="left")

    # ------------------------------------------------------------------ #
    #  Admin dashboard                                                   #
    # ------------------------------------------------------------------ #

    def _show_admin_dashboard(self):
        session = auth.get_session()

        # ── ADDED: "students" item to admin sidebar ────────────
        content = self._sidebar_layout(
            items=[
                ("roles",     "📋", "Roles",      lambda: None),
                ("questions", "❓", "Questions",  lambda: None),
                ("students",  "👨🎓", "Students",   lambda: None),
                ("analytics", "📈", "Analytics",  lambda: None),
                ("profile",   "👤", "My Profile", lambda: None),
                ("logout",    "→",  "Logout",     self._logout),
            ],
            active_key="roles",
            logo_line1="Admin",
            logo_line2="Console",
            user_label=session["user"]["username"],
        )

        def build_roles(parent):
            for w in parent.winfo_children(): w.destroy()
            scroll = ctk.CTkScrollableFrame(parent, fg_color="transparent")
            scroll.pack(fill="both", expand=True, padx=28, pady=24)
            self._section_title(scroll, "Job Roles", "Add, rename or remove roles.")
            form = ctk.CTkFrame(scroll, fg_color="transparent")
            form.pack(anchor="w", pady=(0, 16))
            self.new_role_entry = self._entry(form, width=280)
            self.new_role_entry.pack(side="left", padx=(0, 10))
            self.new_role_coding_var = tk.IntVar(value=0)
            ctk.CTkCheckBox(form, text="Include Coding Round", variable=self.new_role_coding_var, font=F("md")).pack(side="left", padx=(0, 15))
            self._btn(form, "+ Add Role", self._add_role,
                      color=SUCCESS, width=120).pack(side="left")

            _style_tree()
            self.roles_tree = ttk.Treeview(scroll, columns=("id", "name", "coding"),
                                           show="headings", height=10)
            self.roles_tree.heading("id",   text="ID")
            self.roles_tree.heading("name", text="Role Name")
            self.roles_tree.heading("coding", text="Coding Round")
            self.roles_tree.column("id",   width=50,  anchor="center")
            self.roles_tree.column("name", width=250)
            self.roles_tree.column("coding", width=140, anchor="center")
            self.roles_tree.pack(fill="x", pady=(0, 12))

            btns = ctk.CTkFrame(scroll, fg_color="transparent")
            btns.pack(anchor="w")
            self._ghost_btn(btns, "Refresh", self._refresh_roles,
                            width=100).pack(side="left", padx=(0, 8))
            self._btn(btns, "Rename/Edit", self._update_role,
                      width=110).pack(side="left", padx=(0, 8))
            self._btn(btns, "Delete", self._delete_role,
                      color=DANGER, width=100).pack(side="left")
            self._refresh_roles()

        def build_questions(parent):
            for w in parent.winfo_children(): w.destroy()
            scroll = ctk.CTkScrollableFrame(parent, fg_color="transparent")
            scroll.pack(fill="both", expand=True, padx=28, pady=24)
            self._section_title(scroll, "Question Bank",
                                "Add MCQ, Technical, Coding or HR questions.")
            self._admin_questions_panel(scroll)

        def build_analytics(parent):
            for w in parent.winfo_children(): w.destroy()
            
            # Header
            header = ctk.CTkFrame(parent, fg_color="transparent")
            header.pack(fill="x", padx=28, pady=(24, 0))
            
            self._section_title(header, "Analytics Console", "Real-time statistics, question summaries, activity logs, and performance metrics.")
            
            btn_container = ctk.CTkFrame(header, fg_color="transparent")
            btn_container.pack(side="right", anchor="n")
            
            self._btn(btn_container, "📥 Export Results", self._export_all_results,
                      color=TEAL, width=155).pack(side="left", padx=(0, 10))
            self._ghost_btn(btn_container, "🔄 Refresh", lambda: build_analytics(parent),
                            width=100).pack(side="left")
            
            # Content scroll frame
            scroll = ctk.CTkScrollableFrame(parent, fg_color="transparent")
            scroll.pack(fill="both", expand=True, padx=28, pady=(10, 24))
            
            self._admin_analytics_panel(scroll)

        # ── CHANGE 5: Admin profile page ────────────────────────────────
        def build_admin_profile(parent):
            for w in parent.winfo_children():
                w.destroy()

            scroll = ctk.CTkScrollableFrame(parent, fg_color="transparent")
            scroll.pack(fill="both", expand=True, padx=32, pady=28)

            # Get current admin from session
            current_admin = admin.load_admin_profile(
                auth.get_session()["user"]["id"]
            )

            self._section_title(
                scroll,
                "My Profile",
                "Update your admin account details."
            )

            card, body = self._card_frame(scroll)
            card.pack(fill="x")

            username_entry = self._form_row(
                body,
                "Username",
                width=400
            )
            username_entry.insert(0, current_admin["username"])

            password_entry = self._form_row(
                body,
                "New Password",
                show="*",
                width=400
            )

            confirm_entry = self._form_row(
                body,
                "Confirm Password",
                show="*",
                width=400
            )

            # def save_profile():
            #     username = username_entry.get().strip()
            #     password = password_entry.get()
            #     confirm = confirm_entry.get()

            #     if not username:
            #         messagebox.showwarning(
            #             "Missing Fields",
            #             "Username is required."
            #         )
            #         return

            #     admin_id = auth.get_session()["user"]["id"]

            #     admin.update_admin_profile(
            #         admin_id,
            #         username
            #     )

            #     if password:
            #         if password != confirm:
            #             messagebox.showerror(
            #                 "Password",
            #                 "Passwords do not match."
            #             )
            #             return

            #         admin.update_admin_password(
            #             admin_id,
            #             password
            #         )

            #     messagebox.showinfo(
            #         "Success",
            #         "Profile updated successfully."
            #     )

            #     # Refresh page
            #     self._show_admin_dashboard()

            # self._btn(
            #     body,
            #     "Save Changes",
            #     save_profile,
            #     color=SUCCESS,
            #     width=180
            # ).pack(anchor="w", pady=20)

            def save_admin_profile():
                username = username_entry.get().strip()
                password = password_entry.get()
                confirm  = confirm_entry.get()

                if not username:
                    messagebox.showwarning(
                        "Missing Fields",
                        "Username is required."
                    )
                    return

                current_uid = auth.get_session()["user"]["id"]
                ok, msg = admin.update_admin_profile(current_uid, username)
                if not ok:
                    messagebox.showerror("Update Failed", msg)
                    return

                if password:
                    if password != confirm:
                        messagebox.showerror(
                            "Password Mismatch",
                            "Passwords do not match."
                        )
                        return
                    p_ok, p_msg = admin.update_admin_password(current_uid, password)
                    if not p_ok:
                        messagebox.showerror("Password Error", p_msg)
                        return

                messagebox.showinfo(
                    "Success",
                    "Admin profile updated successfully."
                )
                # Refresh dashboard so sidebar user label updates
                self._show_admin_dashboard()

            self._btn(
                body,
                "Save Changes",
                save_admin_profile,
                color=SUCCESS,
                width=180
            ).pack(anchor="w", pady=20)

        def build_students(parent):
            for w in parent.winfo_children(): w.destroy()
            scroll = ctk.CTkScrollableFrame(parent, fg_color="transparent")
            scroll.pack(fill="both", expand=True, padx=28, pady=24)
            self._section_title(scroll, "Student Management", "View student profiles, interview history, or delete accounts.")
            
            # Search Frame
            search_row = ctk.CTkFrame(scroll, fg_color="transparent")
            search_row.pack(fill="x", pady=(0, 15))
            
            ctk.CTkLabel(search_row, text="Search Student:", font=F("sm", "bold"), text_color=TEXT_MUTED).pack(side="left", padx=(0, 10))
            self.student_search_entry = self._entry(search_row, width=280, placeholder="Name or email...")
            self.student_search_entry.pack(side="left", padx=(0, 15))
            self.student_search_entry.bind("<KeyRelease>", lambda e: self._refresh_students_list())
            
            self._ghost_btn(search_row, "Clear Search", lambda: (
                self.student_search_entry.delete(0, tk.END),
                self._refresh_students_list()
            ), width=110).pack(side="left")
            
            # Treeview Setup
            _style_tree()
            self.students_tree = ttk.Treeview(scroll, columns=("id", "name", "email", "joined", "interviews"),
                                               show="headings", height=12)
            self.students_tree.heading("id",         text="ID")
            self.students_tree.heading("name",       text="Name")
            self.students_tree.heading("email",      text="Email")
            self.students_tree.heading("joined",     text="Joined Date")
            self.students_tree.heading("interviews", text="Total Interviews")
            
            self.students_tree.column("id",         width=50,  anchor="center")
            self.students_tree.column("name",       width=200)
            self.students_tree.column("email",      width=250)
            self.students_tree.column("joined",     width=180, anchor="center")
            self.students_tree.column("interviews", width=130, anchor="center")
            
            self.students_tree.pack(fill="x", pady=(0, 15))
            
            # Actions Row
            btns = ctk.CTkFrame(scroll, fg_color="transparent")
            btns.pack(anchor="w")
            
            self._btn(btns, "👤 View Profile", self._view_student_profile_action,
                      color=ACCENT, width=140).pack(side="left", padx=(0, 10))
            self._btn(btns, "📜 View History", self._view_student_history_action,
                      color=TEAL, width=140).pack(side="left", padx=(0, 10))
            self._btn(btns, "🗑 Delete Student", self._delete_student_action,
                      color=DANGER, width=150).pack(side="left", padx=(0, 10))
            self._ghost_btn(btns, "🔄 Refresh", self._refresh_students_list,
                            width=100).pack(side="left")
            
            self._refresh_students_list()

        # Wire all sidebar buttons
        self._sidebar_btns["roles"].configure(
            command=lambda: self._sidebar_nav("roles",     build_roles))
        self._sidebar_btns["questions"].configure(
            command=lambda: self._sidebar_nav("questions", build_questions))
        self._sidebar_btns["students"].configure(
            command=lambda: self._sidebar_nav("students",  build_students))
        self._sidebar_btns["analytics"].configure(
            command=lambda: self._sidebar_nav("analytics", build_analytics))
        # ── CHANGE 5 (cont.): Wire admin profile button ─────────────────
        self._sidebar_btns["profile"].configure(
            command=lambda: self._sidebar_nav("profile", build_admin_profile))

        build_roles(content)

    # ------------------------------------------------------------------ #
    #  Admin — students UI action helpers                                #
    # ------------------------------------------------------------------ #

    def _refresh_students_list(self):
        if not hasattr(self, "students_tree") or not self.students_tree.winfo_exists():
            return
        for i in self.students_tree.get_children():
            self.students_tree.delete(i)
        
        q = ""
        if hasattr(self, "student_search_entry") and self.student_search_entry.winfo_exists():
            q = self.student_search_entry.get().strip()
            
        for s in admin.view_students(q):
            self.students_tree.insert("", "end", values=(
                s["id"], s["name"], s["email"], s["created_at"], s["total_interviews"]
            ))

    def _delete_student_action(self):
        sel = self.students_tree.selection()
        if not sel:
            messagebox.showwarning("Select Student", "Please select a student from the list.")
            return
        vals = self.students_tree.item(sel[0])["values"]
        student_id, name, email = vals[0], vals[1], vals[2]
        
        if messagebox.askyesno("Confirm Delete", f"Are you sure you want to delete this student and all their interview results?\n\nStudent: {name} ({email})"):
            ok, msg = admin.delete_student(student_id)
            if ok:
                messagebox.showinfo("Success", "Student and all associated interview results deleted successfully.")
                self._refresh_students_list()
            else:
                messagebox.showerror("Error", f"Failed to delete student: {msg}")

    def _view_student_profile_action(self):
        sel = self.students_tree.selection()
        if not sel:
            messagebox.showwarning("Select Student", "Please select a student from the list.")
            return
        student_id = self.students_tree.item(sel[0])["values"][0]
        stats = admin.get_student_profile_stats(student_id)
        if not stats:
            messagebox.showerror("Error", "Could not load student profile stats.")
            return
            
        win = ctk.CTkToplevel(self)
        win.title(f"Student Profile - {stats['name']}")
        win.geometry("500x380")
        win.resizable(False, False)
        win.grab_set()
        
        win.configure(fg_color=CONTENT_BG[1] if self._get_appearance_mode() == "dark" else CONTENT_BG[0])
        
        card, body = self._card_frame(win, padx=25, pady=25)
        card.pack(fill="both", expand=True, padx=20, pady=20)
        
        ctk.CTkLabel(body, text=f"👤 {stats['name']}", font=F("xl", "bold")).pack(anchor="w", pady=(0, 5))
        ctk.CTkLabel(body, text=stats["email"], font=F("md"), text_color=TEXT_MUTED).pack(anchor="w", pady=(0, 20))
        
        grid = ctk.CTkFrame(body, fg_color="transparent")
        grid.pack(fill="x", pady=(0, 20))
        grid.columnconfigure((0, 1), weight=1)
        
        details = [
            ("Joined Date", stats["created_at"]),
            ("Total Interviews", f"{stats['total_interviews']} completed"),
            ("Average Overall Score", f"{stats['avg_score']}%"),
            ("Best Overall Score", f"{stats['best_score']}%"),
        ]
        
        for idx, (label, val) in enumerate(details):
            r = idx // 2
            c = idx % 2
            box = ctk.CTkFrame(grid, fg_color=("gray95", "gray20"), corner_radius=8)
            box.grid(row=r, column=c, padx=5, pady=5, sticky="ew")
            ctk.CTkLabel(box, text=label, font=F("xs", "bold"), text_color=TEXT_MUTED).pack(padx=10, pady=(8, 2), anchor="w")
            ctk.CTkLabel(box, text=val, font=F("md", "bold")).pack(padx=10, pady=(0, 8), anchor="w")
            
        self._btn(body, "Close", win.destroy, width=120).pack(anchor="e")

    def _view_student_history_action(self):
        sel = self.students_tree.selection()
        if not sel:
            messagebox.showwarning("Select Student", "Please select a student from the list.")
            return
        vals = self.students_tree.item(sel[0])["values"]
        student_id, name = vals[0], vals[1]
        
        results = admin.view_student_scores(student_id)
        
        win = ctk.CTkToplevel(self)
        win.title(f"Interview History - {name}")
        win.geometry("900x650")
        win.minsize(800, 550)
        win.grab_set()
        
        win.configure(fg_color=CONTENT_BG[1] if self._get_appearance_mode() == "dark" else CONTENT_BG[0])
        
        main_frame = ctk.CTkFrame(win, fg_color="transparent")
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        ctk.CTkLabel(main_frame, text=f"Interview History: {name}", font=F("lg", "bold")).pack(anchor="w", pady=(0, 10))
        
        _style_tree()
        history_tree = ttk.Treeview(
            main_frame,
            columns=("role", "diff", "date", "mcq", "tech", "coding", "hr", "overall", "status"),
            show="headings", height=8
        )
        
        for col, h, w in [
            ("role",     "Role",        180),
            ("diff",     "Difficulty",  90),
            ("date",     "Date",        140),
            ("mcq",      "MCQ",         60),
            ("tech",     "Tech",        60),
            ("coding",   "Coding",      60),
            ("hr",       "HR",          60),
            ("overall",  "Overall",     70),
            ("status",   "Status",      70),
        ]:
            history_tree.heading(col, text=h)
            history_tree.column(col, width=w, anchor="center" if col not in ("role") else "w")
            
        history_tree.pack(fill="x", pady=(0, 15))
        
        for r in results:
            status_val = r.get("status")
            if not status_val:
                status_val = "PASS" if r["overall_score"] >= 60 else "FAIL"
                
            has_coding_session = False
            try:
                if r.get("session_data"):
                    snap = json.loads(r["session_data"])
                    if snap.get("coding") and len(snap["coding"]) > 0:
                        has_coding_session = True
            except Exception:
                pass
            coding_val = f"{r['coding_score']}%" if has_coding_session else "—"

            history_tree.insert("", "end", iid=str(r["id"]), values=(
                r["role_name"],
                r["difficulty"] or "—",
                r["created_at"],
                f"{r['mcq_score']}%",
                f"{r['technical_score']}%",
                coding_val,
                f"{r['hr_score']}%",
                f"{r['overall_score']}%",
                status_val
            ))
            
        detail_card, detail_body = self._card_frame(main_frame, padx=15, pady=15)
        detail_card.pack(fill="both", expand=True)
        
        placeholder = ctk.CTkLabel(detail_body, text="Select an interview record above to view complete details.",
                                   font=F("md", "bold"), text_color=TEXT_MUTED)
        placeholder.pack(expand=True)
                     
        def on_select_history(event):
            selected = history_tree.selection()
            if not selected:
                return
            result_id = int(selected[0])
            r = next((x for x in results if x["id"] == result_id), None)
            if not r:
                return
                
            for w in detail_body.winfo_children():
                w.destroy()
                
            top_row = ctk.CTkFrame(detail_body, fg_color="transparent")
            top_row.pack(fill="x", pady=(0, 10))
            
            ctk.CTkLabel(top_row, text=f"Role: {r['role_name']} ({r['difficulty'] or '—'})", 
                         font=F("md", "bold")).pack(side="left")
            ctk.CTkLabel(top_row, text=f"Date: {r['created_at']}", 
                         font=F("sm"), text_color=TEXT_MUTED).pack(side="right")
                         
            scores_frame = ctk.CTkFrame(detail_body, fg_color="transparent")
            scores_frame.pack(fill="x", pady=(0, 10))
            scores_frame.columnconfigure((0, 1, 2, 3), weight=1)
            
            status_val = r.get("status")
            if not status_val:
                status_val = "PASS" if r["overall_score"] >= 60 else "FAIL"
                
            scores = [
                ("MCQ Score", f"{r['mcq_score']}%", ACCENT),
                ("Technical Score", f"{r['technical_score']}%", TEAL),
                ("HR Score", f"{r['hr_score']}%", PINK),
                ("Overall Score", f"{r['overall_score']}% ({status_val})", AMBER if status_val == "PASS" else DANGER),
            ]
            
            for i, (label, val, color) in enumerate(scores):
                box = ctk.CTkFrame(scores_frame, fg_color=color, corner_radius=6)
                box.grid(row=0, column=i, padx=4, sticky="ew")
                ctk.CTkLabel(box, text=label, font=F("xs", "bold"), text_color="white").pack(pady=(4, 0))
                ctk.CTkLabel(box, text=val, font=F("sm", "bold"), text_color="white").pack(pady=(0, 4))
                
            areas_frame = ctk.CTkFrame(detail_body, fg_color="transparent")
            areas_frame.pack(fill="x", pady=(0, 10))
            areas_frame.columnconfigure((0, 1), weight=1)
            
            import feedback_engine
            strong_areas = r.get("strong_areas") or ", ".join(feedback_engine.find_strengths(r["mcq_score"], r["technical_score"], r["hr_score"]))
            weak_areas = r.get("weak_areas") or ", ".join(feedback_engine.find_weaknesses(r["mcq_score"], r["technical_score"], r["hr_score"]))
            
            s_box = ctk.CTkFrame(areas_frame, fg_color=("gray95", "gray20"), corner_radius=6)
            s_box.grid(row=0, column=0, padx=(0, 5), sticky="ew")
            ctk.CTkLabel(s_box, text="Strong Areas", font=F("xs", "bold"), text_color=SUCCESS).pack(padx=8, pady=(4, 0), anchor="w")
            ctk.CTkLabel(s_box, text=strong_areas, font=F("xs"), wraplength=350, justify="left").pack(padx=8, pady=(0, 6), anchor="w")
            
            w_box = ctk.CTkFrame(areas_frame, fg_color=("gray95", "gray20"), corner_radius=6)
            w_box.grid(row=0, column=1, padx=(5, 0), sticky="ew")
            ctk.CTkLabel(w_box, text="Weak Areas", font=F("xs", "bold"), text_color=DANGER).pack(padx=8, pady=(4, 0), anchor="w")
            ctk.CTkLabel(w_box, text=weak_areas, font=F("xs"), wraplength=350, justify="left").pack(padx=8, pady=(0, 6), anchor="w")
            
            ctk.CTkLabel(detail_body, text="Detailed Feedback Report:", font=F("sm", "bold"), text_color=TEXT_MUTED).pack(anchor="w", pady=(5, 2))
            fb_text = ctk.CTkTextbox(detail_body, height=120, wrap="word", font=F("xs"))
            fb_text.pack(fill="both", expand=True)
            fb_text.insert("1.0", r["feedback"])
            fb_text.configure(state="disabled")
            
        history_tree.bind("<<TreeviewSelect>>", on_select_history)

    # ------------------------------------------------------------------ #
    #  Admin — roles CRUD                                                #
    # ------------------------------------------------------------------ #

    def _refresh_roles(self):
        for i in self.roles_tree.get_children():
            self.roles_tree.delete(i)
        for r in admin.fetch_roles():
            coding_str = "Enabled" if r.get("has_coding_round") else "Disabled"
            self.roles_tree.insert("", "end", values=(r["id"], r["role_name"], coding_str))

    def _add_role(self):
        has_coding = getattr(self, "new_role_coding_var", None)
        coding_val = has_coding.get() if has_coding else 0
        ok, msg = admin.add_role(self.new_role_entry.get(), coding_val)
        if ok:
            self.new_role_entry.delete(0, tk.END)
            if has_coding:
                has_coding.set(0)
            self._refresh_roles()
        messagebox.showinfo("Role", msg)

    def _update_role(self):
        sel = self.roles_tree.selection()
        if not sel:
            messagebox.showwarning("Select", "Select a role first.")
            return
        vals = self.roles_tree.item(sel[0])["values"]
        rid = vals[0]
        name = vals[1]
        coding_str = vals[2] if len(vals) > 2 else "Disabled"
        has_coding_init = 1 if coding_str == "Enabled" else 0

        # Custom Toplevel dialog
        dialog = ctk.CTkToplevel(self)
        dialog.title("Edit Role")
        dialog.geometry("400x250")
        dialog.transient(self)
        dialog.grab_set()
        # Center the dialog relative to parent self
        dialog.update_idletasks()
        x = self.winfo_x() + (self.winfo_width() - dialog.winfo_width()) // 2
        y = self.winfo_y() + (self.winfo_height() - dialog.winfo_height()) // 2
        dialog.geometry(f"+{x}+{y}")

        # Widgets
        content = ctk.CTkFrame(dialog, fg_color="transparent")
        content.pack(fill="both", expand=True, padx=20, pady=20)

        ctk.CTkLabel(content, text="Role Name:", font=F("md", "bold")).pack(anchor="w", pady=(0, 5))
        name_entry = self._entry(content, width=360)
        name_entry.insert(0, name)
        name_entry.pack(anchor="w", pady=(0, 15))

        coding_var = tk.IntVar(value=has_coding_init)
        cb = ctk.CTkCheckBox(content, text="Include Coding Round", variable=coding_var, font=F("md"))
        cb.pack(anchor="w", pady=(0, 20))

        btn_row = ctk.CTkFrame(content, fg_color="transparent")
        btn_row.pack(anchor="w")

        def save():
            new_name = name_entry.get().strip()
            if not new_name:
                messagebox.showerror("Error", "Role name cannot be empty.")
                return
            ok, msg = admin.update_role(rid, new_name, coding_var.get())
            if ok:
                self._refresh_roles()
                dialog.destroy()
                messagebox.showinfo("Success", msg)
            else:
                messagebox.showerror("Error", msg)

        self._btn(btn_row, "Save", save, width=100).pack(side="left", padx=(0, 10))
        self._ghost_btn(btn_row, "Cancel", dialog.destroy, width=100).pack(side="left")

    def _delete_role(self):
        sel = self.roles_tree.selection()
        if not sel:
            return
        rid = self.roles_tree.item(sel[0])["values"][0]
        if messagebox.askyesno("Confirm", "Delete this role and all its questions?"):
            ok, msg = admin.delete_role(rid)
            self._refresh_roles()
            messagebox.showinfo("Role", msg)

    # ------------------------------------------------------------------ #
    #  Admin — questions panel                                           #
    # ------------------------------------------------------------------ #

    def _admin_questions_panel(self, parent):
        tab_row = ctk.CTkFrame(parent, fg_color="transparent")
        tab_row.pack(anchor="w", pady=(0, 16))
        q_content = ctk.CTkFrame(parent, fg_color="transparent")
        q_content.pack(fill="both", expand=True)

        def clear_q():
            for w in q_content.winfo_children(): w.destroy()

        for label, cmd_fn, color in [
            ("MCQ",       lambda: (clear_q(), self._build_mcq_form(q_content)),  ACCENT),
            ("Technical", lambda: (clear_q(), self._build_tech_form(q_content)), TEAL),
            ("Coding",    lambda: (clear_q(), self._build_coding_form(q_content)), AMBER),
            ("HR",        lambda: (clear_q(), self._build_hr_form(q_content)),   PINK),
            ("Bulk Upload", lambda: (clear_q(), self._build_bulk_upload(q_content)), AMBER),
            ("View All",  lambda: (clear_q(), self._build_q_list(q_content)),
             ("gray40", "gray50")),
        ]:
            self._btn(tab_row, label, cmd_fn,
                      color=color, width=110, height=34).pack(
                side="left", padx=(0, 8))

        clear_q()
        self._build_mcq_form(q_content)

    def _role_combo_widget(self, parent):
        roles = admin.fetch_roles()
        var   = tk.StringVar()
        vals  = [f"{r['id']}: {r['role_name']}" for r in roles]
        if not vals:
            vals = ["No roles available"]
        cb = ctk.CTkComboBox(parent, variable=var, values=vals,
                             width=340, state="readonly", font=F("md"))
        var.set(vals[0])
        return var, cb

    def _get_role_id_from_var(self, var):
        value = var.get().strip()
        if not value or ":" not in value:
            return None
        try:
            return int(value.split(":")[0])
        except ValueError:
            return None

    def _grid_label(self, parent, text, row):
        ctk.CTkLabel(parent, text=text, font=F("sm", "bold"),
                     text_color=TEXT_MUTED, anchor="e").grid(
            row=row, column=0, sticky="e", padx=(0, 10), pady=5)

    def _grid_entry(self, parent, row):
        e = self._entry(parent, width=340)
        e.grid(row=row, column=1, pady=5, sticky="ew")
        return e

    def _build_mcq_form(self, parent):
        parent.columnconfigure(1, weight=1)
        self.mcq_role_var, cb = self._role_combo_widget(parent)
        self._grid_label(parent, "Role:", 0)
        cb.grid(row=0, column=1, pady=5, sticky="ew")
        self._grid_label(parent, "Difficulty:", 1)
        self.mcq_diff = ctk.CTkComboBox(parent,
                                         values=["Easy", "Medium", "Hard"],
                                         width=340, state="readonly", font=F("md"))
        self.mcq_diff.set("Easy")
        self.mcq_diff.grid(row=1, column=1, pady=5)
        self.mcq_q   = self._grid_labeled_entry(parent, "Question:", 2)
        self.mcq_a   = self._grid_labeled_entry(parent, "Option A:", 3)
        self.mcq_b   = self._grid_labeled_entry(parent, "Option B:", 4)
        self.mcq_c   = self._grid_labeled_entry(parent, "Option C:", 5)
        self.mcq_d   = self._grid_labeled_entry(parent, "Option D:", 6)
        self.mcq_ans = self._grid_labeled_entry(parent, "Correct (A/B/C/D):", 7)
        self._btn(parent, "Add MCQ", self._add_mcq,
                  width=130).grid(row=8, column=1, pady=14, sticky="w")

    def _build_coding_form(self, parent):
        parent.columnconfigure(1, weight=1)
        self.coding_role_var, cb = self._role_combo_widget(parent)
        self._grid_label(parent, "Role:", 0)
        cb.grid(row=0, column=1, pady=5, sticky="ew")
        
        self._grid_label(parent, "Difficulty:", 1)
        self.coding_diff = ctk.CTkComboBox(parent,
                                            values=["Easy", "Medium", "Hard"],
                                            width=340, state="readonly", font=F("md"))
        self.coding_diff.set("Easy")
        self.coding_diff.grid(row=1, column=1, pady=5, sticky="w")
        
        self.coding_q = self._grid_labeled_entry(parent, "Question:", 2)
        
        self._grid_label(parent, "Code Snippet:", 3)
        self.coding_code = ctk.CTkTextbox(parent, height=120, font=ctk.CTkFont(family="Courier", size=12), corner_radius=8)
        self.coding_code.grid(row=3, column=1, pady=5, sticky="ew")
        
        self.coding_a   = self._grid_labeled_entry(parent, "Option A:", 4)
        self.coding_b   = self._grid_labeled_entry(parent, "Option B:", 5)
        self.coding_c   = self._grid_labeled_entry(parent, "Option C:", 6)
        self.coding_d   = self._grid_labeled_entry(parent, "Option D:", 7)
        self.coding_ans = self._grid_labeled_entry(parent, "Correct (A/B/C/D):", 8)
        
        self._btn(parent, "Add Coding Question", self._add_coding_question_action,
                  color=AMBER, width=180).grid(row=9, column=1, pady=14, sticky="w")

    def _build_tech_form(self, parent):
        parent.columnconfigure(1, weight=1)
        self.tech_role_var, cb = self._role_combo_widget(parent)
        self._grid_label(parent, "Role:", 0)
        cb.grid(row=0, column=1, pady=5, sticky="ew")
        self._grid_label(parent, "Difficulty:", 1)
        self.tech_diff = ctk.CTkComboBox(parent,
                                          values=["Easy", "Medium", "Hard"],
                                          width=340, state="readonly", font=F("md"))
        self.tech_diff.set("Easy")
        self.tech_diff.grid(row=1, column=1, pady=5)
        self.tech_q  = self._grid_labeled_entry(parent, "Question:", 2)
        self.tech_kw = self._grid_labeled_entry(parent,
                                                 "Keywords (comma-sep, min 3):", 3)
        self._btn(parent, "Add Technical", self._add_technical,
                  color=TEAL, width=150).grid(row=4, column=1, pady=14, sticky="w")

    def _build_hr_form(self, parent):
        parent.columnconfigure(1, weight=1)
        self.hr_role_var, cb = self._role_combo_widget(parent)
        self._grid_label(parent, "Role:", 0)
        cb.grid(row=0, column=1, pady=5, sticky="ew")
        self._grid_label(parent, "Difficulty:", 1)
        self.hr_diff = ctk.CTkComboBox(
            parent, values=["Easy", "Medium", "Hard"],
            width=340, state="readonly", font=F("md"))
        self.hr_diff.set("Easy")
        self.hr_diff.grid(row=1, column=1, pady=5, sticky="ew")
        self.hr_q  = self._grid_labeled_entry(parent, "Question:", 2)
        self.hr_kw = self._grid_labeled_entry(parent, "Expected Keywords:", 3)
        self._btn(parent, "Add HR", self._add_hr,
                  color=PINK, width=120).grid(row=4, column=1, pady=14, sticky="w")

    def _download_template(self):
        import csv
        import json
        from docx import Document

        win = ctk.CTkToplevel(self)
        win.title("Download Template")
        win.geometry("420x220")
        win.resizable(False, False)
        win.grab_set()

        ctk.CTkLabel(win, text="Choose Template Format", font=F("lg")).pack(pady=(20, 10))
        ctk.CTkLabel(win, text="Click any format to download instantly.",
                     font=F("sm"), text_color=TEXT_MUTED).pack(pady=(0, 20))

        def download_csv():
            filepath = filedialog.asksaveasfilename(
                defaultextension=".csv", filetypes=[("CSV File", "*.csv")],
                initialfile="questions_template.csv")
            if not filepath:
                return
            rows = [
                ["type", "question", "option_a", "option_b", "option_c",
                 "option_d", "correct_answer", "keywords", "ideal_answer"],
                ["mcq", "What is polymorphism?", "Inheritance", "Method Overloading",
                 "Abstraction", "Encapsulation", "B", "", ""],
                ["technical", "Explain deadlock in OS.", "", "", "", "", "",
                 "deadlock, locking, process", "Should explain circular wait."],
                ["hr", "Tell me about a challenge you faced.", "", "", "", "", "",
                 "teamwork, communication", "Should explain challenge and learning."]
            ]
            with open(filepath, "w", newline="", encoding="utf-8") as f:
                csv.writer(f).writerows(rows)
            win.destroy()
            messagebox.showinfo("Success", "CSV template downloaded.")

        def download_json():
            filepath = filedialog.asksaveasfilename(
                defaultextension=".json", filetypes=[("JSON File", "*.json")],
                initialfile="questions_template.json")
            if not filepath:
                return
            import json
            data = [
                {"type": "mcq", "question": "What is polymorphism?",
                 "option_a": "Inheritance", "option_b": "Method Overloading",
                 "option_c": "Abstraction", "option_d": "Encapsulation",
                 "correct_answer": "B", "keywords": "", "ideal_answer": ""},
                {"type": "technical", "question": "Explain deadlock in OS.",
                 "option_a": "", "option_b": "", "option_c": "", "option_d": "",
                 "correct_answer": "", "keywords": "deadlock, locking, process",
                 "ideal_answer": "Should explain circular wait."},
                {"type": "hr", "question": "Tell me about a challenge you faced.",
                 "option_a": "", "option_b": "", "option_c": "", "option_d": "",
                 "correct_answer": "", "keywords": "teamwork, communication",
                 "ideal_answer": "Should explain challenge and learning."}
            ]
            with open(filepath, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=4)
            win.destroy()
            messagebox.showinfo("Success", "JSON template downloaded.")

        def download_docx():
            filepath = filedialog.asksaveasfilename(
                defaultextension=".docx", filetypes=[("Word File", "*.docx")],
                initialfile="questions_template.docx")
            if not filepath:
                return
            doc = Document()
            doc.add_heading("Question Import Template", level=1)
            doc.add_paragraph("Separate each question with ----")
            doc.add_paragraph(
                "TYPE: mcq\nQUESTION: What is polymorphism?\n"
                "A: Inheritance\nB: Method Overloading\nC: Abstraction\n"
                "D: Encapsulation\nANSWER: B\nKEYWORDS:\nIDEAL_ANSWER:\n----\n"
                "TYPE: technical\nQUESTION: Explain deadlock in OS.\n"
                "KEYWORDS: deadlock, locking, process\n"
                "IDEAL_ANSWER: Should explain circular wait.\n----\n"
                "TYPE: hr\nQUESTION: Tell me about a challenge you faced.\n"
                "KEYWORDS: teamwork, communication\n"
                "IDEAL_ANSWER: Should explain challenge and learning.\n----"
            )
            doc.save(filepath)
            win.destroy()
            messagebox.showinfo("Success", "DOCX template downloaded.")

        btn_row = ctk.CTkFrame(win, fg_color="transparent")
        btn_row.pack(pady=10)
        self._btn(btn_row, "CSV Template",  download_csv,  width=120).grid(
            row=0, column=0, padx=8)
        self._btn(btn_row, "JSON Template", download_json, width=120).grid(
            row=0, column=1, padx=8)
        self._btn(btn_row, "DOCX Template", download_docx, width=120).grid(
            row=0, column=2, padx=8)

    def _build_bulk_upload(self, parent):
        parent.columnconfigure(0, weight=0)
        parent.columnconfigure(1, weight=0)
        parent.columnconfigure(2, weight=1)

        self.bulk_role_var, cb = self._role_combo_widget(parent)
        self._grid_label(parent, "Job role:", 0)
        cb.grid(row=0, column=1, columnspan=2, padx=(12, 0), pady=(8, 12), sticky="ew")

        self.bulk_diff_var = tk.StringVar(value="Easy")
        self._grid_label(parent, "Difficulty:", 1)
        diff_cb = ctk.CTkComboBox(parent, variable=self.bulk_diff_var,
                                   values=["Easy", "Medium", "Hard"],
                                   width=180, state="readonly", font=F("md"))
        diff_cb.grid(row=1, column=1, padx=(12, 0), pady=(6, 12), sticky="w")

        self._grid_label(parent, "Question file:", 2)
        self._btn(parent, "Choose CSV / JSON / DOCX", self._pick_bulk_file,
                  width=190).grid(row=2, column=1, padx=(12, 10), pady=8, sticky="w")
        self.bulk_file_label = ctk.CTkLabel(parent, text="No file selected",
                                             font=F("sm"), text_color=TEXT_MUTED,
                                             anchor="w")
        self.bulk_file_label.grid(row=2, column=2, padx=(8, 0), pady=8, sticky="w")

        self.bulk_replace_var = tk.BooleanVar(value=False)
        ctk.CTkCheckBox(
            parent,
            text="Replace all existing questions for this role + difficulty",
            variable=self.bulk_replace_var, font=F("sm"),
        ).grid(row=3, column=1, columnspan=2, padx=(12, 0), pady=(10, 16), sticky="w")

        hint = (
            "Upload one file per role + difficulty (.csv, .json, .docx).\n"
            "Columns: type, question, option_a, option_b,\n"
            "option_c, option_d, correct_answer,\n"
            "keywords, ideal_answer\n"
            "Difficulty is selected above.\n"
            "Template: assets/questions_template.csv"
        )
        ctk.CTkLabel(parent, text=hint, font=F("xs"), text_color=TEXT_MUTED,
                     justify="left", wraplength=520).grid(
            row=4, column=1, columnspan=2, padx=(12, 0), pady=(4, 18), sticky="w")

        btn_row = ctk.CTkFrame(parent, fg_color="transparent")
        btn_row.grid(row=5, column=1, columnspan=2, padx=(12, 0), pady=8, sticky="w")
        self._btn(btn_row, "Upload questions", self._run_bulk_upload,
                  color=AMBER, width=170).grid(row=0, column=0, padx=(0, 12))
        self._ghost_btn(btn_row, "Open Template", self._open_bulk_template,
                        width=150).grid(row=0, column=1, padx=(0, 12))
        self._ghost_btn(btn_row, "Download Template", self._download_template,
                        width=170).grid(row=0, column=2)

        self._bulk_filepath = None

    def _pick_bulk_file(self):
        path = filedialog.askopenfilename(
            title="Select question bank",
            filetypes=[
                ("Question Files", "*.csv *.json *.docx"),
                ("CSV files", "*.csv"),
                ("JSON files", "*.json"),
                ("Word files", "*.docx"),
            ],
        )
        if path:
            self._bulk_filepath = path
            self.bulk_file_label.configure(text=os.path.basename(path))

    def _open_bulk_template(self):
        import subprocess
        template = "assets/questions_template.csv"
        if not os.path.exists(template):
            messagebox.showerror("Template Missing",
                                 "questions_template.csv not found in assets folder.")
            return
        try:
            if sys.platform.startswith("win"):
                os.startfile(template)
            elif sys.platform == "darwin":
                subprocess.Popen(["open", template])
            else:
                subprocess.Popen(["xdg-open", template])
        except Exception as exc:
            messagebox.showerror("Open Failed", f"Could not open template:\n{exc}")

    def _run_bulk_upload(self):
        role_id = self._get_role_id_from_var(self.bulk_role_var)
        if not role_id:
            messagebox.showwarning("Role", "Select a valid job role.")
            return

        difficulty = self.bulk_diff_var.get().strip()
        if difficulty not in ["Easy", "Medium", "Hard"]:
            messagebox.showwarning("Difficulty", "Select a valid difficulty level.")
            return

        filepath = getattr(self, "_bulk_filepath", None)
        if not filepath:
            messagebox.showwarning("File", "Choose a CSV, JSON, or DOCX file first.")
            return

        ext = filepath.lower().split(".")[-1]
        if ext not in {"csv", "json", "docx"}:
            messagebox.showerror("Invalid File", "Supported formats: CSV, JSON, DOCX")
            return

        replace = self.bulk_replace_var.get()
        if replace and not messagebox.askyesno(
            "Replace questions",
            f"This will delete ALL existing '{difficulty}' questions for the selected role "
            f"and import from the file.\n\nOther difficulties are NOT affected.\n\nContinue?"
        ):
            return

        try:
            ok, msg = admin.import_questions_from_file(
                role_id, filepath,
                difficulty=difficulty,
                replace_existing=replace,
            )
            if ok:
                self._refresh_questions()
                self._bulk_filepath = None
                if hasattr(self, "bulk_file_label"):
                    self.bulk_file_label.configure(text="No file selected")
                # Reset the checkbox so the next upload doesn't accidentally replace
                if hasattr(self, "bulk_replace_var"):
                    self.bulk_replace_var.set(False)
            messagebox.showinfo("Bulk Upload", msg)
        except Exception as exc:
            messagebox.showerror("Bulk Upload Failed", f"Error: {exc}")

    def _build_q_list(self, parent):
        # Search panel
        sf = ctk.CTkFrame(parent, fg_color="transparent")
        sf.pack(fill="x", pady=(0, 10))
        
        # Search Entry
        ctk.CTkLabel(sf, text="Search:", font=F("sm", "bold"), text_color=TEXT_MUTED).grid(row=0, column=0, padx=(0, 5), sticky="w")
        self.q_search_entry = self._entry(sf, width=200, placeholder="Question text or keywords...")
        self.q_search_entry.grid(row=0, column=1, padx=(0, 15), sticky="w")
        self.q_search_entry.bind("<KeyRelease>", lambda e: self._refresh_questions())
        
        # Role Filter
        ctk.CTkLabel(sf, text="Role:", font=F("sm", "bold"), text_color=TEXT_MUTED).grid(row=0, column=2, padx=(0, 5), sticky="w")
        roles = admin.fetch_roles()
        role_vals = ["All Roles"] + [f"{r['id']}: {r['role_name']}" for r in roles]
        self.q_filter_role = ctk.CTkComboBox(sf, values=role_vals, width=150, state="readonly", font=F("sm"))
        self.q_filter_role.set("All Roles")
        self.q_filter_role.grid(row=0, column=3, padx=(0, 15), sticky="w")
        self.q_filter_role.configure(command=lambda v: self._refresh_questions())
        
        # Difficulty Filter
        ctk.CTkLabel(sf, text="Difficulty:", font=F("sm", "bold"), text_color=TEXT_MUTED).grid(row=0, column=4, padx=(0, 5), sticky="w")
        self.q_filter_diff = ctk.CTkComboBox(sf, values=["All Difficulties", "Easy", "Medium", "Hard"], width=130, state="readonly", font=F("sm"))
        self.q_filter_diff.set("All Difficulties")
        self.q_filter_diff.grid(row=0, column=5, padx=(0, 15), sticky="w")
        self.q_filter_diff.configure(command=lambda v: self._refresh_questions())
        
        # Type Filter
        ctk.CTkLabel(sf, text="Type:", font=F("sm", "bold"), text_color=TEXT_MUTED).grid(row=0, column=6, padx=(0, 5), sticky="w")
        self.q_filter_type = ctk.CTkComboBox(sf, values=["All Types", "MCQ", "Technical", "Coding", "HR"], width=110, state="readonly", font=F("sm"))
        self.q_filter_type.set("All Types")
        self.q_filter_type.grid(row=0, column=7, padx=(0, 15), sticky="w")
        self.q_filter_type.configure(command=lambda v: self._refresh_questions())

        _style_tree()
        tree_frame = tk.Frame(parent, bg="#1e1e2e")
        tree_frame.pack(fill="both", expand=True, pady=(0, 10))

        vsb = ttk.Scrollbar(tree_frame, orient="vertical")
        vsb.pack(side="right", fill="y")

        self.q_tree = ttk.Treeview(
            tree_frame,
            columns=("id", "role", "type", "diff", "question"),
            show="headings", height=14,
            yscrollcommand=vsb.set,
        )
        vsb.configure(command=self.q_tree.yview)
        for col, h, w in [
            ("id",       "ID",       40),
            ("role",     "Role",    140),
            ("type",     "Type",     90),
            ("diff",     "Diff",     60),
            ("question", "Question", 300),
        ]:
            self.q_tree.heading(col, text=h)
            self.q_tree.column(col, width=w)
        self.q_tree.pack(side="left", fill="both", expand=True)

        bf = ctk.CTkFrame(parent, fg_color="transparent")
        bf.pack(anchor="w")
        self._ghost_btn(bf, "Refresh", self._refresh_questions,
                        width=100).pack(side="left", padx=(0, 8))
        self._btn(bf, "Delete", self._delete_question,
                  color=DANGER, width=100).pack(side="left")
        self._refresh_questions()

    def _grid_labeled_entry(self, parent, label, row):
        self._grid_label(parent, label, row)
        e = self._entry(parent, width=340)
        e.grid(row=row, column=1, pady=5, sticky="ew")
        return e

    def _add_mcq(self):
        ok, msg = admin.add_mcq_question(
            self._get_role_id_from_var(self.mcq_role_var),
            self.mcq_diff.get(),
            self.mcq_q.get(), self.mcq_a.get(), self.mcq_b.get(),
            self.mcq_c.get(), self.mcq_d.get(), self.mcq_ans.get(),
        )
        messagebox.showinfo("MCQ", msg)
        if ok: self._refresh_questions()

    def _add_coding_question_action(self):
        ok, msg = admin.add_coding_question(
            self._get_role_id_from_var(self.coding_role_var),
            self.coding_diff.get(),
            self.coding_q.get(),
            self.coding_code.get("1.0", "end-1c"),
            self.coding_a.get(),
            self.coding_b.get(),
            self.coding_c.get(),
            self.coding_d.get(),
            self.coding_ans.get(),
        )
        messagebox.showinfo("Coding Question", msg)
        if ok:
            self.coding_q.delete(0, tk.END)
            self.coding_code.delete("1.0", tk.END)
            self.coding_a.delete(0, tk.END)
            self.coding_b.delete(0, tk.END)
            self.coding_c.delete(0, tk.END)
            self.coding_d.delete(0, tk.END)
            self.coding_ans.delete(0, tk.END)
            self._refresh_questions()

    def _add_technical(self):
        ok, msg = admin.add_technical_question(
            self._get_role_id_from_var(self.tech_role_var),
            self.tech_diff.get(), self.tech_q.get(), self.tech_kw.get(),
        )
        messagebox.showinfo("Technical", msg)
        if ok: self._refresh_questions()

    def _add_hr(self):
        ok, msg = admin.add_hr_question(
            self._get_role_id_from_var(self.hr_role_var),
            self.hr_diff.get(),
            self.hr_q.get(),
            self.hr_kw.get(),
        )
        messagebox.showinfo("HR", msg)
        if ok: self._refresh_questions()

    def _refresh_questions(self):
        if not hasattr(self, "q_tree") or not self.q_tree.winfo_exists():
            return
        for i in self.q_tree.get_children():
            self.q_tree.delete(i)
            
        search_query = ""
        if hasattr(self, "q_search_entry") and self.q_search_entry.winfo_exists():
            search_query = self.q_search_entry.get().strip()
            
        role_id = None
        if hasattr(self, "q_filter_role") and self.q_filter_role.winfo_exists():
            sel = self.q_filter_role.get()
            if sel != "All Roles" and ":" in sel:
                try:
                    role_id = int(sel.split(":")[0])
                except ValueError:
                    pass
                    
        difficulty = None
        if hasattr(self, "q_filter_diff") and self.q_filter_diff.winfo_exists():
            sel = self.q_filter_diff.get()
            if sel != "All Difficulties":
                difficulty = sel
                
        q_type = None
        if hasattr(self, "q_filter_type") and self.q_filter_type.winfo_exists():
            sel = self.q_filter_type.get()
            if sel != "All Types":
                q_type = sel.lower()
                
        for q in admin.search_questions(search_query, role_id, difficulty, q_type):
            self.q_tree.insert("", "end", values=(
                q["id"], q["role_name"], q["question_type"],
                q["difficulty"] or "—", (q["question"] or "")[:55],
            ))

    def _delete_question(self):
        sel = self.q_tree.selection()
        if not sel: return
        qid = self.q_tree.item(sel[0])["values"][0]
        if messagebox.askyesno("Confirm", "Delete this question?"):
            ok, msg = admin.delete_question(qid)
            self._refresh_questions()
            messagebox.showinfo("Question", msg)

    def _export_all_results(self):
        filepath = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV File", "*.csv")],
            initialfile="all_interview_results.csv"
        )
        if not filepath:
            return
        try:
            admin.export_all_results_csv(filepath)
            messagebox.showinfo("Success", "Results exported successfully.")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to export results: {e}")

    def _admin_analytics_panel(self, parent):
        stats     = admin.get_dashboard_stats()
        avg       = admin.get_average_scores()
        role_perf = admin.get_role_wise_performance()
        recent    = admin.get_recent_interviews()
        
        diff_sum  = admin.get_question_difficulty_summary()
        top_stud  = admin.get_top_performers(5)
        failed_roles = admin.get_most_failed_roles()
        activities = admin.get_recent_activities(10)
        q_per_role = admin.get_question_count_per_role()

        # 1. Row 1: Statistics Cards (Total Students, Total Roles, Total Questions, Total Interviews Taken)
        cards = ctk.CTkFrame(parent, fg_color="transparent")
        cards.pack(fill="x", pady=(0, 20))
        for i in range(4):
            cards.columnconfigure(i, weight=1)
            
        stat_items = [
            ("👥 Total Students", stats["students"], ACCENT),
            ("📋 Total Roles",    stats["roles"],    TEAL),
            ("❓ Total Questions",stats["questions"], PINK),
            ("📊 Total Interviews",stats["interviews"], AMBER),
        ]
        
        for col, (label, val, color) in enumerate(stat_items):
            box = ctk.CTkFrame(cards, corner_radius=10, border_color=CARD_BORDER[1], border_width=1)
            box.grid(row=0, column=col, padx=6, sticky="ew")
            
            lbl = ctk.CTkLabel(box, text=label, font=F("sm", "bold"), text_color=color)
            lbl.pack(pady=(12, 2), padx=10, anchor="w")
            
            val_lbl = ctk.CTkLabel(box, text=str(val), font=F("xxl", "bold"))
            val_lbl.pack(pady=(0, 12), padx=10, anchor="w")

        # 2. Row 2: Question Difficulty Summary Cards
        diff_frame = ctk.CTkFrame(parent, fg_color="transparent")
        diff_frame.pack(fill="x", pady=(0, 20))
        diff_frame.columnconfigure((0, 1, 2), weight=1)
        
        diff_items = [
            ("🟢 Easy Questions",   diff_sum["easy"],   SUCCESS),
            ("🟡 Medium Questions", diff_sum["medium"], AMBER),
            ("🔴 Hard Questions",   diff_sum["hard"],   DANGER),
        ]
        
        for col, (label, val, color) in enumerate(diff_items):
            box = ctk.CTkFrame(diff_frame, corner_radius=8, fg_color=("gray95", "#222232"))
            box.grid(row=0, column=col, padx=6, sticky="ew")
            
            lbl = ctk.CTkLabel(box, text=label, font=F("xs", "bold"), text_color=TEXT_MUTED)
            lbl.pack(side="left", padx=15, pady=10)
            
            val_lbl = ctk.CTkLabel(box, text=str(val), font=F("md", "bold"), text_color=color)
            val_lbl.pack(side="right", padx=15, pady=10)

        # 2b. Row 2.5: Coding Output Round Analytics Row
        coding_stats = admin.get_coding_analytics_stats()
        
        coding_stats_frame = ctk.CTkFrame(parent, fg_color="transparent")
        coding_stats_frame.pack(fill="x", pady=(0, 20))
        for i in range(4):
            coding_stats_frame.columnconfigure(i, weight=1)
            
        coding_stat_items = [
            ("💻 Coding Attempts", coding_stats["attempts"], AMBER),
            ("📈 Avg Coding Score",       f"{coding_stats['avg_score']}%", TEAL),
            ("🎯 Coding Success Rate",     f"{coding_stats['success_rate']}%", SUCCESS),
            ("🏆 Best Coding Performer",  coding_stats["best_performer"], ACCENT),
        ]
        
        for col, (label, val, color) in enumerate(coding_stat_items):
            box = ctk.CTkFrame(coding_stats_frame, corner_radius=10, border_color=CARD_BORDER[1], border_width=1)
            box.grid(row=0, column=col, padx=6, sticky="ew")
            
            lbl = ctk.CTkLabel(box, text=label, font=F("sm", "bold"), text_color=color)
            lbl.pack(pady=(12, 2), padx=10, anchor="w")
            
            val_lbl = ctk.CTkLabel(box, text=str(val), font=F("sm" if col == 3 else "xxl", "bold"))
            val_lbl.pack(pady=(0, 12), padx=10, anchor="w")

        # 3. Row 3: Two-Column Split Layout
        split_frame = ctk.CTkFrame(parent, fg_color="transparent")
        split_frame.pack(fill="both", expand=True)
        split_frame.columnconfigure(0, weight=1, minsize=400)
        split_frame.columnconfigure(1, weight=1, minsize=400)
        
        left_col = ctk.CTkFrame(split_frame, fg_color="transparent")
        left_col.grid(row=0, column=0, padx=(0, 10), sticky="nwe")
        
        right_col = ctk.CTkFrame(split_frame, fg_color="transparent")
        right_col.grid(row=0, column=1, padx=(10, 0), sticky="nwe")

        # ------------------- LEFT COLUMN PANELS -------------------
        
        # Panel A: Average Scores
        score_card, score_body = self._card_frame(left_col, padx=15, pady=15)
        score_card.pack(fill="x", pady=(0, 15))
        ctk.CTkLabel(score_body, text="Average Round Scores", font=F("md", "bold")).pack(anchor="w", pady=(0, 10))
        
        score_grid = ctk.CTkFrame(score_body, fg_color="transparent")
        score_grid.pack(fill="x")
        score_grid.columnconfigure((0, 1, 2, 3, 4), weight=1)
        
        avg_scores = [
            ("MCQ", f"{avg['mcq']}%", ACCENT),
            ("Technical", f"{avg['technical']}%", TEAL),
            ("Coding", f"{avg['coding']}%", AMBER),
            ("HR", f"{avg['hr']}%", PINK),
            ("Overall", f"{avg['overall']}%", AMBER),
        ]
        for idx, (lbl, val, color) in enumerate(avg_scores):
            box = ctk.CTkFrame(score_grid, fg_color=color, corner_radius=6)
            box.grid(row=0, column=idx, padx=4, sticky="ew")
            ctk.CTkLabel(box, text=lbl, font=F("xs", "bold"), text_color="white").pack(pady=(6, 2))
            ctk.CTkLabel(box, text=val, font=F("sm", "bold"), text_color="white").pack(pady=(0, 6))

        # Panel B: Role Performance
        role_card, role_body = self._card_frame(left_col, padx=15, pady=15)
        role_card.pack(fill="x", pady=(0, 15))
        ctk.CTkLabel(role_body, text="Role Wise Performance Metrics", font=F("md", "bold")).pack(anchor="w", pady=(0, 8))
        
        _style_tree()
        role_tree = ttk.Treeview(
            role_body, columns=("role", "attempts", "score"),
            show="headings", height=6)
        role_tree.heading("role",     text="Role")
        role_tree.heading("attempts", text="Attempts")
        role_tree.heading("score",    text="Avg Score")
        role_tree.column("role",     width=180)
        role_tree.column("attempts", width=90, anchor="center")
        role_tree.column("score",    width=90, anchor="center")
        role_tree.pack(fill="x")
        for row in role_perf:
            role_tree.insert("", "end", values=(
                row["role_name"], row["total_attempts"],
                f"{round(row['avg_score'] or 0.0, 2)}%"))

        # Panel C: Question Count Per Role
        qcount_card, qcount_body = self._card_frame(left_col, padx=15, pady=15)
        qcount_card.pack(fill="x", pady=(0, 15))
        ctk.CTkLabel(qcount_body, text="Question Count Per Role", font=F("md", "bold")).pack(anchor="w", pady=(0, 8))
        
        qcount_tree = ttk.Treeview(
            qcount_body, columns=("role", "count"),
            show="headings", height=5)
        qcount_tree.heading("role",  text="Job Role")
        qcount_tree.heading("count", text="Total Questions")
        qcount_tree.column("role",  width=250)
        qcount_tree.column("count", width=110, anchor="center")
        qcount_tree.pack(fill="x")
        for row in q_per_role:
            qcount_tree.insert("", "end", values=(row["role_name"], row["question_count"]))
        if not q_per_role:
            qcount_tree.insert("", "end", values=("No roles created yet", "0"))

        # Panel D: Most Failed Roles
        fail_card, fail_body = self._card_frame(left_col, padx=15, pady=15)
        fail_card.pack(fill="x", pady=(0, 15))
        ctk.CTkLabel(fail_body, text="Most Failed Roles", font=F("md", "bold")).pack(anchor="w", pady=(0, 8))
        
        fail_tree = ttk.Treeview(
            fail_body, columns=("role", "fails"),
            show="headings", height=5)
        fail_tree.heading("role",  text="Job Role")
        fail_tree.heading("fails", text="Total Fails")
        fail_tree.column("role",  width=250)
        fail_tree.column("fails", width=110, anchor="center")
        fail_tree.pack(fill="x")
        for row in failed_roles:
            fail_tree.insert("", "end", values=(row["role_name"], row["fail_count"]))
        if not failed_roles:
            fail_tree.insert("", "end", values=("No failed interviews recorded", "0"))

        # ------------------- RIGHT COLUMN PANELS -------------------
        
        # Panel E: Recent Interviews
        recent_card, recent_body = self._card_frame(right_col, padx=15, pady=15)
        recent_card.pack(fill="x", pady=(0, 15))
        ctk.CTkLabel(recent_body, text="Recent Interview Attempts", font=F("md", "bold")).pack(anchor="w", pady=(0, 8))
        
        recent_tree = ttk.Treeview(
            recent_body, columns=("name", "role", "score", "date"),
            show="headings", height=6)
        recent_tree.heading("name",  text="Student")
        recent_tree.heading("role",  text="Role")
        recent_tree.heading("score", text="Score")
        recent_tree.heading("date",  text="Date")
        recent_tree.column("name",  width=120)
        recent_tree.column("role",  width=140)
        recent_tree.column("score", width=60, anchor="center")
        recent_tree.column("date",  width=100, anchor="center")
        recent_tree.pack(fill="x")
        for row in recent:
            recent_tree.insert("", "end", values=(
                row["name"], row["role_name"],
                f"{row['overall_score']}%", row["created_at"][:10]))

        # Panel F: Top Students
        top_card, top_body = self._card_frame(right_col, padx=15, pady=15)
        top_card.pack(fill="x", pady=(0, 15))
        ctk.CTkLabel(top_body, text="🏆 Top Students", font=F("md", "bold")).pack(anchor="w", pady=(0, 8))
        
        top_tree = ttk.Treeview(
            top_body, columns=("rank", "name", "score"),
            show="headings", height=5)
        top_tree.heading("rank",  text="Rank")
        top_tree.heading("name",  text="Student Name")
        top_tree.heading("score", text="Average Score")
        top_tree.column("rank",  width=60, anchor="center")
        top_tree.column("name",  width=200)
        top_tree.column("score", width=120, anchor="center")
        top_tree.pack(fill="x")
        for idx, row in enumerate(top_stud, 1):
            top_tree.insert("", "end", values=(f"#{idx}", row["name"], f"{round(row['avg_score'], 2)}%"))
        if not top_stud:
            top_tree.insert("", "end", values=("—", "No scores recorded yet", "—"))

        # Panel G: Recent Activities
        act_card, act_body = self._card_frame(right_col, padx=15, pady=15)
        act_card.pack(fill="x", pady=(0, 15))
        ctk.CTkLabel(act_body, text="Recent Activities Log", font=F("md", "bold")).pack(anchor="w", pady=(0, 8))
        
        act_tree = ttk.Treeview(
            act_body, columns=("timestamp", "activity"),
            show="headings", height=8)
        act_tree.heading("timestamp", text="Timestamp")
        act_tree.heading("activity",  text="Activity Log Description")
        act_tree.column("timestamp", width=130, anchor="center")
        act_tree.column("activity",  width=270)
        act_tree.pack(fill="x")
        for act in activities:
            act_tree.insert("", "end", values=(act["created_at"], act["activity_text"]))
        if not activities:
            act_tree.insert("", "end", values=("—", "No activities logged yet"))

    def _clear(self):
        for w in self.winfo_children(): w.destroy()


#---------------------------------------------------------------------------
# Seed & entry point
#---------------------------------------------------------------------------
def _seed_questions_for_role(role_id, role_name):
    import database
    conn = database.get_connection()
    cur = conn.cursor()
    cur.execute("SELECT COUNT(*) FROM questions WHERE role_id = ?", (role_id,))
    if cur.fetchone()[0] > 0:
        conn.close()
        return

    if role_name == "Software Developer":
        admin.add_mcq_question(
            role_id, "Easy",
            "What does OOP stand for?",
            "Object Oriented Programming", "Open Object Protocol",
            "Operational Output Process", "Ordered Operand Pointer", "A")
        admin.add_mcq_question(
            role_id, "Easy",
            "Which keyword defines a function in Python?",
            "func", "def", "function", "fn", "B")
        admin.add_technical_question(
            role_id, "Easy",
            "Explain the concept of inheritance in OOP.",
            "inheritance,class,object,parent,child")
        admin.add_hr_question(
            role_id,
            "Tell us about a time you worked in a team.",
            "teamwork,collaboration,communication,goals,learning")
    elif role_name == "Web Developer":
        admin.add_mcq_question(
            role_id, "Easy",
            "Which language runs in the browser?",
            "Python", "JavaScript", "Java", "C++", "B")
        admin.add_technical_question(
            role_id, "Easy",
            "What is the difference between HTML and CSS?",
            "structure,style,markup,presentation,layout")
        admin.add_hr_question(
            role_id,
            "How do you handle tight deadlines on a project?",
            "prioritize,communication,deadline,planning,team")
    elif role_name == "Data Analyst":
        admin.add_mcq_question(
            role_id, "Easy",
            "Which tool is commonly used for data visualization?",
            "Git", "Tableau", "Docker", "Nginx", "B")
        admin.add_mcq_question(
            role_id, "Easy",
            "What does SQL primarily manage?",
            "Images", "Relational data", "Network packets", "Audio", "B")
        admin.add_technical_question(
            role_id, "Easy",
            "What is the difference between mean and median?",
            "average,outlier,central,skew,distribution")
        admin.add_hr_question(
            role_id,
            "How do you explain insights to non-technical stakeholders?",
            "communication,simple,visual,story,business")
    conn.close()


def _seed_coding_questions_for_role(role_id, role_name):
    if role_name == "Software Developer":
        # Easy
        admin.add_coding_question(role_id, "Easy", 
            "What will be the output of this Python code?", 
            "x = [1, 2, 3]\ny = x\ny.append(4)\nprint(len(x))",
            "3", "4", "Error", "5", "B")
        admin.add_coding_question(role_id, "Easy", 
            "What is the output of the following function call?", 
            "def f(a, b=[]):\n    b.append(a)\n    return b\nprint(f(1) + f(2))",
            "[1, 2]", "[1, 1, 2]", "[1, 1, 2, 2]", "[1, 2, 1, 2]", "D")
        admin.add_coding_question(role_id, "Easy", 
            "What does this division return in Python 3?", 
            "print(type(1 / 2))",
            "<class 'int'>", "<class 'float'>", "<class 'double'>", "<class 'number'>", "B")
        admin.add_coding_question(role_id, "Easy", 
            "What happens when executing this tuple assignment?", 
            "a = (1, 2, 3)\na[0] = 4\nprint(a)",
            "(4, 2, 3)", "TypeError", "[4, 2, 3]", "ValueError", "B")
        admin.add_coding_question(role_id, "Easy", 
            "What is the output of this slice operation?", 
            "print(\"Python\"[::-1])",
            "nohtyP", "Python", "Pytho", "PythonP", "A")

        # Medium
        admin.add_coding_question(role_id, "Medium", 
            "What is the output when key is missing in dictionary get method?", 
            "d = {1: 'a', 2: 'b'}\nprint(d.get(3, 'c'))",
            "None", "c", "KeyError", "b", "B")
        admin.add_coding_question(role_id, "Medium", 
            "What will this boolean expression evaluate to?", 
            "print(all([]) and any([]))",
            "True", "False", "Error", "None", "B")
        admin.add_coding_question(role_id, "Medium", 
            "What is the output when f() modifies the global variable x?", 
            "x = 10\ndef f():\n    global x\n    x = 20\nf()\nprint(x)",
            "10", "20", "UnboundLocalError", "None", "B")
        admin.add_coding_question(role_id, "Medium", 
            "What is the output of this map and lambda snippet?", 
            "print(list(map(lambda x: x*2, [1, 2]))[1])",
            "2", "4", "1", "Error", "B")
        admin.add_coding_question(role_id, "Medium", 
            "What does find method return when substring is not found?", 
            "x = \"Hello\"\nprint(x.find(\"z\"))",
            "None", "-1", "ValueError", "0", "B")

        # Hard
        admin.add_coding_question(role_id, "Hard", 
            "What is the output when consuming this generator?", 
            "def g():\n    yield 1\n    yield 2\nx = g()\nprint(next(x) + next(x))",
            "3", "12", "TypeError", "StopIteration", "A")
        admin.add_coding_question(role_id, "Hard", 
            "What will be printed after appending to a nested multiplied list?", 
            "a = [[]] * 3\na[0].append(1)\nprint(a)",
            "[[1], [], []]", "[[1], [1], [1]]", "[[1]]", "Error", "B")
        admin.add_coding_question(role_id, "Hard", 
            "What is the output of this sum operation with start value?", 
            "print(sum([1, 2, 3], 10))",
            "6", "16", "TypeError", "13", "B")
        admin.add_coding_question(role_id, "Hard", 
            "What is the length of set x after adding a duplicate?", 
            "x = {1, 2, 3}\nx.add(1)\nprint(len(x))",
            "4", "3", "2", "Error", "B")
        admin.add_coding_question(role_id, "Hard", 
            "Is True an instance of int in Python?", 
            "print(isinstance(True, int))",
            "True", "False", "TypeError", "None", "A")

    elif role_name == "Web Developer":
        # Easy
        admin.add_coding_question(role_id, "Easy", 
            "What is the output of typeof NaN in JavaScript?", 
            "console.log(typeof NaN);",
            "\"number\"", "\"NaN\"", "\"undefined\"", "\"object\"", "A")
        admin.add_coding_question(role_id, "Easy", 
            "What is the output of addition with a string operand?", 
            "console.log(1 + \"2\");",
            "3", "\"12\"", "NaN", "TypeError", "B")
        admin.add_coding_question(role_id, "Easy", 
            "What is the comparison output of spread array copies?", 
            "let a = [1, 2, 3];\nlet b = [...a];\nconsole.log(a === b);",
            "true", "false", "undefined", "TypeError", "B")
        admin.add_coding_question(role_id, "Easy", 
            "What is the loose equality comparison in JS?", 
            "console.log(0 == false);",
            "true", "false", "undefined", "null", "A")
        admin.add_coding_question(role_id, "Easy", 
            "What does this comparison return in JavaScript?", 
            "console.log(3 > 2 > 1);",
            "true", "false", "NaN", "Error", "B")

        # Medium
        admin.add_coding_question(role_id, "Medium", 
            "What is the nested typeof output in JavaScript?", 
            "console.log(typeof typeof 1);",
            "\"number\"", "\"string\"", "\"object\"", "\"undefined\"", "B")
        admin.add_coding_question(role_id, "Medium", 
            "What happens when object is used as a key in another object?", 
            "const a = {};\nconst b = { key: 'b' };\nconst c = { key: 'c' };\na[b] = 123;\na[c] = 456;\nconsole.log(a[b]);",
            "123", "456", "undefined", "ReferenceError", "B")
        admin.add_coding_question(role_id, "Medium", 
            "What is the output of accessing block-scoped variable before declaration?", 
            "let x = 1;\nfunction f() {\n  console.log(x);\n  let x = 2;\n}\nf();",
            "1", "2", "ReferenceError", "undefined", "C")
        admin.add_coding_question(role_id, "Medium", 
            "What is the output of array addition in JavaScript?", 
            "console.log([] + []);",
            "\"\"", "[]", "undefined", "NaN", "A")
        admin.add_coding_question(role_id, "Medium", 
            "What is the output of parseInt with units?", 
            "console.log(parseInt(\"16px\"));",
            "16", "NaN", "16px", "Error", "A")

        # Hard
        admin.add_coding_question(role_id, "Hard", 
            "What is the output sequence of microtasks and synchronous code?", 
            "const p = new Promise((res) => {\n  res(\"done\");\n});\np.then(console.log);\nconsole.log(\"start\");",
            "done start", "start done", "start", "done", "B")
        admin.add_coding_question(role_id, "Hard", 
            "What is the comparison output of empty array and falsey value?", 
            "console.log([] == ![]);",
            "true", "false", "TypeError", "undefined", "A")
        admin.add_coding_question(role_id, "Hard", 
            "What is the output of typeof null in JavaScript?", 
            "console.log(typeof null);",
            "\"null\"", "\"object\"", "\"undefined\"", "\"value\"", "B")
        admin.add_coding_question(role_id, "Hard", 
            "What is the length of array after assigning out of bounds index?", 
            "const a = [1, 2, 3];\na[10] = 99;\nconsole.log(a.length);",
            "3", "11", "10", "4", "B")
        admin.add_coding_question(role_id, "Hard", 
            "What is the output of this immediate invoked closure?", 
            "console.log((() => {\n  let a = 1;\n  return () => ++a;\n})()());",
            "1", "2", "3", "undefined", "B")


def seed_sample_data():
    import database
    conn = database.get_connection()
    cur  = conn.cursor()
    cur.execute("SELECT id FROM admins WHERE username = ?", ("admin",))
    if not cur.fetchone():
        auth.create_admin("admin", "admin123")
    cur.execute("SELECT COUNT(*) FROM roles")
    if cur.fetchone()[0] == 0:
        for role in ("Software Developer", "Web Developer", "Data Analyst"):
            has_coding = 1 if role in ("Software Developer", "Web Developer") else 0
            admin.add_role(role, has_coding)
    else:
        cur.execute("UPDATE roles SET has_coding_round = 1 WHERE role_name IN ('Software Developer', 'Web Developer')")
        conn.commit()

    # Seed coding questions helper
    for role in admin.fetch_roles():
        if role["has_coding_round"] == 1:
            cur.execute("SELECT COUNT(*) FROM questions WHERE role_id = ? AND question_type = 'coding'", (role["id"],))
            if cur.fetchone()[0] == 0:
                _seed_coding_questions_for_role(role["id"], role["role_name"])

    cur.execute("SELECT COUNT(*) FROM questions")
    if cur.fetchone()[0] == 0:
        roles = {r["role_name"]: r["id"] for r in admin.fetch_roles()}
        dev  = roles.get("Software Developer")
        web  = roles.get("Web Developer")
        data = roles.get("Data Analyst")

        if dev:
            admin.add_mcq_question(
                dev, "Easy", "What does OOP stand for?",
                "Object Oriented Programming", "Open Object Protocol",
                "Operational Output Process", "Ordered Operand Pointer", "A")
            admin.add_mcq_question(
                dev, "Easy", "Which keyword defines a function in Python?",
                "func", "def", "function", "fn", "B")
            admin.add_mcq_question(
                dev, "Medium", "Which data structure uses LIFO?",
                "Queue", "Stack", "Tree", "Graph", "B")
            admin.add_technical_question(
                dev, "Easy", "Explain the concept of inheritance in OOP.",
                "inheritance,class,object,parent,child")
            admin.add_technical_question(
                dev, "Medium", "What is polymorphism? Give an example.",
                "polymorphism,method,override,interface,class")
            admin.add_hr_question(
                dev, "Tell us about a time you worked in a team.",
                "teamwork,collaboration,communication,goals,learning")
            admin.add_hr_question(
                dev, "Where do you see yourself in five years?",
                "growth,learning,skills,goals,career")

        if web:
            admin.add_mcq_question(
                web, "Easy", "Which language runs in the browser?",
                "Python", "JavaScript", "Java", "C++", "B")
            admin.add_mcq_question(
                web, "Easy", "What does CSS stand for?",
                "Computer Style Sheets", "Cascading Style Sheets",
                "Creative Style System", "Colorful Style Sheets", "B")
            admin.add_technical_question(
                web, "Easy", "What is the difference between HTML and CSS?",
                "structure,style,markup,presentation,layout")
            admin.add_technical_question(
                web, "Medium", "Explain responsive web design.",
                "responsive,mobile,media,queries,layout")
            admin.add_hr_question(
                web, "How do you handle tight deadlines on a project?",
                "prioritize,communication,deadline,planning,team")
            admin.add_hr_question(
                web, "Describe a bug you fixed recently.",
                "debugging,testing,problem,solution,learning")

        if data:
            admin.add_mcq_question(
                data, "Easy", "Which tool is commonly used for data visualization?",
                "Git", "Tableau", "Docker", "Nginx", "B")
            admin.add_mcq_question(
                data, "Easy", "What does SQL primarily manage?",
                "Images", "Relational data", "Network packets", "Audio", "B")
            admin.add_mcq_question(
                data, "Medium", "Which measure shows data spread around the mean?",
                "Median", "Mode", "Standard deviation", "Count", "C")
            admin.add_technical_question(
                data, "Easy", "What is the difference between mean and median?",
                "average,outlier,central,skew,distribution")
            admin.add_technical_question(
                data, "Medium", "When would you use a JOIN in SQL?",
                "tables,relationship,keys,query,combine")
            admin.add_hr_question(
                data, "How do you explain insights to non-technical stakeholders?",
                "communication,simple,visual,story,business")
            admin.add_hr_question(
                data, "Tell us about a data project you are proud of.",
                "analysis,impact,metrics,learning,results")

    for role in admin.fetch_roles():
        _seed_questions_for_role(role["id"], role["role_name"])

    conn.close()


def main():
    ctk.set_appearance_mode("dark")
    ctk.set_default_color_theme("blue")
    init_db()
    seed_sample_data()
    app = InterviewApp()
    app.mainloop()


if __name__ == "__main__":
    main()

